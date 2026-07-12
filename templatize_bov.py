"""Convert the polished BOV Report (CLIENT) (1).docx into a reusable template.

This preserves the original design 100% (cover text boxes, blue bars, tables,
comparable layout) and only swaps the concrete sample values for {{placeholders}}
that both4.py fills at generation time.

Run:  python templatize_bov.py
Output: bov_template.docx
"""
from docx import Document
from docx.oxml.ns import qn

SOURCE = "BOV Report (CLIENT) (1).docx"
OUTPUT = "bov_template.docx"

W_T = qn("w:t")


def set_runs_text(paragraph, text):
    """Replace a paragraph's text with `text`, preserving the first run's format."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def set_cell(cell, text):
    """Set a table cell's value, preserving formatting of its first run."""
    paragraph = cell.paragraphs[0]
    set_runs_text(paragraph, text)
    # Clear any extra paragraphs in the cell
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


# ---- Text box replacements (exact run-text match inside w:txbxContent) ----
TEXTBOX_MAP = {
    "May 09, 2026": "{{Date}}",
    "Steele Oral Surgery": "{{property_name}}",
    "225 TX-121 Coppell TX 75019": "{{address}}",
    "Darrell Steele": "{{prepared_for}}",
    "Owner": "{{prepared_for_company}}",
    "Braden Fisher": "{{prepared_by}}",
    "VASCRE": "{{prepared_by_company}}",
    "123 Example Address": "{{prepared_by_address}}",
}


def replace_in_textboxes(doc, mapping):
    body = doc.element.body
    count = 0
    for txbx in body.iter(qn("w:txbxContent")):
        for t in txbx.iter(W_T):
            if t.text and t.text.strip() in mapping:
                t.text = mapping[t.text.strip()]
                count += 1
    return count


def preserve_prepared_label_spaces(doc):
    """Keep the trailing space after PREPARED BY:/FOR: (Word drops it without xml:space)."""
    fixed = 0
    for t in doc.element.body.iter(W_T):
        if not t.text:
            continue
        stripped = t.text.strip()
        if stripped in ("PREPARED BY:", "PREPARED FOR:"):
            t.text = stripped + " "
            t.set(qn("xml:space"), "preserve")
            fixed += 1
    return fixed


# ---- Table cell replacements by (table, row, col) ----
TABLE_MAP = {
    # Valuation summary (Table 1)
    (1, 1, 1): "{{market_price_psf}}/PSF X {{market_building_sf}}",
    (1, 1, 2): "{{market_value}}",
    (1, 2, 2): "{{market_value_rounded}}",
    (1, 5, 0): "{{value_aggressive}}",
    (1, 5, 1): "{{market_value}}",
    (1, 5, 2): "{{value_conservative}}",
    # General Information (Table 2)
    (2, 1, 1): "{{property_name}}",
    (2, 2, 1): "{{property_type}}",
    (2, 3, 1): "{{state}}",
    (2, 4, 1): "{{county}}",
    (2, 5, 1): "{{longitude}}",
    (2, 6, 1): "{{latitude}}",
    (2, 8, 1): "{{Topography}}",
    (2, 9, 1): "{{shape}}",
    (2, 10, 1): "{{Access}}",
    (2, 11, 1): "{{Exposure}}",
    (2, 13, 1): "{{lot_area}}",
    (2, 14, 1): "{{acres}}",
    (2, 15, 1): "{{recorded_sale_date}}",
    (2, 16, 1): "{{zoning}}",
    (2, 17, 1): "{{apn}}",
    (2, 18, 1): "{{current_owner}}",
    # Population (Table 3)
    (3, 2, 1): "{{pop_2010_us}}", (3, 2, 2): "{{pop_2010_state}}", (3, 2, 3): "{{pop_2010_county}}",
    (3, 3, 1): "{{pop_2020_us}}", (3, 3, 2): "{{pop_2020_state}}", (3, 3, 3): "{{pop_2020_county}}",
    (3, 4, 1): "{{pop_2025_us}}", (3, 4, 2): "{{pop_2025_state}}", (3, 4, 3): "{{pop_2025_county}}",
    # Population density (Table 4)
    (4, 2, 1): "{{density_2020_us}}", (4, 2, 2): "{{density_2020_state}}", (4, 2, 3): "{{density_2020_county}}",
    (4, 3, 1): "{{density_2025_us}}", (4, 3, 2): "{{density_2025_state}}", (4, 3, 3): "{{density_2025_county}}",
    # Households (Table 5)
    (5, 2, 1): "{{hh_2024_us}}", (5, 2, 2): "{{hh_2024_state}}", (5, 2, 3): "{{hh_2024_county}}",
    (5, 3, 1): "{{hh_2029_us}}", (5, 3, 2): "{{hh_2029_state}}", (5, 3, 3): "{{hh_2029_county}}",
    (5, 4, 1): "{{hh_cagr_us}}", (5, 4, 2): "{{hh_cagr_state}}", (5, 4, 3): "{{hh_cagr_county}}",
    # Average household size (Table 6)
    (6, 2, 1): "{{hhsize_2024_us}}", (6, 2, 2): "{{hhsize_2024_state}}", (6, 2, 3): "{{hhsize_2024_county}}",
    (6, 3, 1): "{{hhsize_2029_us}}", (6, 3, 2): "{{hhsize_2029_state}}", (6, 3, 3): "{{hhsize_2029_county}}",
    (6, 4, 1): "{{hhsize_cagr_us}}", (6, 4, 2): "{{hhsize_cagr_state}}", (6, 4, 3): "{{hhsize_cagr_county}}",
    # Housing tenure (Table 7)
    (7, 2, 1): "{{owner_us}}", (7, 2, 2): "{{owner_state}}", (7, 2, 3): "{{owner_county}}",
    (7, 3, 1): "{{renter_us}}", (7, 3, 2): "{{renter_state}}", (7, 3, 3): "{{renter_county}}",
    # Local Area Demographics rings (Table 8) - rows mapped to generator data
    (8, 5, 1): "{{r1_pop_2024}}", (8, 5, 2): "{{r3_pop_2024}}", (8, 5, 3): "{{r5_pop_2024}}",
    (8, 6, 1): "{{r1_pop_2029}}", (8, 6, 2): "{{r3_pop_2029}}", (8, 6, 3): "{{r5_pop_2029}}",
    (8, 13, 1): "{{r1_hh_2024}}", (8, 13, 2): "{{r3_hh_2024}}", (8, 13, 3): "{{r5_hh_2024}}",
    (8, 14, 1): "{{r1_hh_2029}}", (8, 14, 2): "{{r3_hh_2029}}", (8, 14, 3): "{{r5_hh_2029}}",
    (8, 19, 1): "{{r1_owner_pct}}", (8, 19, 2): "{{r3_owner_pct}}", (8, 19, 3): "{{r5_owner_pct}}",
    (8, 20, 1): "{{r1_renter_pct}}", (8, 20, 2): "{{r3_renter_pct}}", (8, 20, 3): "{{r5_renter_pct}}",
    (8, 3, 6): "{{r1_avg_hh_income}}", (8, 3, 7): "{{r3_avg_hh_income}}", (8, 3, 8): "{{r5_avg_hh_income}}",
    (8, 7, 6): "{{r1_median_hh_income}}", (8, 7, 7): "{{r3_median_hh_income}}", (8, 7, 8): "{{r5_median_hh_income}}",
    (8, 11, 6): "{{r1_per_capita_income}}", (8, 11, 7): "{{r3_per_capita_income}}", (8, 11, 8): "{{r5_per_capita_income}}",
    # Reconciliation (Table 10)
    (10, 1, 1): "{{market_value_rounded}}",
    (10, 1, 2): "{{reconciliation_summary}}",
    # Opinions of value (Table 11)
    (11, 1, 1): "{{market_price_psf}} SF",
    (11, 1, 3): "{{market_building_sf}}",
    (11, 1, 4): "{{market_value}}",
    (11, 2, 4): "{{market_value_rounded}}",
}


# ---- Narrative paragraph replacements by text prefix ----
# First match -> placeholder; later duplicates -> cleared (avoid repeating text)
NARRATIVE_MAP = [
    ("The property is located at 225 TX-121, Coppell, Texas, within Dallas County", "{{property_summary}}"),
    ("The subject property is located at 225 TX-121 in Coppell", "{{location_summary}}"),
    ("The retail property at 225 TX-121 in Coppell, Texas, benefits from a robust demographic", "{{demographic_analysis}}"),
    ("The retail property located at 225 TX-121, Coppell, Texas, is situated on relatively flat", "{{size_and_topography}}"),
    ("Coppell, located in Dallas County, Texas, has experienced consistent population", "{{population_analysis}}"),
    ("The area surrounding 225 TX-121 in Coppell, Texas, exhibits a balanced owner/renter", "{{household_trends}}"),
    ("Coppell, Texas, located in Dallas County, benefits from a robust labor market", "{{employment_analysis}}"),
    ("The retail property at 225 TX-121 in Coppell, Texas, benefits significantly from its strategic location", "{{economic_factors}}"),
    ("The property at 225 TX-121 in Coppell, Texas, benefits from strong community services", "{{community_services}}"),
    ("Dallas County  Retail Market Report", "{{county}} {{property_type}} Market Report — {{market_quarter}}"),
    ("Dallas County sits within the DFW retail market", "{{market_overview}}"),
]


def replace_narratives(doc):
    seen = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for prefix, placeholder in NARRATIVE_MAP:
            if text.startswith(prefix):
                if placeholder in seen:
                    set_runs_text(paragraph, "")  # clear duplicate
                else:
                    set_runs_text(paragraph, placeholder)
                    seen.add(placeholder)
                break


def main():
    doc = Document(SOURCE)

    tb = replace_in_textboxes(doc, TEXTBOX_MAP)
    labels = preserve_prepared_label_spaces(doc)

    cells = 0
    for (ti, ri, ci), placeholder in TABLE_MAP.items():
        try:
            set_cell(doc.tables[ti].rows[ri].cells[ci], placeholder)
            cells += 1
        except IndexError:
            print(f"  WARN: cell {(ti, ri, ci)} out of range")

    replace_narratives(doc)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}: {tb} textbox values, {labels} prepared-label spaces, {cells} table cells, narratives parameterized")
    print(f"Tables preserved: {len(doc.tables)} | paragraphs: {len(doc.paragraphs)}")

    # Remove any remaining static market/comps sample blocks
    from fix_bov_template_static import fix_template
    fix_template(OUTPUT)


if __name__ == "__main__":
    main()
