"""Generate BOV-style Word templates with placeholders and professional styling.

Produces two templates:
  - bov_template.docx          (BOV Client: full 20-25 page report)
  - bov_prospect_template.docx (BOV Prospect: condensed 5-7 page report)

Cells and paragraphs contain {{placeholders}} that both4.py fills at generation
time. Run:  python build_bov_template.py
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# Brand palette (matches the app accent)
ACCENT = RGBColor(0x0D, 0xB2, 0xFF)
ACCENT_HEX = "0DB2FF"
DARK = RGBColor(0x1F, 0x29, 0x37)
LIGHT_HEX = "EAF7FF"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x55, 0x5F, 0x6D)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def style_styles(doc):
    """Apply document-wide font and heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11.5, GREY),
    ]:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = color
            st.font.bold = True


def add_accent_rule(paragraph):
    """Add a bottom accent border to a paragraph (used under H1)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), ACCENT_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        add_accent_rule(h)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(6)
    return h


def add_body(doc, text):
    names = [s.name for s in doc.styles]
    p = doc.add_paragraph(text)
    if "Body Text" in names:
        p.style = doc.styles["Body Text"]
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width


def make_table(doc, header, rows, widths=None, label_col=True):
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_background(cell, ACCENT_HEX)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    # Data rows with zebra striping
    for r_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row_values):
            cells[i].text = text
            if r_index % 2 == 1:
                set_cell_background(cells[i], LIGHT_HEX)
            for paragraph in cells[i].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and label_col) else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    if i == 0 and label_col:
                        run.bold = True

    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_source(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)


def add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BROKER OPINION OF VALUE")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("{{property_name}}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK

    for text, size in [("{{property_type}}", 13), ("{{county}}, {{state}}", 13), ("{{Date}}", 12)]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph()

    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep.add_run("PREPARED BY\n").bold = True
    prep.add_run("{{prepared_by}}\n{{prepared_by_company}}\n{{prepared_by_address}}")

    doc.add_paragraph()
    prep_for = doc.add_paragraph()
    prep_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep_for.add_run("PREPARED FOR\n").bold = True
    prep_for.add_run("{{prepared_for}}\n{{prepared_for_company}}\n{{prepared_for_address}}")
    doc.add_page_break()


def general_info_table(doc):
    make_table(
        doc,
        ["General Information", ""],
        [
            ["Property Name", "{{property_name}}"],
            ["Property Type", "{{property_type}}"],
            ["State", "{{state}}"],
            ["County", "{{county}}"],
            ["Longitude", "{{longitude}}"],
            ["Latitude", "{{latitude}}"],
            ["Topography", "{{Topography}}"],
            ["Shape", "{{shape}}"],
            ["Access", "{{Access}}"],
            ["Exposure", "{{Exposure}}"],
            ["Lot Area", "{{lot_area}}"],
            ["Acres", "{{acres}}"],
            ["Recorded Sale Date", "{{recorded_sale_date}}"],
            ["Zoning", "{{zoning}}"],
            ["APN", "{{apn}}"],
            ["Current Owner", "{{current_owner}}"],
            ["Marketing Period", "{{marketing_period}}"],
        ],
        widths=[Inches(2.6), Inches(3.4)],
    )


def valuation_tables(doc):
    make_table(
        doc,
        ["Aggressive", "Market Value", "Conservative"],
        [["{{value_aggressive}}", "{{market_value}}", "{{value_conservative}}"]],
        widths=[Inches(2.0), Inches(2.0), Inches(2.0)],
        label_col=False,
    )
    make_table(
        doc,
        ["Metric", "Calculation", "Value"],
        [
            ["Market Sale Price", "{{market_price_psf}}/SF X {{market_building_sf}}", "{{market_value}}"],
            ["Market Sales Price (Rounded)", "", "{{market_value_rounded}}"],
        ],
        widths=[Inches(2.4), Inches(2.4), Inches(1.8)],
    )


def population_tables(doc):
    g3 = [Inches(2.7), Inches(1.1), Inches(1.1), Inches(1.1)]
    make_table(
        doc,
        ["Population", "U.S.", "State", "County"],
        [
            ["2010 Total Population", "{{pop_2010_us}}", "{{pop_2010_state}}", "{{pop_2010_county}}"],
            ["2020 Total Population", "{{pop_2020_us}}", "{{pop_2020_state}}", "{{pop_2020_county}}"],
            ["2025 Total Population", "{{pop_2025_us}}", "{{pop_2025_state}}", "{{pop_2025_county}}"],
        ],
        widths=g3,
    )
    make_table(
        doc,
        ["Population Density (per sq mi)", "U.S.", "State", "County"],
        [
            ["2020", "{{density_2020_us}}", "{{density_2020_state}}", "{{density_2020_county}}"],
            ["2025", "{{density_2025_us}}", "{{density_2025_state}}", "{{density_2025_county}}"],
        ],
        widths=g3,
    )
    add_source(doc, "Source: Esri GeoEnrichment")


def household_tables(doc):
    g3 = [Inches(2.7), Inches(1.1), Inches(1.1), Inches(1.1)]
    make_table(
        doc,
        ["Number of Households", "U.S.", "State", "County"],
        [
            ["2024 Households", "{{hh_2024_us}}", "{{hh_2024_state}}", "{{hh_2024_county}}"],
            ["2029 Households", "{{hh_2029_us}}", "{{hh_2029_state}}", "{{hh_2029_county}}"],
            ["CAGR", "{{hh_cagr_us}}", "{{hh_cagr_state}}", "{{hh_cagr_county}}"],
        ],
        widths=g3,
    )
    make_table(
        doc,
        ["Average Household Size", "U.S.", "State", "County"],
        [
            ["2024", "{{hhsize_2024_us}}", "{{hhsize_2024_state}}", "{{hhsize_2024_county}}"],
            ["2029", "{{hhsize_2029_us}}", "{{hhsize_2029_state}}", "{{hhsize_2029_county}}"],
            ["2024-2029 CAGR", "{{hhsize_cagr_us}}", "{{hhsize_cagr_state}}", "{{hhsize_cagr_county}}"],
        ],
        widths=g3,
    )
    make_table(
        doc,
        ["Housing Units", "U.S.", "State", "County"],
        [
            ["Owner Occupied", "{{owner_us}}", "{{owner_state}}", "{{owner_county}}"],
            ["Renter Occupied", "{{renter_us}}", "{{renter_state}}", "{{renter_county}}"],
        ],
        widths=g3,
    )
    add_source(doc, "Source: Esri GeoEnrichment")


def ring_table(doc):
    widths = [Inches(2.7), Inches(1.1), Inches(1.1), Inches(1.1)]
    make_table(
        doc,
        ["Description", "1 mile", "3 miles", "5 miles"],
        [
            ["2024 Population", "{{r1_pop_2024}}", "{{r3_pop_2024}}", "{{r5_pop_2024}}"],
            ["2029 Population", "{{r1_pop_2029}}", "{{r3_pop_2029}}", "{{r5_pop_2029}}"],
            ["2024 Households", "{{r1_hh_2024}}", "{{r3_hh_2024}}", "{{r5_hh_2024}}"],
            ["2029 Households", "{{r1_hh_2029}}", "{{r3_hh_2029}}", "{{r5_hh_2029}}"],
            ["Avg Household Income", "{{r1_avg_hh_income}}", "{{r3_avg_hh_income}}", "{{r5_avg_hh_income}}"],
            ["Median Household Income", "{{r1_median_hh_income}}", "{{r3_median_hh_income}}", "{{r5_median_hh_income}}"],
            ["Per Capita Income", "{{r1_per_capita_income}}", "{{r3_per_capita_income}}", "{{r5_per_capita_income}}"],
            ["Owner Occupied", "{{r1_owner_pct}}", "{{r3_owner_pct}}", "{{r5_owner_pct}}"],
            ["Renter Occupied", "{{r1_renter_pct}}", "{{r3_renter_pct}}", "{{r5_renter_pct}}"],
        ],
        widths=widths,
    )
    add_source(doc, "Source: {{demographics_source}}")


def certification(doc):
    add_heading(doc, "CERTIFICATION AND DISCLAIMERS", 1)
    add_body(
        doc,
        "This Broker Opinion of Value (BOV) has been prepared for informational purposes only "
        "and is based on the broker's professional knowledge of the market, publicly available "
        "data, and the comparable sales analyzed herein.",
    )
    for clause in [
        "Not an Appraisal — This valuation is an estimate of market value and does not conform to USPAP.",
        "Assumptions & Limitations — The opinion is based on information deemed reliable but not guaranteed.",
        "No Legal or Financial Liability — The broker assumes no liability for errors, omissions, or reliance on this opinion.",
        "Market Conditions — This valuation reflects data available as of the preparation date.",
        "Confidentiality — This BOV is confidential and intended solely for the recipient.",
    ]:
        add_body(doc, clause)
    doc.add_paragraph()
    for line in ["Broker Signature: ____________________", "Prepared By: {{prepared_by}}", "Position:", "Phone:", "Email:"]:
        doc.add_paragraph(line)


def build_client():
    """Full 20-25 page BOV (Client) report."""
    doc = Document()
    style_styles(doc)
    add_page_numbers(doc)
    cover(doc)

    add_heading(doc, "TABLE OF CONTENTS", 1)
    make_table(
        doc,
        ["Section", "Page"],
        [
            ["Executive Summary", "3"],
            ["Property & Location Summary", "4"],
            ["Aerial Map & Subject Photos", "5"],
            ["Demographic Analysis", "6"],
            ["Employment & Economy", "10"],
            ["Market Analysis", "12"],
            ["Property Comparables", "15"],
            ["Reconciliation & Sales Conclusion", "18"],
            ["Certification & Disclaimers", "20"],
        ],
        widths=[Inches(5.0), Inches(1.0)],
    )
    doc.add_page_break()

    add_heading(doc, "EXECUTIVE SUMMARY", 1)
    add_body(doc, "{{executive_summary}}")
    add_heading(doc, "Opinion of Value", 2)
    valuation_tables(doc)
    add_heading(doc, "SWOT Summary", 2)
    make_table(doc, ["Strengths", "Weaknesses"], [["{{swot_strengths}}", "{{swot_weaknesses}}"]],
               widths=[Inches(3.0), Inches(3.0)], label_col=False)
    make_table(doc, ["Opportunities", "Threats"], [["{{swot_opportunities}}", "{{swot_threats}}"]],
               widths=[Inches(3.0), Inches(3.0)], label_col=False)
    doc.add_page_break()

    add_heading(doc, "PROPERTY SUMMARY", 1)
    add_body(doc, "{{property_summary}}")
    add_heading(doc, "General Information", 2)
    general_info_table(doc)

    add_heading(doc, "LOCATION SUMMARY", 1)
    add_body(doc, "{{location_summary}}")
    add_heading(doc, "AERIAL MAP", 2)
    doc.add_paragraph("{{aerial_map}}")
    add_heading(doc, "SUBJECT PHOTOS", 2)
    doc.add_paragraph("{{subject_photos}}")
    doc.add_page_break()

    add_heading(doc, "REGIONAL ANALYSIS", 1)
    add_body(doc, "{{regional_analysis}}")
    add_heading(doc, "DEMOGRAPHIC ANALYSIS", 1)
    add_heading(doc, "Overview", 2)
    add_body(doc, "{{demographic_analysis}}")
    add_heading(doc, "Size and Topography", 2)
    add_body(doc, "{{size_and_topography}}")
    add_heading(doc, "Population", 2)
    add_body(doc, "{{population_analysis}}")
    population_tables(doc)
    add_heading(doc, "Household Trends", 2)
    add_body(doc, "{{household_trends}}")
    household_tables(doc)
    add_heading(doc, "Local Area Demographics (Drive-Time Rings)", 2)
    add_body(doc, "Demographic profile within 1-, 3-, and 5-mile radii of the subject property.")
    ring_table(doc)
    doc.add_page_break()

    add_heading(doc, "EMPLOYMENT", 1)
    add_body(doc, "{{employment_analysis}}")
    make_table(
        doc,
        ["Employment & Unemployment", "U.S.", "State", "County"],
        [
            ["Total Employment", "{{emp_total_us}}", "{{emp_total_state}}", "{{emp_total_county}}"],
            ["Unemployment Rate", "{{unemp_us}}", "{{unemp_state}}", "{{unemp_county}}"],
        ],
        widths=[Inches(2.7), Inches(1.1), Inches(1.1), Inches(1.1)],
    )
    add_source(doc, "Source: BLS (Bureau of Labor Statistics)")
    add_heading(doc, "Economic Factors", 2)
    add_body(doc, "{{economic_factors}}")
    add_heading(doc, "Community Services", 2)
    add_body(doc, "{{community_services}}")
    doc.add_page_break()

    add_heading(doc, "MARKET ANALYSIS", 1)
    p = doc.add_paragraph()
    p.add_run("{{county}} {{property_type}} Market Report — {{market_quarter}}").bold = True
    doc.add_paragraph("Prepared by {{prepared_by_company}} — Independent Analysis Based on Multiple Public Data Sources").italic = True
    add_heading(doc, "1. Market Overview", 2)
    add_body(doc, "{{market_overview}}")
    add_heading(doc, "2. Key Market Metrics", 2)
    doc.add_paragraph("Vacancy Rates (Based on Public Data & Market Trends)").bold = True
    add_body(doc, "{{vacancy_rates}}")
    doc.add_paragraph("Lease Rates (Based on Public Listings & Industry Trends)").bold = True
    add_body(doc, "{{lease_rates}}")
    doc.add_paragraph("Construction Activity").bold = True
    add_body(doc, "{{construction_activity}}")
    add_heading(doc, "3. Trends & Forecast", 2)
    add_body(doc, "{{market_trends}}")
    add_heading(doc, "4. Investment Insights", 2)
    add_body(doc, "{{investment_insights}}")
    add_heading(doc, "5. Recommendations", 2)
    add_body(doc, "{{market_recommendations}}")
    add_heading(doc, "6. Data Sources & Disclaimer", 2)
    add_body(doc, "{{market_data_sources}}")
    doc.add_page_break()

    add_heading(doc, "PROPERTY COMPARABLES", 1)
    add_body(doc, "Comparable sales are summarized below. (Populated from the comparables module.)")
    make_table(
        doc,
        ["#", "Property", "Address", "Sale Price", "Price/SF"],
        [[str(i), "", "", "", ""] for i in range(1, 7)],
        widths=[Inches(0.5), Inches(1.8), Inches(2.2), Inches(1.0), Inches(1.0)],
        label_col=False,
    )
    add_heading(doc, "RECONCILIATION TABLE", 1)
    add_body(doc, "{{reconciliation_summary}}")
    make_table(
        doc,
        ["Valuation Method", "Value Estimate", "Notes"],
        [["Sales Comparison Approach", "{{market_value_rounded}}", "Primary approach for this property type."]],
        widths=[Inches(2.2), Inches(1.8), Inches(2.6)],
    )
    make_table(
        doc,
        ["Opinions of Value", "PSF", "X", "SF", "Value"],
        [
            ["Market Sale Price", "{{market_price_psf}}", "X", "{{market_building_sf}}", "{{market_value}}"],
            ["Market Sales Price (Rounded)", "", "", "", "{{market_value_rounded}}"],
        ],
        widths=[Inches(2.2), Inches(1.0), Inches(0.5), Inches(1.0), Inches(1.6)],
    )
    add_heading(doc, "SALES CONCLUSION", 1)
    add_body(doc, "{{sales_conclusion}}")
    doc.add_page_break()

    certification(doc)
    doc.save("bov_template.docx")
    print(f"Wrote bov_template.docx ({len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs)")


def build_prospect():
    """Condensed 5-7 page BOV (Prospect) report."""
    doc = Document()
    style_styles(doc)
    add_page_numbers(doc)
    cover(doc)

    add_heading(doc, "EXECUTIVE SUMMARY", 1)
    add_body(doc, "{{executive_summary}}")
    add_heading(doc, "Opinion of Value", 2)
    valuation_tables(doc)
    doc.add_page_break()

    add_heading(doc, "PROPERTY SUMMARY", 1)
    add_body(doc, "{{property_summary}}")
    add_heading(doc, "General Information", 2)
    general_info_table(doc)

    add_heading(doc, "LOCATION SUMMARY", 1)
    add_body(doc, "{{location_summary}}")
    add_heading(doc, "AERIAL MAP", 2)
    doc.add_paragraph("{{aerial_map}}")
    doc.add_page_break()

    add_heading(doc, "DEMOGRAPHIC SNAPSHOT", 1)
    add_body(doc, "{{demographic_analysis}}")
    population_tables(doc)
    add_heading(doc, "Local Area Demographics (Drive-Time Rings)", 2)
    ring_table(doc)
    doc.add_page_break()

    add_heading(doc, "MARKET OVERVIEW", 1)
    add_body(doc, "{{market_overview}}")
    add_heading(doc, "Recommendations", 2)
    add_body(doc, "{{market_recommendations}}")
    doc.add_page_break()

    certification(doc)
    doc.save("bov_prospect_template.docx")
    print(f"Wrote bov_prospect_template.docx ({len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs)")


if __name__ == "__main__":
    # NOTE: The canonical high-fidelity client template (bov_template.docx) is now
    # produced by templatize_bov.py from the original BOV report to preserve its
    # exact design. This script only builds the condensed prospect variant.
    build_prospect()
