"""Diagnose TOC bleed and cover broken-image issues."""
import zipfile
import re
from docx import Document
from docx.oxml.ns import qn

PATH = r"property_reports\BOV_River_Walk_Retail_Center_2026-07-11_2029.docx"
TMPL = "bov_template.docx"
ORIG = "BOV Report (CLIENT) (1).docx"

W_T = qn("w:t")
W_DRAWING = qn("w:drawing")


def dump_cover(path, label):
    print("\n" + "=" * 70)
    print(label, path)
    doc = Document(path)
    # textboxes
    for i, txbx in enumerate(doc.element.body.iter(qn("w:txbxContent"))):
        texts = [t.text.strip() for t in txbx.iter(W_T) if t.text and t.text.strip()]
        if texts:
            joined = " | ".join(texts)
            if any(k in joined for k in ["Table of Contents", "EXECUTIVE", "Prepared", "Date", "Commerce", "River", "{{"]):
                print(f"  TB{i}: {joined[:140]}")

    # early paragraphs
    print("  --- early paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:25]):
        t = p.text.strip()
        if not t:
            continue
        print(f"  P{i}: {t[:100]!r}")

    # TOC table
    if doc.tables:
        print("  --- Table 0 (TOC) ---")
        for ri, row in enumerate(doc.tables[0].rows):
            print("   ", [c.text.strip()[:40] for c in row.cells])

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
        # leftover placeholders
        phs = sorted(set(re.findall(r"\{\{[^}]+\}\}", xml)))
        print("  placeholders in XML:", phs[:20] if phs else "NONE")
        # broken image descrs
        for kw in ["main_img", "aerial", "Subject_photo", "Picture", "descr="]:
            if kw in ["main_img", "aerial", "Subject_photo"]:
                print(f"  contains {kw}:", kw in xml)
        # count drawings
        print("  drawing count:", xml.count("<w:drawing>"))
        # docPr descrs
        descrs = re.findall(r'descr="([^"]*)"', xml)
        print("  image descrs:", descrs[:15])
        # check for empty blips / missing relationships
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        print("  media files:", len(media))
        # relationships
        try:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8", "ignore")
            targets = re.findall(r'Target="([^"]+)"', rels)
            print("  rel targets sample:", targets[:8])
            broken = [t for t in targets if t.startswith("media/") and f"word/{t}" not in z.namelist() and f"word/{t.replace('%20',' ')}" not in [n for n in z.namelist()]]
            missing = []
            for t in targets:
                if "media/" in t:
                    path_m = "word/" + t.replace("\\", "/")
                    if path_m not in z.namelist():
                        missing.append(t)
            print("  missing media targets:", missing[:10] if missing else "NONE")
        except Exception as e:
            print("  rels err", e)


dump_cover(PATH, "GENERATED")
dump_cover(TMPL, "TEMPLATE")

# Compare executive summary paragraph location vs TOC
print("\n" + "=" * 70)
print("EXECUTIVE SUMMARY placement analysis")
doc = Document(PATH)
for i, p in enumerate(doc.paragraphs):
    if "EXECUTIVE SUMMARY" in p.text or "executive summary" in p.text.lower() or (p.text.strip().startswith("The retail property located at 849")):
        # check if inside textbox ancestor
        el = p._element
        in_txbx = False
        parent = el.getparent()
        while parent is not None:
            if "txbxContent" in (parent.tag or ""):
                in_txbx = True
                break
            parent = parent.getparent()
        print(f"  P{i} in_textbox={in_txbx} text={p.text.strip()[:80]!r}")
