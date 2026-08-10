import openai
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
import json
import requests
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
import logging
from pathlib import Path
import re
from location_services import HybridLocationService

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class PropertyReportData:
    """Complete data structure for property report"""
    # Basic Info
    address: str
    date: str = field(default_factory=lambda: datetime.now().strftime("%B %d, %Y"))
    
    # Preparer Info
    prepared_by: str = ""
    prepared_by_title: str = ""
    prepared_by_company: str = ""
    prepared_by_address: str = ""
    
    # Client Info  
    prepared_for: str = ""
    prepared_for_title: str = ""
    prepared_for_company: str = ""
    prepared_for_address: str = ""
    
    # Property Details
    property_name: str = ""
    property_type: str = ""
    state: str = ""
    county: str = ""
    longitude: str = ""
    latitude: str = ""
    
    # Physical Characteristics (mostly static)
    topography: str = "Level at street Grade"
    shape: str = "Irregular"
    access: str = "Good"  # Average, Average/Good, Good, Good/Excellent, Excellent
    exposure: str = "Average/Good"  # Average, Average/Good, Good, Good/Excellent, Excellent
    
    # Property Specific
    lot_area: str = ""
    acres: str = ""
    recorded_sale_date: str = ""
    zoning: str = ""
    apn: str = ""
    current_owner: str = ""
    
    # Transaction context
    lease_or_sale: str = ""

    # Market Analysis (static)
    marketing_period: str = "Six months or less"
    
    # SWOT Analysis
    swot_strengths: str = ""
    swot_weaknesses: str = ""
    swot_opportunities: str = ""
    swot_threats: str = ""
    
    # Generated Content
    property_summary: str = ""
    location_summary: str = ""
    demographic_analysis: str = ""
    size_and_topography: str = ""
    population_analysis: str = ""
    household_trends: str = ""
    housing_tenure: str = ""
    local_area_analysis: str = ""
    employment_analysis: str = ""
    economic_factors: str = ""
    community_services: str = ""
    
    # Market Analysis fields
    market_overview: str = ""
    vacancy_rates: str = ""
    lease_rates: str = ""
    construction_activity: str = ""
    market_trends: str = ""
    investment_insights: str = ""
    market_recommendations: str = ""
    market_data_sources: str = ""
    market_quarter: str = field(default_factory=lambda: f"Q{(datetime.now().month-1)//3 + 1} {datetime.now().year}")
    
    # Image paths
    aerial_image_path: Optional[str] = None
    street_view_image_path: Optional[str] = None

    # BOV table values (population, households, rings, employment, valuation, etc.)
    table_values: Dict[str, str] = field(default_factory=dict)

    # Additional BOV narrative sections
    executive_summary: str = ""
    regional_analysis: str = ""
    sales_conclusion: str = ""
    reconciliation_summary: str = ""
    reconciliation_notes: str = ""

    # Comparable sales extracted from uploaded CoStar PDF (optional)
    comps: List[Any] = field(default_factory=list)
    # Raw demographic dataset used for employment table refresh
    bov_dataset: Dict = field(default_factory=dict)

class ComprehensivePropertyReportGenerator:
    def __init__(
        self,
        openai_api_key: str,
        template_path: str,
        output_dir: str = "output",
        google_api_key: Optional[str] = None,
        esri_api_key: Optional[str] = None,
    ):
        """
        Initialize the Comprehensive Property Report Generator
        
        Args:
            openai_api_key: OpenAI API key
            template_path: Path to the Word document template
            output_dir: Directory to save generated reports
            google_api_key: Google Maps API key (Street View; optional fallback)
            esri_api_key: Esri API key (geocoding, aerial imagery, demographics)
        """
        self.openai_client = openai.OpenAI(api_key=openai_api_key)
        self.location_service = HybridLocationService(
            esri_api_key=esri_api_key,
            google_api_key=google_api_key,
        )
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Create images subdirectory
        self.images_dir = self.output_dir / "images"
        self.images_dir.mkdir(exist_ok=True)
        
        # Verify template exists
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

    def get_property_images(self, address: str, lat: float, lng: float) -> Tuple[Optional[str], Optional[str]]:
        """
        Get aerial and street view images for the property
        
        Args:
            address: The property address
            lat: Latitude
            lng: Longitude
            
        Returns:
            Tuple of (aerial_image_path, street_view_image_path)
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        aerial_path = None
        street_view_path = None
        
        try:
            logger.info("Fetching aerial image for %s @ (%.6f, %.6f)...", address, lat, lng)
            aerial_filename = f"aerial_{timestamp}.jpg"
            aerial_path = self.location_service.get_aerial_image(
                lat, lng, self.images_dir / aerial_filename
            )
            if aerial_path:
                logger.info("Aerial image saved: %s", aerial_path)

            logger.info("Fetching Street View for address %r @ (%.6f, %.6f)...", address, lat, lng)
            street_view_filename = f"street_view_{timestamp}.jpg"
            street_view_path = self.location_service.get_street_view_image(
                address,
                self.images_dir / street_view_filename,
                lat=lat,
                lng=lng,
            )
            if street_view_path:
                logger.info("Street view image saved: %s", street_view_path)
            else:
                logger.warning("No Street View image available; SUBJECT PHOTOS will be empty")
                
        except Exception as e:
            logger.error(f"Error fetching images: {e}")
            
        return aerial_path, street_view_path

    def get_coordinates_and_details(self, address: str) -> Tuple[float, float, Dict]:
        """
        Get latitude, longitude, and location details (Esri first, Google fallback).
        """
        try:
            return self.location_service.geocode(address)
        except Exception as e:
            logger.error(f"Error getting coordinates for {address}: {e}")
            raise

    def get_census_data(self, lat: float, lng: float, county: str, state: str) -> Dict:
        """
        Get demographic data from Esri GeoEnrichment, with AI fallback.
        """
        try:
            demographics = self.location_service.get_demographics(lat, lng)
            if demographics:
                logger.info("Using Esri GeoEnrichment demographics")
                return demographics

            census_prompt = f"""
            Generate realistic and current demographic data for coordinates {lat}, {lng} in {county}, {state}.
            Include population statistics, household data, employment data, and economic factors.
            Format as JSON with keys: population_2020, population_growth_rate, households_2020, 
            avg_household_size, employment_rate, major_industries, median_income.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a demographic data analyst. Provide realistic census-like data in JSON format."},
                    {"role": "user", "content": census_prompt}
                ],
                temperature=0.1
            )
            
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"Error getting census data: {e}")
            return {}

    def _search_current_market_data(self, county: str, state: str, property_type: str) -> Dict:
        """Search for current market data using AI to simulate real market conditions"""
        
        # Get current quarter
        current_date = datetime.now()
        current_quarter = f"Q{(current_date.month-1)//3 + 1} {current_date.year}"
        
        prompt = f"""
        Generate realistic current {current_quarter} market data for {property_type} in {county}, {state}.
        
        Include the following metrics with realistic values:
        1. Direct vacancy rate with QoQ and YoY changes
        2. Sublease vacancy rate with QoQ change
        3. Total vacancy rate
        4. Average lease rates overall and by class (A, B, C)
        5. Construction pipeline (square footage under construction and type)
        6. Market absorption rates
        7. Population growth projections
        8. Employment statistics
        9. Major market trends
        
        Format as JSON with this structure:
        {{
            "quarter": "{current_quarter}",
            "direct_vacancy": 12.71,
            "direct_qoq": "+1.02",
            "direct_yoy": "+1.68",
            "sublease_vacancy": 5.51,
            "sublease_qoq": "-0.39",
            "total_vacancy": 18.22,
            "avg_lease_rate": 24.33,
            "lease_rate_yoy": "-0.04",
            "cap_rate": 6.7,
            "class_a_rate": 27.13,
            "class_b_rate": 22.03,
            "class_c_rate": 19.44,
            "construction_sf": 24000,
            "construction_type": "medical office",
            "population_projection_2060": 5500000,
            "unemployment_rate": 3.8,
            "national_unemployment": 4.4,
            "absorption_rate_sf": 15000,
            "major_trends": ["hybrid work adoption", "flight to quality", "suburban growth"]
        }}
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": f"You are a commercial real estate market analyst. Provide current, realistic {current_quarter} market data specific to {county}, {state}. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                response_format={"type": "json_object"},
            )

            content = response.choices[0].message.content.strip()
            # Strip accidental markdown fences before parsing
            if content.startswith("```"):
                content = content.split("```", 2)[1].lstrip("json").strip()
            return json.loads(content)
        except Exception as e:
            logger.error(f"Error generating market data: {e}")
            # Return default data as fallback
            return {
                "quarter": current_quarter,
                "direct_vacancy": 12.71,
                "direct_qoq": "+1.02",
                "direct_yoy": "+1.68",
                "sublease_vacancy": 5.51,
                "sublease_qoq": "-0.39",
                "total_vacancy": 18.22,
                "avg_lease_rate": 24.33,
                "lease_rate_yoy": "-0.04",
                "cap_rate": 6.7,
                "class_a_rate": 27.13,
                "class_b_rate": 22.03,
                "class_c_rate": 19.44,
                "construction_sf": 24000,
                "construction_type": "medical office",
                "population_projection_2060": 5500000,
                "unemployment_rate": 3.8,
                "national_unemployment": 4.4
            }

    def _generate_market_overview(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate a location-specific market overview via AI (not a static template)."""
        county = self._ctx_value(context, 'County')
        state = self._ctx_value(context, 'State')
        quarter = market_data.get('quarter', self._current_quarter())

        prompt = f"""Write a professional 2-paragraph market overview for the {property_type} market
in {county}, {state}, as of {quarter}.

Ground everything specifically in {county}, {state}. Do NOT reference any other state, county, or city.
Incorporate these current metrics naturally in prose: total vacancy about {market_data.get('total_vacancy', 'n/a')}%,
average asking lease rate about ${market_data.get('avg_lease_rate', 'n/a')}/SF, local unemployment about
{market_data.get('unemployment_rate', 'n/a')}% versus the national average of about {market_data.get('national_unemployment', 'n/a')}%.

Cover local population and employment growth drivers, demand for {property_type} space, and current leasing
dynamics (e.g., sublease availability, flight to quality, hybrid-work effects where relevant).
Do NOT cite or mention CoStar / Costar / Co-Star by name. Attribute market conditions only to public
brokerage reports (e.g. Colliers, CBRE, JLL) or generic public sources if needed.
Plain text only. No markdown, no bullet points, no headings."""
        return self._get_ai_response(prompt)

    def _generate_vacancy_rates(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate Key Market Metrics bullets in CLIENT-report style (proxy metrics + Note)."""
        county = self._ctx_value(context, "County")
        state = self._ctx_value(context, "State")
        quarter = market_data.get("quarter", self._current_quarter())
        vacancy = market_data.get("total_vacancy", "n/a")
        rent = market_data.get("avg_lease_rate", "n/a")
        construction_sf = market_data.get("construction_sf", 0) or 0
        cap_rate = market_data.get("cap_rate", "n/a")

        prompt = f"""Write the Key Market Metrics block under "Vacancy Rates" for a {property_type}
property in {county}, {state}, as of {quarter}.

Match this EXACT structure and tone (plain text only — no markdown, no **bold** markers):

• Vacancy (metro proxy): {vacancy}% ({quarter}); <one short sentence on metro trend and how {county} typically tracks it>.
• Asking Rents (metro proxy): ~${rent}/SF market asking rent ({quarter}), with brief rent-growth context. Note that {property_type} reporting is not commonly split into Class A/B.
• Construction Pipeline (metro): ~{int(construction_sf):,} SF under construction as of {quarter}; one short pipeline/delivery note.
• Investment Indicators (metro): Average {property_type} capitalization rate ~{cap_rate}% ({quarter}).
Note: Reliable {county}-only breakouts for {property_type} vacancy and asking rents are limited; metro-level metrics are the accepted proxy in most public reports and broker opinions of value.

Rules:
- One bullet per line, each starting with "• " (bullet character + space).
- After the four bullets, one "Note:" paragraph (no bullet).
- Separate each item with a blank line.
- Do NOT output Direct/Sublease/Total vacancy line items.
- Do NOT invent other counties/states. Ground everything in {county}, {state} / its metro.
- NEVER cite or mention CoStar / Costar / Co-Star.
"""
        return self._get_ai_response(prompt)

    def _generate_lease_rates(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate lease-rates prose (CLIENT style) — not Class A/B/C bullet lists."""
        county = self._ctx_value(context, "County")
        state = self._ctx_value(context, "State")
        quarter = market_data.get("quarter", self._current_quarter())
        rent = market_data.get("avg_lease_rate", "n/a")
        rent_yoy = market_data.get("lease_rate_yoy", "n/a")

        prompt = f"""Write ONE short paragraph for "Lease Rates (Based on Public Listings & Industry Trends)"
for {property_type} in {county}, {state}, as of {quarter}.

Style example:
"{property_type} lease rates in {county} generally align with broader metro trends, averaging around ${rent}/SF as of {quarter}. Annual rent growth has remained steady{'' if rent_yoy == 'n/a' else f' ({rent_yoy} YoY)'}, supported by disciplined supply and healthy tenant demand in well-located corridors."

Rules:
- Plain text only, single paragraph, no bullets, no Class A/B/C splits, no markdown.
- Keep it to 2–4 sentences. Do not reference other markets.
"""
        return self._get_ai_response(prompt)

    def _generate_construction_activity(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate a location-specific construction activity section via AI."""
        county = self._ctx_value(context, 'County')
        state = self._ctx_value(context, 'State')
        quarter = market_data.get('quarter', self._current_quarter())
        construction_sf = market_data.get('construction_sf', 0) or 0

        prompt = f"""Write a short "Construction Activity" section for the {property_type} market in
{county}, {state}, as of {quarter}.

Use these figures: approximately {int(construction_sf):,} SF under construction, primarily
{market_data.get('construction_type', property_type)} product.
Reference realistic, plausible recent deliveries and pipeline trends specific to {county}, {state}.
Do NOT reference any other market (no Salt Lake, Lehi, Provo, etc. unless that is the actual county/state).

Return 2-3 concise bullet points, each starting with "• ". Plain text only. Separate bullets with blank lines."""
        return self._get_ai_response(prompt)

    def _generate_market_trends(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate location-specific market trends via AI."""
        county = self._ctx_value(context, 'County')
        state = self._ctx_value(context, 'State')
        quarter = market_data.get('quarter', self._current_quarter())
        trends = ", ".join(market_data.get('major_trends', []) or [])

        prompt = f"""Write a "Trends & Forecast" section for the {property_type} market in {county}, {state}, as of {quarter}.
{f'Relevant current themes to weave in: {trends}.' if trends else ''}
Cover 3-4 distinct trends (e.g., population/employment growth, leasing flexibility/sublease, supply pipeline,
hybrid-work or e-commerce effects) as they specifically apply to {county}, {state} and to {property_type}.
Do NOT reference any other market.
Format: each trend is one bullet starting with "• " and a short label then colon (e.g. "• Demand skew to small shops: ...").
Separate bullets with blank lines. Plain text only — no markdown (no **, *, #, or backticks)."""
        return self._get_ai_response(prompt)

    def _generate_investment_insights(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate location-specific investment insights via AI."""
        county = self._ctx_value(context, 'County')
        state = self._ctx_value(context, 'State')

        prompt = f"""Write "Investment Insights" for a {property_type} property in {county}, {state}.
Provide 3 concise, actionable insights grounded in the local {county}, {state} market and {property_type} fundamentals.
Do NOT reference any other market.
Return 3 bullets, each starting with "• " and a short label then colon. Separate with blank lines.
Plain text only — no markdown (no **, *, #, or backticks)."""
        return self._get_ai_response(prompt)

    def _generate_market_recommendations(self, context: str, property_type: str, market_data: Dict) -> str:
        """Generate location-specific recommendations via AI (CLIENT short-bullet style)."""
        county = self._ctx_value(context, "County")
        state = self._ctx_value(context, "State")

        prompt = f"""Write "Recommendations" for a {property_type} Broker Opinion of Value in {county}, {state}.

Return EXACTLY 3 short bullets in this style (plain text only):
• Prioritize <one concrete action grounded in {county} {property_type} conditions>.
• Target <one concrete leasing / tenant / location action>.
• For value-add, focus on <one concrete underwriting / lease-up / mark-to-market action>.

Rules:
- Each bullet is ONE sentence (max ~35 words). Start with "• ".
- Separate bullets with a blank line.
- Do NOT use markdown (no **, no *, no #, no backticks).
- Do NOT label bullets as Investors / Tenants / Developers.
- Do NOT reference any other market.
"""
        return self._get_ai_response(prompt)

    @staticmethod
    def _strip_markdown(text: str) -> str:
        """Remove common markdown markers AI sometimes emits into Word body text."""
        if not text:
            return text or ""
        # Bold / italic
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # Headings / list markers that are not our bullets
        text = re.sub(r"(?m)^#{1,6}\s*", "", text)
        return text

    @staticmethod
    def _scrub_costar_citations(text: str) -> str:
        """Remove CoStar attributions from report body text (legal risk).

        User-uploaded CoStar comps are fine; citing CoStar as a market-data
        source in disclaimers / narrative is not.
        """
        if not text:
            return text or ""
        # Replace common "CoStar – …" / "CoStar - …" source bullets with Colliers
        text = re.sub(
            r"(?im)^(\s*•\s*)CoStar\b([^\n]*)",
            r"\1Colliers International\2",
            text,
        )
        text = re.sub(
            r"(?i)\bCoStar\b(\s*[–—-]\s*)",
            r"Colliers International\1",
            text,
        )
        # Any remaining bare CoStar mentions in prose → generic public brokerage
        text = re.sub(r"(?i)\bCo-?Star\b", "public brokerage market reports", text)
        return text

    def _generate_data_sources(self, context: str = "", property_type: str = "", market_data: Dict = None) -> str:
        """Generate a current, market-relevant data sources list via AI."""
        market_data = market_data or {}
        county = self._ctx_value(context, 'County')
        state = self._ctx_value(context, 'State')
        quarter = market_data.get('quarter', self._current_quarter())
        year = datetime.now().year

        prompt = f"""Write the "Data Sources & Disclaimer" sources block for a {property_type} Broker Opinion
of Value in {county}, {state}, as of {quarter}.

Match this EXACT structure (plain text only — no markdown):

This analysis relies on multiple public sources:

• Colliers International – {county}/{state} metro {property_type} Market Report {quarter} (vacancy rates, rental rates, absorption)
• CBRE / JLL / similar public brokerage – {property_type} Market Report {quarter} (market trends and leasing activity)
• U.S. Census Bureau – latest population, income, and housing statistics for {county}, {state}
• Esri Business Analyst / GeoEnrichment – current demographic and consumer spending data
• U.S. Bureau of Labor Statistics – current employment and unemployment data
• Local county appraisal / assessor records – public property records for {county}

Rules:
- Start with the italic-style intro line exactly: "This analysis relies on multiple public sources:"
- Then 5–6 bullets, each starting with "• " (bullet character). Separate bullets with blank lines.
- Do NOT use a numbered list (no "1.", "2.", "Sources Used:").
- Keep each bullet to one line. Sources must be current ({year} / {quarter}), not older than two years.
- Do NOT invent Utah-specific sources unless the property is in Utah.
- NEVER cite, name, or attribute anything to CoStar / Costar / Co-Star. CoStar is forbidden as a source.
- Only cite publicly available brokerage quarterly reports (Colliers, CBRE, JLL, Newmark, etc.), census/BLS/Esri, and public assessor records.
"""
        text = self._get_ai_response(prompt)
        # Force CLIENT-style bullets even if the model returns "1. 2. 3."
        text = text.replace("Sources Used:", "This analysis relies on multiple public sources:")
        text = re.sub(r"(?m)^\s*\d+\.\s+", "• ", text)
        return self._scrub_costar_citations(text)

    @staticmethod
    def _ctx_value(context: str, label: str) -> str:
        """Safely pull a labelled value (e.g. 'County') out of the context block."""
        try:
            return context.split(f'{label}:')[1].split('\n')[0].strip()
        except (IndexError, AttributeError):
            return ""

    @staticmethod
    def _current_quarter() -> str:
        now = datetime.now()
        return f"Q{(now.month - 1) // 3 + 1} {now.year}"

    def _generate_market_analysis_sections(self, context: str, property_type: str) -> Dict[str, str]:
        """Generate all market analysis sections"""
        
        # Get current market data
        county = context.split('County:')[1].split('\n')[0].strip() if 'County:' in context else 'Utah County'
        state = context.split('State:')[1].split('\n')[0].strip() if 'State:' in context else 'Utah'
        
        market_data = self._search_current_market_data(county, state, property_type)
        
        sections = {
            'market_overview': self._generate_market_overview(context, property_type, market_data),
            'vacancy_rates': self._generate_vacancy_rates(context, property_type, market_data),
            'lease_rates': self._generate_lease_rates(context, property_type, market_data),
            'construction_activity': self._generate_construction_activity(context, property_type, market_data),
            'market_trends': self._generate_market_trends(context, property_type, market_data),
            'investment_insights': self._generate_investment_insights(context, property_type, market_data),
            'market_recommendations': self._generate_market_recommendations(context, property_type, market_data),
            'market_data_sources': self._generate_data_sources(context, property_type, market_data)
        }
        # Word shows literal **bold** if markdown slips through — strip it.
        # Also scrub any CoStar attributions (never cite CoStar as a market source).
        return {
            k: self._scrub_costar_citations(self._strip_markdown(v))
            for k, v in sections.items()
        }

    def generate_comprehensive_content(self, address: str, property_data: PropertyReportData) -> PropertyReportData:
        """
        Generate all content sections using AI with the exact formatting requirements
        """
        logger.info(f"Generating comprehensive content for: {address}")
        
        # Create detailed context for AI
        transaction = ""
        if property_data.lease_or_sale:
            kind = property_data.lease_or_sale.strip().lower()
            if "lease" in kind:
                transaction = "Transaction Context: The property is being evaluated for LEASE."
            else:
                transaction = "Transaction Context: The property is being evaluated for SALE."
        context = f"""
        Property Address: {address}
        Property Type: {property_data.property_type}
        County: {property_data.county}
        State: {property_data.state}
        Coordinates: {property_data.latitude}, {property_data.longitude}
        {transaction}
        """
        
        # Generate property analysis sections
        sections = {
            'property_summary': self._generate_property_summary(context),
            'location_summary': self._generate_location_summary(context),
            'demographic_analysis': self._generate_demographic_analysis(context),
            'size_and_topography': self._generate_size_topography(context),
            'population_analysis': self._generate_population_analysis(context),
            'household_trends': self._generate_household_trends(context),
            'employment_analysis': self._generate_employment_analysis(context),
            'economic_factors': self._generate_economic_factors(context),
            'community_services': self._generate_community_services(context),
            'swot_analysis': self._generate_swot_analysis(context)
        }
        
        # Update property data with generated content
        property_data.property_summary = sections['property_summary']
        property_data.location_summary = sections['location_summary']
        property_data.demographic_analysis = sections['demographic_analysis']
        property_data.size_and_topography = sections['size_and_topography']
        property_data.population_analysis = sections['population_analysis']
        property_data.household_trends = sections['household_trends']
        property_data.employment_analysis = sections['employment_analysis']
        property_data.economic_factors = sections['economic_factors']
        property_data.community_services = sections['community_services']
        
        # SWOT Analysis
        swot = sections['swot_analysis']
        property_data.swot_strengths = swot.get('strengths', '')
        property_data.swot_weaknesses = swot.get('weaknesses', '')
        property_data.swot_opportunities = swot.get('opportunities', '')
        property_data.swot_threats = swot.get('threats', '')
        
        # Generate market analysis sections
        market_sections = self._generate_market_analysis_sections(context, property_data.property_type)
        
        # Update property data with market analysis
        property_data.market_overview = market_sections['market_overview']
        property_data.vacancy_rates = market_sections['vacancy_rates']
        property_data.lease_rates = market_sections['lease_rates']
        property_data.construction_activity = market_sections['construction_activity']
        property_data.market_trends = market_sections['market_trends']
        property_data.investment_insights = market_sections['investment_insights']
        property_data.market_recommendations = market_sections['market_recommendations']
        property_data.market_data_sources = market_sections['market_data_sources']

        comp_context = self._format_comp_context(property_data.comps)

        # Generate additional BOV narrative sections
        property_data.executive_summary = self._get_ai_response(
            f"Write a concise 1-paragraph executive summary for a Broker Opinion of Value of a "
            f"{property_data.property_type} property. Context:\n{context}\n"
            f"{comp_context}"
            f"Summarize location, key value drivers, and the value conclusion. Plain text only.",
        )
        property_data.regional_analysis = self._shorten_regional_analysis(
            self._get_ai_response(
                f"Write a SHORT but worthy REGIONAL ANALYSIS for the subject in "
                f"{property_data.county}, {property_data.state} ({datetime.now().year}).\n\n"
                f"Match this style (plain text, no markdown, no bullets):\n"
                f"Paragraph 1: population growth/size and household/income character that support "
                f"{property_data.property_type} demand (2 sentences max).\n"
                f"Paragraph 2: key employers/industries plus major roads/airport access and why that "
                f"helps the local {property_data.property_type} market (2 sentences max).\n\n"
                f"HARD LIMITS: exactly 2 paragraphs, blank line between them, 85-100 words TOTAL. "
                f"Name concrete employers/roads when plausible. Never mention another county/state.\n"
                f"Context:\n{context}",
            )
        )
        # CLIENT layout: SALES CONCLUSION heading + OPINIONS OF VALUE table only
        # (no long narrative between/after those elements)
        property_data.sales_conclusion = ""
        stats = self._comp_psf_stats(property_data.comps)
        recon_extra = ""
        if stats:
            recon_extra = (
                f" You MUST cite the comparable $/SF range as "
                f"${stats['min']:.2f} to ${stats['max']:.2f} and the average "
                f"${stats['avg']:.2f}/SF from {int(stats['count'])} uploaded comps. "
                f"Do not invent other $/SF figures."
            )
        property_data.reconciliation_summary = self._get_ai_response(
            f"Write 2 short sentences for the RECONCILIATION TABLE narrative (above the valuation grid). "
            f"Mention the comparable $/SF range when comps are available and that the sales comparison "
            f"approach supports the opinion of value for this {property_data.property_type}. "
            f"{recon_extra}\n"
            f"Context:\n{context}\n{comp_context}"
            f"Plain text only. Keep under 60 words.",
        )
        # Short NOTES cell for the valuation grid (CLIENT style one-liner)
        property_data.reconciliation_notes = self._build_reconciliation_notes(property_data)

        return property_data

    @staticmethod
    def _format_comp_context(comps: List[Any]) -> str:
        if not comps:
            return ""
        # Do not name CoStar in AI context — models can echo it into report prose.
        lines = ["Comparable sales from user-uploaded comps:"]
        for comp in sorted(comps, key=lambda c: getattr(c, "comp_number", 0))[:6]:
            lines.append(
                f"- Comp {getattr(comp, 'comp_number', '?')}: {getattr(comp, 'address', '')}, "
                f"Sale {getattr(comp, 'sale_price', 'N/A')}, "
                f"{getattr(comp, 'sale_price_sf', 'N/A')}/SF, "
                f"{getattr(comp, 'comp_sf', '')} SF"
            )
        stats = ComprehensivePropertyReportGenerator._comp_psf_stats(comps)
        if stats:
            lines.append(
                f"Computed from uploaded comps: avg ${stats['avg']:.2f}/SF "
                f"(range ${stats['min']:.2f}–${stats['max']:.2f}, n={stats['count']}). "
                f"Use ONLY these figures for any $/SF discussion."
            )
        return "\n".join(lines) + "\n"

    @staticmethod
    def _comp_psf_stats(comps: List[Any]) -> Optional[Dict[str, float]]:
        """Average / min / max Sale Price/SF from user-uploaded comps.

        Prefers extracted Sale Price/SF; falls back to Sale Price ÷ Comp SF when
        needed. Returns None when no usable $/SF values are present.
        """
        values: List[float] = []
        for comp in comps or []:
            psf = ComprehensivePropertyReportGenerator._to_number(
                getattr(comp, "sale_price_sf", None)
            )
            if psf is None or psf <= 0:
                price = ComprehensivePropertyReportGenerator._to_number(
                    getattr(comp, "sale_price", None)
                )
                sf = ComprehensivePropertyReportGenerator._to_number(
                    getattr(comp, "comp_sf", None)
                )
                if price and sf and sf > 0:
                    psf = price / sf
            if psf is not None and psf > 0:
                values.append(float(psf))

        if not values:
            return None

        avg = sum(values) / len(values)
        return {
            "avg": round(avg, 2),
            "min": round(min(values), 2),
            "max": round(max(values), 2),
            "count": float(len(values)),
        }

    def _build_reconciliation_notes(self, property_data: "PropertyReportData") -> str:
        """One-line NOTES cell for the valuation grid (CLIENT style)."""
        comps = getattr(property_data, "comps", None) or []
        n_comps = len(comps) if comps else 0

        rounded = None
        # Prefer finalized valuation from table_values when available
        tv = getattr(property_data, "table_values", None) or {}
        rounded = tv.get("{{market_value_rounded}}")
        if not rounded and getattr(property_data, "bov_dataset", None):
            val = (property_data.bov_dataset or {}).get("valuation") or {}
            mv = val.get("market_value_rounded") or val.get("market_value")
            if mv is not None:
                try:
                    rounded = f"${int(round(float(mv))):,}"
                except (TypeError, ValueError):
                    rounded = str(mv)

        if not rounded:
            rounded = "the concluded market value"

        if n_comps > 0:
            stats = self._comp_psf_stats(comps)
            if stats:
                return (
                    f"The sales comparison approach yields a value of {rounded} "
                    f"based on the average ${stats['avg']:,.2f}/SF from "
                    f"{int(stats['count'])} comparables."
                )
            return (
                f"The sales comparison approach yields a value of {rounded} "
                f"based on the average $/SF from {n_comps} comparables."
            )
        return (
            f"The sales comparison approach yields a value of {rounded} "
            f"based on comparable sales analysis."
        )

    @staticmethod
    def _shorten_regional_analysis(text: str, max_words: int = 100) -> str:
        """Keep regional analysis to ~2 short paragraphs (CLIENT density)."""
        if not text:
            return ""
        text = ComprehensivePropertyReportGenerator._strip_markdown(text).strip()
        parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        if len(parts) == 1:
            sentences = re.split(r"(?<=[.!?])\s+", parts[0])
            sentences = [s.strip() for s in sentences if s.strip()]
            if len(sentences) >= 2:
                mid = max(1, len(sentences) // 2)
                parts = [" ".join(sentences[:mid]), " ".join(sentences[mid:])]
            else:
                parts = [parts[0]]
        parts = parts[:2]

        def _clip(paragraph: str, budget: int) -> str:
            words = paragraph.split()
            if len(words) <= budget:
                return paragraph
            clipped = " ".join(words[:budget]).rstrip(",;:")
            if not clipped.endswith((".", "!", "?")):
                clipped += "."
            return clipped

        budget_each = max(35, max_words // max(len(parts), 1))
        parts = [_clip(p, budget_each) for p in parts]
        joined = "\n\n".join(parts)
        words = joined.split()
        if len(words) > max_words:
            joined = " ".join(words[:max_words]).rstrip(",;:")
            if not joined.endswith((".", "!", "?")):
                joined += "."
        return joined

    def _style_regional_analysis(self, doc: Document) -> None:
        """CLIENT style: italic blue body under REGIONAL ANALYSIS."""
        from docx.shared import RGBColor, Pt

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().upper() != "REGIONAL ANALYSIS":
                continue
            for j in range(i + 1, len(doc.paragraphs)):
                nxt = doc.paragraphs[j]
                style_name = (nxt.style.name if nxt.style else "") or ""
                upper = nxt.text.strip().upper()
                if "Heading" in style_name or upper in (
                    "DEMOGRAPHIC ANALYSIS",
                    "LOCATION SUMMARY",
                    "PROPERTY SUMMARY",
                ):
                    break
                if not nxt.text.strip():
                    continue
                for run in nxt.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                    run.font.size = Pt(12)
            break

    def _generate_property_summary(self, context: str) -> str:
        """Generate property summary matching the exact format"""
        example = """The subject is located in Salem, in Utah County. It is part of the Provo-Orem MSA. The subject property is located in northern Utah within the official boundaries of Utah County. The county is situated directly south of Salt Lake County. This area is generally called the Provo/Orem metropolitan area and is approximately 45 miles south of metropolitan Salt Lake, which is the financial center for the Intermountain Region. This region encompasses all of Utah, southern Idaho, southwestern Wyoming, and eastern Nevada. Utah County is part of a four-county area that is commonly known as the Wasatch Front. Provo is the Utah County seat."""
        
        prompt = f"""
        Generate a property summary paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Maintain the same structure: location, county, MSA/metropolitan area, regional context, broader region description, and county seat information.
        """
        
        return self._get_ai_response(prompt)

    def _generate_location_summary(self, context: str) -> str:
        """Generate location summary matching the exact format"""
        example = """The subject is located on the corner of Hwy 198 and Elk Ridge Drive with good access and exposure. A major thoroughfare in the area is Hwy 198 which partially fronts the subject. The location also offers very close proximity to Salem Pond, Salem High School, Salem Community Center, Salem City Recreation, with limited retail areas in close proximity. The subject is surrounded by vacant land and residential uses. Utah County is broken up into three sectors.  North County (Lindon to Lehi) Central County (Provo/Orem) and South County (Springville to Payson). Central county accounts for a lot of the class B office buildings. The following is taken from Reis, it shows the market area with an arrow pointing to the subject."""
        
        prompt = f"""
        Generate a DETAILED location summary following this FORMAT and style (adapt fully to the real location):
        {example}

        Use this context: {context}

        Write a thorough 6-9 sentence description. Include, when applicable to the actual location:
        - the specific corridor/intersection and major thoroughfares fronting or serving the site
        - access and visibility/exposure characteristics
        - nearby anchors, retail, employers, schools, parks, and other notable amenities (name plausible real ones for that area)
        - surrounding land uses and development character
        - how the site sits within the broader county/metro and its submarket subdivisions
        Match the older template's level of detail. Plain text only, no markdown.
        """

        return self._get_ai_response(prompt)

    def _generate_demographic_analysis(self, context: str) -> str:
        """Qualitative demographic overview — no invented statistics.

        Exact population / household / tenure figures are filled later from
        Esri tables via `_sync_narratives_to_table_data`.
        """
        prompt = f"""
        Write a short DEMOGRAPHIC OVERVIEW (2-3 sentences) for this property.
        Context:
        {context}

        Rules:
        - Describe location character only (families, workforce, retail/office demand drivers).
        - Do NOT invent or cite any specific numbers, percentages, dollars, or years.
        - Do NOT invent median income, population totals, growth rates, or vacancy.
        - Plain text only, no markdown.
        Example tone: "The subject benefits from a solid demographic profile within the
        county, with a mix of households and employment bases that support local
        commercial demand. Accessibility to major corridors further strengthens the
        draw for nearby residents and workers."
        """
        return self._get_ai_response(prompt)

    def _generate_size_topography(self, context: str) -> str:
        """Generate size and topography description"""
        example = """The surrounding mountains form a valley about 30 miles wide and 50 miles long. Utah Lake is located centrally to the valley and is Utah's largest freshwater lake. The Wasatch Mountains, which provide a beautiful background to the county on the east, nearly converge with Utah Lake on the west to form the southern boundary south of Santaquin City. The northern boundary is considered the "point of the mountain" which is just north of Lehi City. The elevation varies from 4,480 to 11,928 feet (Mt. Nebo) above sea level. Utah Lake and Mt. Timpanogos present a mountainous scenic backdrop within this metropolitan setting."""
        
        prompt = f"""
        Generate a size and topography paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: geographical features, valley dimensions, major landmarks, elevation ranges, and scenic elements.
        """
        
        return self._get_ai_response(prompt)

    def _generate_population_analysis(self, context: str) -> str:
        """Generate population analysis with specific statistics"""
        example = """According to Pitney Bowes/Gadberry Group - GroundView®, a Geographic Information System (GIS) Company, Utah County had a 2020 total population of 649,258 and experienced an annual growth rate of 2.3%, which was higher than the Utah annual growth rate of 1.6%. The county accounted for 20.0% of the total Utah population (3,254,284). Within the county the population density was 304 people per square mile compared to the lower Utah population density of 38 people per square mile and the lower United States population density of 92 people per square mile."""
        
        prompt = f"""
        Generate a population analysis paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: data source attribution, specific population numbers, growth rates, state comparisons, population density comparisons. Use current 2024 data where available.
        """
        
        return self._get_ai_response(prompt)

    def _generate_household_trends(self, context: str) -> str:
        """Generate household trends analysis"""
        example = """The 2020 number of households in the county was 178,689. The number of households in the county is projected to grow by 2.0% annually, increasing the number of households to 197,669 by 2025. The 2020 average household size for the county was 3.55, which was 37.61% larger than the United States average household size of 2.58 for 2020. The average household size in the county is anticipated to retract by 0.06% annually, reducing the average household size to 3.54 by 2025."""
        
        prompt = f"""
        Generate a household trends paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: specific household numbers, growth projections, average household size, national comparisons, and future projections. Use current 2024 data and project to 2029.
        """
        
        return self._get_ai_response(prompt)

    def _generate_employment_analysis(self, context: str) -> str:
        """Generate employment analysis using the most recent years available."""
        year = datetime.now().year
        example = f"""Total employment has increased annually over the past three years in the state by roughly 2.5% and by roughly 3.9% in the county. From {year-1} to {year} unemployment fell in the state by about 0.4% and by about 0.4% in the county, with the county's rate remaining below the national average. Over the most recent month, unemployment declined by about 0.3% at both the state and county level."""

        prompt = f"""
        Generate an employment analysis paragraph following this FORMAT and style (do not copy its numbers):
        {example}

        Use this context: {context}

        Requirements:
        - Use the MOST RECENT data available (reference {year-2}-{year}); never cite a decade-old range such as 2010-2019.
        - Include: recent employment growth rates, unemployment trends, state vs county comparisons, and the latest month-over-month change.
        - Base facts on the specific state and county in the context. Plain text only.
        """

        return self._get_ai_response(prompt)

    def _generate_economic_factors(self, context: str) -> str:
        """Generate economic factors analysis"""
        example = """Salem is a suburb of Payson and Provo/Orem market area. Salem is still considered somewhat of a rural area but over the years has begun to be built out. A majority of resident's commute to other cities within the metropolitan area for employment. The largest industries in the city include manufacturing, public administration agricultural uses and retail trade. The local economy consists of commercial and industrial businesses on the main arterials. The city's commercial area is on Hwy 198, featuring retail, office, residential, and financial services."""
        
        prompt = f"""
        Generate an economic factors paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: suburban/rural character, development patterns, commuting patterns, major industries, commercial areas, and business types.
        """
        
        return self._get_ai_response(prompt)

    def _generate_community_services(self, context: str) -> str:
        """Generate community services description"""
        example = """Community services and facilities are readily available in the surrounding area. These include public services such as fire stations, hospitals, police stations, and schools (all ages). GreatSchools.org is an on-line tool that rates every school on a scale of one to ten based on test scores. They also track parents rating of the school on a one to five scale."""
        
        prompt = f"""
        Generate a community services paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: availability of services, specific service types, educational resources, and rating systems.
        """
        
        return self._get_ai_response(prompt)

    def _generate_swot_analysis(self, context: str) -> Dict[str, str]:
        """Generate SWOT analysis components"""
        prompt = f"""
        Generate a SWOT analysis for this property location. Return as JSON with keys: strengths, weaknesses, opportunities, threats.
        
        Context: {context}
        
        Examples:
        - Strengths: "Easy access from Highway 198 Within close proximity to residential developments"
        - Weaknesses: "The subject has average to weak visibility"  
        - Opportunities: "Opportunity for development of improvement on property"
        - Threats: "There is excess land around the property that could possibility be developed."
        
        Keep each section short and concise, focusing on location-specific factors.
        """
        
        response = self._get_ai_response(prompt, json_response=True)
        try:
            return json.loads(response)
        except:
            return {
                'strengths': 'Good location with development potential',
                'weaknesses': 'Limited visibility from main roads',
                'opportunities': 'Opportunity for future development',
                'threats': 'Competition from nearby available land'
            }

    def _get_ai_response(self, prompt: str, json_response: bool = False) -> str:
        """Get response from OpenAI API"""
        try:
            kwargs = {
                "model": "gpt-4o",
                "messages": [
                    {"role": "system", "content": "You are a professional commercial real estate analyst. Use the example only for FORMAT/style; base all facts on the specific location provided in the context. Never reference Utah, Salt Lake, Provo, or Lehi unless the subject property is actually there. Provide factual, current information." + (" Respond in valid JSON format." if json_response else "")},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.3,
            }
            if json_response:
                kwargs["response_format"] = {"type": "json_object"}
            response = self.openai_client.chat.completions.create(**kwargs)
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error getting AI response: {e}")
            return "Content generation failed"

    def create_property_report(self, 
                             address: str,
                             prepared_by: str = "",
                             prepared_by_title: str = "",
                             prepared_by_company: str = "",
                             prepared_by_address: str = "",
                             prepared_for: str = "",
                             prepared_for_title: str = "",
                             prepared_for_company: str = "",
                             prepared_for_address: str = "",
                             property_name: str = "",
                             property_type: str = "Vacant Land",
                             **kwargs) -> PropertyReportData:
        """
        Create complete property report data
        """
        logger.info(f"Creating property report for: {address}")

        comps = kwargs.pop("comps", None) or []
        
        # Get coordinates and location details
        lat, lng, address_details = self.get_coordinates_and_details(address)
        
        # Initialize property data
        property_data = PropertyReportData(
            address=address,
            prepared_by=prepared_by,
            prepared_by_title=prepared_by_title,
            prepared_by_company=prepared_by_company,
            prepared_by_address=prepared_by_address,
            prepared_for=prepared_for,
            prepared_for_title=prepared_for_title,
            prepared_for_company=prepared_for_company,
            prepared_for_address=prepared_for_address,
            property_name=property_name or f"{address_details.get('city', 'Property')} {property_type}",
            property_type=(property_type or "").strip().title(),
            state=address_details.get('state', ''),
            county=address_details.get('county', ''),
            latitude=str(lat),
            longitude=str(lng),
            **kwargs
        )
        
        # Get property images
        aerial_path, street_view_path = self.get_property_images(address, lat, lng)
        property_data.aerial_image_path = aerial_path
        property_data.street_view_image_path = street_view_path
        property_data.comps = comps
        
        # Generate all content sections (including market analysis)
        property_data = self.generate_comprehensive_content(address, property_data)

        # Build BOV table values (Esri GeoEnrichment first, AI fallback)
        try:
            lat_f = float(property_data.latitude)
            lng_f = float(property_data.longitude)
        except (TypeError, ValueError):
            lat_f, lng_f = 0.0, 0.0
        property_data.table_values = self.build_bov_dataset(
            lat_f, lng_f, property_data
        )
        # Rewrite demo/employment narratives from the SAME numbers shown in tables
        # (Esri admin + employment_history) so verbiage never invents different figures.
        self._sync_narratives_to_table_data(property_data)
        # Rebuild short valuation NOTES after GBA/valuation math is final
        property_data.reconciliation_notes = self._build_reconciliation_notes(property_data)
        property_data.table_values["{{reconciliation_notes}}"] = property_data.reconciliation_notes
        property_data.table_values["{{reconciliation_summary}}"] = (
            property_data.reconciliation_summary or ""
        )
        property_data.table_values["{{sales_conclusion}}"] = (
            property_data.sales_conclusion or ""
        )

        return property_data

    def build_bov_dataset(self, lat: float, lng: float, property_data: "PropertyReportData") -> Dict[str, str]:
        """Build the full set of BOV table placeholder values.

        Strategy: generate a complete realistic dataset with AI, then overlay
        real Esri GeoEnrichment numbers where the Esri key is available.
        """
        county = property_data.county or "the county"
        state = property_data.state or "the state"
        property_type = property_data.property_type or "Commercial"

        dataset = self._generate_bov_demographics_ai(county, state, property_type)

        # Use the entered GBA (Gross Building Area) to drive the valuation math.
        # Market $/SF MUST come from the average Sale Price/SF of uploaded comps
        # when available — never reuse a stale AI default (e.g. 265) across reports.
        gba = self._to_number(property_data.lot_area)
        if gba:
            valuation = dataset.setdefault("valuation", {})
            comp_stats = self._comp_psf_stats(property_data.comps)
            if comp_stats:
                price_psf = float(comp_stats["avg"])
                logger.info(
                    "Valuation $/SF from uploaded comps average: $%.2f "
                    "(min=$%.2f max=$%.2f n=%s)",
                    price_psf,
                    comp_stats["min"],
                    comp_stats["max"],
                    int(comp_stats["count"]),
                )
            else:
                price_psf = self._to_number(valuation.get("price_psf"))
                if not price_psf:
                    price_psf = 265.0
                    logger.warning(
                        "No comps $/SF available — using fallback $%.2f/SF", price_psf
                    )
                else:
                    logger.warning(
                        "No comps $/SF available — using AI valuation $%.2f/SF", price_psf
                    )

            valuation["building_sf"] = gba
            valuation["price_psf"] = round(float(price_psf), 2)
            market_value = float(price_psf) * gba
            valuation["market_value"] = market_value
            valuation["market_value_rounded"] = round(market_value / 10000) * 10000

            # Opinion band: prefer min/max of uploaded comps when we have a spread
            if (
                comp_stats
                and comp_stats["max"] > comp_stats["min"]
                and int(comp_stats["count"]) >= 2
            ):
                valuation["value_aggressive"] = round(float(comp_stats["max"]) * gba)
                valuation["value_conservative"] = round(float(comp_stats["min"]) * gba)
            else:
                valuation["value_aggressive"] = round(market_value * 1.04)
                valuation["value_conservative"] = round(market_value * 0.96)

        # Overlay real Esri ring demographics when available
        try:
            rings = self.location_service.get_ring_demographics(lat, lng)
            if rings:
                self._overlay_esri_rings(dataset, rings)
                dataset["demographics_source"] = "Esri GeoEnrichment"
                logger.info("Overlaid real Esri ring demographics onto BOV dataset")
        except Exception as exc:
            logger.warning("Could not overlay Esri ring demographics: %s", exc)

        # Overlay US / State / County tables from Esri admin geographies (not point buffer)
        try:
            admin = self.location_service.get_admin_demographics(lat, lng)
            if admin:
                self._overlay_esri_admin(dataset, admin)
                dataset["demographics_source"] = "Esri GeoEnrichment"
                logger.info(
                    "Overlaid Esri admin demographics (us=%s state=%s county=%s)",
                    bool(admin.get("us")),
                    bool(admin.get("state")),
                    bool(admin.get("county")),
                )
        except Exception as exc:
            logger.warning("Could not overlay Esri admin demographics: %s", exc)

        property_data.bov_dataset = dataset
        return self._format_bov_placeholders(dataset, property_data)

    def _overlay_esri_admin(self, dataset: Dict, admin: Dict[str, Dict]) -> None:
        """Replace US/State/County population, density, HH, tenure from Esri admin levels."""
        pop = dataset.setdefault("population", {})
        density = dataset.setdefault("density", {})
        households = dataset.setdefault("households", {})
        hh_size = dataset.setdefault("hh_size", {})
        tenure = dataset.setdefault("tenure", {})

        for year in ("2010", "2020", "2025"):
            pop.setdefault(year, {})
        for year in ("2020", "2025"):
            density.setdefault(year, {})
        for key in ("2024", "2029", "cagr"):
            households.setdefault(key, {})
            hh_size.setdefault(key, {})
        tenure.setdefault("owner", {})
        tenure.setdefault("renter", {})

        for geo in ("us", "state", "county"):
            row = admin.get(geo) or {}
            if not row:
                continue

            if row.get("pop_2010") is not None:
                pop["2010"][geo] = row["pop_2010"]
            if row.get("pop_2020") is not None:
                pop["2020"][geo] = row["pop_2020"]
            if row.get("pop_cy") is not None:
                pop["2025"][geo] = row["pop_cy"]

            if row.get("density_2020") is not None:
                density["2020"][geo] = row["density_2020"]
            if row.get("density_cy") is not None:
                density["2025"][geo] = row["density_cy"]

            if row.get("hh_cy") is not None:
                households["2024"][geo] = row["hh_cy"]
            if row.get("hh_fy") is not None:
                households["2029"][geo] = row["hh_fy"]
            cagr_hh = self._cagr(row.get("hh_cy"), row.get("hh_fy"), years=5)
            if cagr_hh is not None:
                households["cagr"][geo] = cagr_hh

            if row.get("hhsize_cy") is not None:
                hh_size["2024"][geo] = row["hhsize_cy"]
            hhsize_fy = row.get("hhsize_fy")
            if hhsize_fy is None and row.get("pop_fy") and row.get("hh_fy"):
                hhsize_fy = round(float(row["pop_fy"]) / float(row["hh_fy"]), 2)
            if hhsize_fy is not None:
                hh_size["2029"][geo] = hhsize_fy
            cagr_hs = self._cagr(row.get("hhsize_cy"), hhsize_fy, years=5)
            if cagr_hs is not None:
                hh_size["cagr"][geo] = cagr_hs

            if row.get("owner_pct") is not None:
                tenure["owner"][geo] = row["owner_pct"]
            if row.get("renter_pct") is not None:
                tenure["renter"][geo] = row["renter_pct"]

    @staticmethod
    def _cagr(start, end, years: int = 5):
        """Compound annual growth rate (%) over `years`."""
        try:
            start_f = float(start)
            end_f = float(end)
            if start_f <= 0 or end_f <= 0 or years <= 0:
                return None
            return round(((end_f / start_f) ** (1.0 / years) - 1.0) * 100.0, 1)
        except (TypeError, ValueError, ZeroDivisionError):
            return None

    @staticmethod
    def _fmt_int_commas(value) -> str:
        try:
            return f"{int(round(float(value))):,}"
        except (TypeError, ValueError):
            return ""

    @staticmethod
    def _fmt_pct_plain(value, decimals: int = 1) -> str:
        try:
            f = float(value)
            # Keep enough precision for small rates like 0.02%
            if decimals == 1 and 0 < abs(f) < 0.05:
                return f"{f:.2f}%"
            return f"{f:.{decimals}f}%"
        except (TypeError, ValueError):
            return ""

    @staticmethod
    def _fmt_dec(value, decimals: int = 2) -> str:
        try:
            return f"{float(value):.{decimals}f}"
        except (TypeError, ValueError):
            return ""

    def _sync_narratives_to_table_data(self, property_data: "PropertyReportData") -> None:
        """Overwrite AI narratives so they cite the same figures as BOV tables."""
        dataset = getattr(property_data, "bov_dataset", None) or {}
        if not dataset:
            return
        county = property_data.county or "the county"
        state = property_data.state or "the state"
        property_type = property_data.property_type or "commercial"

        pop_text = self._build_population_analysis_from_data(dataset, county, state)
        if pop_text:
            property_data.population_analysis = pop_text

        hh_text = self._build_household_trends_from_data(dataset, county, state)
        if hh_text:
            property_data.household_trends = hh_text

        tenure_text = self._build_housing_tenure_from_data(dataset, county, state)
        if tenure_text:
            property_data.housing_tenure = tenure_text

        local_text = self._build_local_area_analysis_from_data(
            dataset, county, state, property_type
        )
        if local_text:
            property_data.local_area_analysis = local_text

        demo_text = self._build_demographic_overview_from_data(
            dataset, county, state, property_type
        )
        if demo_text:
            property_data.demographic_analysis = demo_text

        emp_text = self._build_employment_analysis_from_data(dataset, county, state)
        if emp_text:
            property_data.employment_analysis = emp_text

        # Keep placeholder map in sync if already built
        if property_data.table_values:
            property_data.table_values["{{population_analysis}}"] = (
                property_data.population_analysis or ""
            )
            property_data.table_values["{{household_trends}}"] = (
                property_data.household_trends or ""
            )
            property_data.table_values["{{housing_tenure}}"] = (
                property_data.housing_tenure or ""
            )
            property_data.table_values["{{local_area_analysis}}"] = (
                property_data.local_area_analysis or ""
            )
            property_data.table_values["{{demographic_analysis}}"] = (
                property_data.demographic_analysis or ""
            )
            property_data.table_values["{{employment_analysis}}"] = (
                property_data.employment_analysis or ""
            )
        logger.info(
            "Synced population/household/tenure/local/demographic/employment narratives"
        )

    def _build_household_trends_from_data(
        self, dataset: Dict, county: str, state: str
    ) -> str:
        """Household Trends paragraph from Esri NUMBER OF HOUSEHOLDS / HH SIZE tables."""
        hh = dataset.get("households") or {}
        hs = dataset.get("hh_size") or {}
        hh_2024 = (hh.get("2024") or {}).get("county")
        hh_2029 = (hh.get("2029") or {}).get("county")
        hh_cagr = (hh.get("cagr") or {}).get("county")
        hs_2024 = (hs.get("2024") or {}).get("county")
        hs_2029 = (hs.get("2029") or {}).get("county")
        hs_cagr = (hs.get("cagr") or {}).get("county")
        us_hs_2024 = (hs.get("2024") or {}).get("us")

        if hh_2024 is None or hh_2029 is None or hh_cagr is None:
            return ""

        county_label = county if county.lower().endswith("county") else f"{county} County"
        hh_cagr_f = float(hh_cagr)
        if hh_cagr_f >= 0:
            hh_trend = (
                f"projected to grow by {self._fmt_pct_plain(hh_cagr_f)} annually, increasing "
                f"the number of households to {self._fmt_int_commas(hh_2029)} by 2029"
            )
        else:
            hh_trend = (
                f"projected to decline by {self._fmt_pct_plain(abs(hh_cagr_f))} annually, "
                f"reducing the number of households to {self._fmt_int_commas(hh_2029)} by 2029"
            )
        parts = [
            f"The 2024 number of households in {county_label} was "
            f"{self._fmt_int_commas(hh_2024)}. The number of households in the county is "
            f"{hh_trend}."
        ]

        if hs_2024 is not None and us_hs_2024 is not None:
            try:
                hs_c = float(hs_2024)
                us_c = float(us_hs_2024)
                if us_c > 0:
                    diff_pct = ((hs_c / us_c) - 1.0) * 100.0
                    if abs(diff_pct) < 0.05:
                        size_cmp = (
                            f"which was about the same as the United States average "
                            f"household size of {self._fmt_dec(us_hs_2024)} for 2024"
                        )
                    elif diff_pct > 0:
                        size_cmp = (
                            f"which was {abs(diff_pct):.2f}% larger than the United States "
                            f"average household size of {self._fmt_dec(us_hs_2024)} for 2024"
                        )
                    else:
                        size_cmp = (
                            f"which was {abs(diff_pct):.2f}% smaller than the United States "
                            f"average household size of {self._fmt_dec(us_hs_2024)} for 2024"
                        )
                    size_sent = (
                        f"The 2024 average household size for the county was "
                        f"{self._fmt_dec(hs_2024)}, {size_cmp}."
                    )
                    if hs_2029 is not None and hs_cagr is not None:
                        cagr_f = float(hs_cagr)
                        if cagr_f < 0:
                            size_sent += (
                                f" The average household size in the county is anticipated "
                                f"to retract by {self._fmt_pct_plain(abs(cagr_f))} annually, "
                                f"reducing the average household size to "
                                f"{self._fmt_dec(hs_2029)} by 2029."
                            )
                        elif cagr_f > 0:
                            size_sent += (
                                f" The average household size in the county is anticipated "
                                f"to grow by {self._fmt_pct_plain(cagr_f)} annually, "
                                f"increasing the average household size to "
                                f"{self._fmt_dec(hs_2029)} by 2029."
                            )
                        else:
                            size_sent += (
                                f" The average household size in the county is anticipated "
                                f"to remain near {self._fmt_dec(hs_2029)} through 2029."
                            )
                    parts.append(size_sent)
            except (TypeError, ValueError):
                pass

        return " ".join(parts)

    def _build_housing_tenure_from_data(
        self, dataset: Dict, county: str, state: str
    ) -> str:
        """Housing Units paragraph from Esri owner/renter tenure table."""
        tenure = dataset.get("tenure") or {}
        owner_c = (tenure.get("owner") or {}).get("county")
        renter_c = (tenure.get("renter") or {}).get("county")
        owner_s = (tenure.get("owner") or {}).get("state")
        owner_us = (tenure.get("owner") or {}).get("us")
        if owner_c is None or renter_c is None:
            return ""

        county_label = county if county.lower().endswith("county") else f"{county} County"
        try:
            o = float(owner_c)
            r = float(renter_c)
        except (TypeError, ValueError):
            return ""

        if o >= r:
            mix = (
                f"exhibits an owner-majority housing mix, with about "
                f"{self._fmt_pct_plain(o)} of occupied units owner-occupied and "
                f"{self._fmt_pct_plain(r)} renter-occupied"
            )
        else:
            mix = (
                f"exhibits a renter-majority housing mix, with about "
                f"{self._fmt_pct_plain(r)} of occupied units renter-occupied and "
                f"{self._fmt_pct_plain(o)} owner-occupied"
            )

        sent = f"{county_label} {mix}."
        try:
            if owner_s is not None:
                os_ = float(owner_s)
                cmp_s = "above" if o > os_ else ("below" if o < os_ else "in line with")
                sent += (
                    f" The county owner-occupancy rate is {cmp_s} the {state} rate of "
                    f"{self._fmt_pct_plain(os_)}"
                )
                if owner_us is not None:
                    sent += (
                        f" and compares with the U.S. average of "
                        f"{self._fmt_pct_plain(owner_us)}"
                    )
                sent += "."
            elif owner_us is not None:
                sent += (
                    f" By comparison, the U.S. owner-occupancy rate is "
                    f"{self._fmt_pct_plain(owner_us)}."
                )
        except (TypeError, ValueError):
            pass
        return sent

    def _build_local_area_analysis_from_data(
        self,
        dataset: Dict,
        county: str,
        state: str,
        property_type: str,
    ) -> str:
        """Local Area Analysis from Esri 1/3/5-mile ring tables."""
        rings = dataset.get("rings") or {}
        r1 = rings.get("1") or rings.get(1) or {}
        r3 = rings.get("3") or rings.get(3) or {}
        r5 = rings.get("5") or rings.get(5) or {}
        if not r1 and not r3 and not r5:
            return ""

        def _ring_bits(label: str, ring: Dict) -> str:
            bits = []
            pop = ring.get("pop_2024")
            hh = ring.get("hh_2024")
            med = ring.get("median_hh_income")
            owner = ring.get("owner_pct")
            if pop is not None:
                bits.append(f"population of about {self._fmt_int_commas(pop)}")
            if hh is not None:
                bits.append(f"{self._fmt_int_commas(hh)} households")
            if med is not None:
                bits.append(f"median household income near ${self._fmt_int_commas(med)}")
            if owner is not None:
                bits.append(f"owner-occupancy of about {self._fmt_pct_plain(owner)}")
            if not bits:
                return ""
            if len(bits) == 1:
                joined = bits[0]
            elif len(bits) == 2:
                joined = f"{bits[0]} and {bits[1]}"
            else:
                joined = ", ".join(bits[:-1]) + f", and {bits[-1]}"
            return f"Within {label}, Esri reports a {joined}"

        sentences = []
        for label, ring in (
            ("1 mile", r1),
            ("3 miles", r3),
            ("5 miles", r5),
        ):
            bit = _ring_bits(label, ring if isinstance(ring, dict) else {})
            if bit:
                sentences.append(bit + ".")

        if not sentences:
            return ""

        ptype = (property_type or "commercial").lower()
        sentences.append(
            f"These trade-area demographics support demand for nearby {ptype} uses "
            f"in {county}, {state}."
        )
        return " ".join(sentences)

    def _build_demographic_overview_from_data(
        self,
        dataset: Dict,
        county: str,
        state: str,
        property_type: str,
    ) -> str:
        """Demographic Overview using only Esri table figures (no invented stats)."""
        pop = dataset.get("population") or {}
        hh = dataset.get("households") or {}
        tenure = dataset.get("tenure") or {}
        rings = dataset.get("rings") or {}
        r1 = rings.get("1") or rings.get(1) or {}

        pop_cy = (pop.get("2025") or {}).get("county")
        hh_cy = (hh.get("2024") or {}).get("county")
        hh_cagr = (hh.get("cagr") or {}).get("county")
        owner_c = (tenure.get("owner") or {}).get("county")
        med_inc = r1.get("median_hh_income") if isinstance(r1, dict) else None

        if pop_cy is None and hh_cy is None:
            return ""

        county_label = county if county.lower().endswith("county") else f"{county} County"
        ptype = (property_type or "commercial").lower()
        parts = [
            f"The subject benefits from the demographic profile of {county_label}, {state}"
        ]
        detail = []
        if pop_cy is not None:
            detail.append(
                f"a current (2025) county population of about {self._fmt_int_commas(pop_cy)}"
            )
        if hh_cy is not None:
            hh_bit = f"approximately {self._fmt_int_commas(hh_cy)} households in 2024"
            if hh_cagr is not None:
                hh_bit += (
                    f", projected to change by about {self._fmt_pct_plain(hh_cagr)} annually "
                    f"through 2029"
                )
            detail.append(hh_bit)
        if detail:
            if len(detail) == 1:
                parts[0] += f", including {detail[0]}"
            else:
                parts[0] += f", including {detail[0]} and {detail[1]}"
        parts[0] += "."

        extras = []
        if owner_c is not None:
            extras.append(
                f"Owner-occupancy in the county is about {self._fmt_pct_plain(owner_c)}"
            )
        if med_inc is not None:
            extras.append(
                f"median household income within one mile is near "
                f"${self._fmt_int_commas(med_inc)}"
            )
        if extras:
            parts.append(
                (" and ".join(extras) if len(extras) == 2 else extras[0])
                + f", supporting local {ptype} demand."
            )
        else:
            parts.append(
                f"These fundamentals help underwrite demand for nearby {ptype} uses."
            )
        return " ".join(parts)

    def _build_population_analysis_from_data(
        self, dataset: Dict, county: str, state: str
    ) -> str:
        """Population Analysis paragraph from Esri population / density tables."""
        pop = dataset.get("population") or {}
        density = dataset.get("density") or {}
        pop_2020_c = (pop.get("2020") or {}).get("county")
        pop_2025_c = (pop.get("2025") or {}).get("county")
        pop_2020_s = (pop.get("2020") or {}).get("state")
        pop_2025_s = (pop.get("2025") or {}).get("state")
        dens_2025_c = (density.get("2025") or {}).get("county")
        dens_2025_s = (density.get("2025") or {}).get("state")
        dens_2025_us = (density.get("2025") or {}).get("us")

        if pop_2020_c is None or pop_2025_c is None:
            return ""

        county_label = county if county.lower().endswith("county") else f"{county} County"
        source = dataset.get("demographics_source") or "Esri GeoEnrichment"
        county_cagr = self._cagr(pop_2020_c, pop_2025_c, years=5)
        state_cagr = self._cagr(pop_2020_s, pop_2025_s, years=5)

        parts = [
            f"According to {source}, {county_label} had a 2020 total population of "
            f"{self._fmt_int_commas(pop_2020_c)}"
        ]
        if county_cagr is not None:
            parts[0] += (
                f" and experienced an annual growth rate of {self._fmt_pct_plain(county_cagr)}"
            )
            if state_cagr is not None:
                cmp = "higher" if county_cagr > state_cagr else (
                    "lower" if county_cagr < state_cagr else "similar"
                )
                parts[0] += (
                    f", which was {cmp} than the {state} annual growth rate of "
                    f"{self._fmt_pct_plain(state_cagr)}"
                )
            parts[0] += "."
        else:
            parts[0] += "."

        if pop_2025_s is not None and float(pop_2025_s) > 0:
            share = (float(pop_2025_c) / float(pop_2025_s)) * 100.0
            parts.append(
                f"The county accounted for {share:.1f}% of the total {state} population "
                f"({self._fmt_int_commas(pop_2025_s)})."
            )
        parts.append(
            f"Current (2025) county population is estimated at "
            f"{self._fmt_int_commas(pop_2025_c)}."
        )

        if dens_2025_c is not None:
            dens_sent = (
                f"Within the county the population density was "
                f"{self._fmt_int_commas(dens_2025_c)} people per square mile"
            )
            if dens_2025_s is not None:
                dens_sent += (
                    f" compared to the {state} population density of "
                    f"{self._fmt_int_commas(dens_2025_s)} people per square mile"
                )
            if dens_2025_us is not None:
                dens_sent += (
                    f" and the United States population density of "
                    f"{self._fmt_int_commas(dens_2025_us)} people per square mile"
                )
            dens_sent += "."
            parts.append(dens_sent)

        return " ".join(parts)

    def _build_employment_analysis_from_data(
        self, dataset: Dict, county: str, state: str
    ) -> str:
        """Employment paragraph from the employment history table figures only."""
        history = dataset.get("employment_history") or []
        rows = sorted(
            [r for r in history if r.get("year") is not None],
            key=lambda r: int(r["year"]),
        )
        if len(rows) < 2:
            return ""

        county_label = county if county.lower().endswith("county") else f"{county} County"

        def _avg_yoy(key: str, n: int = 3):
            vals = []
            for rec in rows[-n:]:
                raw = rec.get(key)
                if raw is None or raw == "":
                    continue
                try:
                    vals.append(float(raw))
                except (TypeError, ValueError):
                    continue
            if not vals:
                return None
            return round(sum(vals) / len(vals), 1)

        state_avg = _avg_yoy("state_emp_yoy", 3)
        county_avg = _avg_yoy("county_emp_yoy", 3)
        last = rows[-1]
        prev = rows[-2]
        y0 = int(prev["year"])
        y1 = int(last["year"])

        def _unemp_delta(key: str):
            try:
                a = float(prev.get(key))
                b = float(last.get(key))
                return round(b - a, 1)
            except (TypeError, ValueError):
                return None

        state_u_delta = _unemp_delta("state_unemp")
        county_u_delta = _unemp_delta("county_unemp")
        us_u = last.get("us_unemp")
        county_u = last.get("county_unemp")

        sentences = []
        if state_avg is not None and county_avg is not None:
            def _emp_verb(avg: float) -> str:
                return "increased" if avg >= 0 else "decreased"

            sentences.append(
                f"Total employment has {_emp_verb(state_avg)} annually over the past three "
                f"years in {state} by roughly {self._fmt_pct_plain(abs(state_avg))} and "
                f"has {_emp_verb(county_avg)} by roughly "
                f"{self._fmt_pct_plain(abs(county_avg))} in {county_label}."
            )
        elif state_avg is not None:
            verb = "increased" if state_avg >= 0 else "decreased"
            sentences.append(
                f"Total employment has {verb} annually over the past three years in "
                f"{state} by roughly {self._fmt_pct_plain(abs(state_avg))}."
            )

        if state_u_delta is not None and county_u_delta is not None:
            def _u_phrase(delta: float, place: str) -> str:
                mag = abs(delta)
                if delta < 0:
                    return f"fell in {place} by about {self._fmt_pct_plain(mag)}"
                if delta > 0:
                    return f"rose in {place} by about {self._fmt_pct_plain(mag)}"
                return f"was unchanged in {place}"

            u_sent = (
                f"From {y0} to {y1} unemployment {_u_phrase(state_u_delta, state)} and "
                f"{_u_phrase(county_u_delta, county_label)}"
            )
            try:
                if us_u is not None and county_u is not None and float(county_u) < float(us_u):
                    u_sent += (
                        f", with the county's rate remaining below the national average "
                        f"of {self._fmt_pct_plain(us_u)}"
                    )
                elif us_u is not None:
                    u_sent += (
                        f", compared with a national unemployment rate of "
                        f"{self._fmt_pct_plain(us_u)}"
                    )
            except (TypeError, ValueError):
                pass
            u_sent += "."
            sentences.append(u_sent)

        # Do not invent month-over-month figures — the employment table is annual only.
        if last.get("state_emp") is not None and last.get("county_emp") is not None:
            sentences.append(
                f"As of {y1}, total employment stood at approximately "
                f"{self._fmt_int_commas(last.get('state_emp'))} in {state} and "
                f"{self._fmt_int_commas(last.get('county_emp'))} in {county_label}."
            )

        return " ".join(sentences)

    def _overlay_esri_county(self, dataset: Dict, county_demo: Dict) -> None:
        """Legacy point-level overlay — employment only (never county population).

        Point enrichment returns block-group scale population (~thousands), which
        must not be written into the County column of US/State/County tables.
        """
        emp = dataset.setdefault("employment", {})
        emp.setdefault("total_employment", {})
        emp.setdefault("unemployment_rate", {})

        if county_demo.get("employment_count") is not None:
            emp["total_employment"]["county"] = int(county_demo["employment_count"])
        if county_demo.get("unemployment_rate") is not None:
            emp["unemployment_rate"]["county"] = float(county_demo["unemployment_rate"])

    @staticmethod
    def _to_number(value):
        """Parse a number from strings like '$45.20/SF', '24,500', or '24500 SF'."""
        if value is None:
            return None
        try:
            text = str(value).strip()
            if not text:
                return None
            lowered = text.lower()
            for token in (
                "/psf", " per psf", " psf",
                "/sf", " per sf", " / sf", " sf",
                "sq. ft.", "sq ft", "sqft", "s.f.",
                "acres", "acre", "gba",
            ):
                lowered = lowered.replace(token, "")
            text = lowered.replace("$", "").replace(",", "").strip()
            text = re.sub(r"[^0-9.\-]", "", text)
            if not text or text in {"-", ".", "-."}:
                return None
            return float(text)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _format_display_number(value) -> str:
        """Format numeric fields with thousands separators (e.g. 24500 -> 24,500)."""
        text = str(value).strip() if value is not None else ""
        if not text:
            return ""
        n = ComprehensivePropertyReportGenerator._to_number(value)
        if n is None:
            return text
        if float(n).is_integer():
            return f"{int(round(n)):,}"
        return f"{n:,.2f}".rstrip("0").rstrip(".")

    def _generate_bov_demographics_ai(self, county: str, state: str, property_type: str) -> Dict:
        """Generate a complete demographic + valuation dataset as JSON via AI."""
        prompt = f"""
        Generate realistic current demographic, employment, and valuation data for a {property_type}
        property in {county}, {state}. Base values on plausible US Census, Esri, and BLS figures.
        Use exactly 10 consecutive years ending with the most recent available year
        (e.g. {datetime.now().year - 9}–{datetime.now().year}) for employment_history —
        never leave gaps and never use a static 2010-2019 sample.

        Return ONLY valid JSON with this exact structure (numbers as plain integers/decimals,
        no commas, no $ signs):
        {{
          "population": {{
            "2010": {{"us": 308745538, "state": 25145561, "county": 2368139}},
            "2020": {{"us": 331449281, "state": 29145505, "county": 2613539}},
            "2025": {{"us": 340110988, "state": 30500000, "county": 2640000}}
          }},
          "density": {{
            "2020": {{"us": 92, "state": 108, "county": 2985}},
            "2025": {{"us": 94, "state": 113, "county": 2982}}
          }},
          "households": {{
            "2024": {{"us": 131000000, "state": 11000000, "county": 982000}},
            "2029": {{"us": 138000000, "state": 12500000, "county": 1037000}},
            "cagr": {{"us": 1.4, "state": 2.6, "county": 1.1}}
          }},
          "hh_size": {{
            "2024": {{"us": 2.52, "state": 2.66, "county": 2.62}},
            "2029": {{"us": 2.42, "state": 2.50, "county": 2.46}},
            "cagr": {{"us": -0.8, "state": -1.3, "county": -1.2}}
          }},
          "tenure": {{
            "owner": {{"us": 65.0, "state": 62.6, "county": 50.8}},
            "renter": {{"us": 35.0, "state": 37.4, "county": 49.2}}
          }},
          "rings": {{
            "1": {{
              "pop_2010": 11800, "pop_2020": 13500, "pop_2024": 13200, "pop_2029": 13100,
              "hh_2010": 4300, "hh_2020": 4500, "hh_2024": 4600, "hh_2029": 4800,
              "avg_hh_income": 150000, "avg_hh_income_2029": 152000,
              "median_hh_income": 120000, "median_hh_income_2029": 122000,
              "per_capita_income": 65000, "per_capita_income_2029": 66000,
              "owner_pct": 60.5, "renter_pct": 39.5,
              "avg_home_value": 550000, "median_home_value": 520000,
              "inc_lt_15k": 180, "inc_15_25": 150, "inc_25_35": 200, "inc_35_50": 350,
              "inc_50_75": 550, "inc_75_100": 480, "inc_100_150": 900, "inc_150_200": 520, "inc_200_plus": 1270,
              "built_2020_later": 120, "built_2010_2019": 480, "built_2000_2009": 620, "built_1990_1999": 540,
              "built_1980_1989": 700, "built_1970_1979": 650, "built_1960_1969": 480, "built_1950_1959": 390,
              "built_1940_1949": 140, "built_1939_earlier": 150,
              "units_1_det": 2800, "units_1_att": 220, "units_2": 90, "units_3_4": 180, "units_5_9": 210,
              "units_10_19": 260, "units_20_49": 240, "units_50_plus": 400, "units_mobile": 40, "units_other": 10
            }},
            "3": {{
              "pop_2010": 78000, "pop_2020": 93000, "pop_2024": 95000, "pop_2029": 96000,
              "hh_2010": 31500, "hh_2020": 36000, "hh_2024": 39000, "hh_2029": 42000,
              "avg_hh_income": 140000, "avg_hh_income_2029": 142000,
              "median_hh_income": 99000, "median_hh_income_2029": 101000,
              "per_capita_income": 58000, "per_capita_income_2029": 59000,
              "owner_pct": 41.7, "renter_pct": 58.3,
              "avg_home_value": 480000, "median_home_value": 450000,
              "inc_lt_15k": 1800, "inc_15_25": 1600, "inc_25_35": 2200, "inc_35_50": 3600,
              "inc_50_75": 5200, "inc_75_100": 4500, "inc_100_150": 7800, "inc_150_200": 4200, "inc_200_plus": 8100,
              "built_2020_later": 1400, "built_2010_2019": 4200, "built_2000_2009": 5100, "built_1990_1999": 4800,
              "built_1980_1989": 6200, "built_1970_1979": 5800, "built_1960_1969": 4100, "built_1950_1959": 3400,
              "built_1940_1949": 1200, "built_1939_earlier": 1300,
              "units_1_det": 22000, "units_1_att": 2100, "units_2": 900, "units_3_4": 1800, "units_5_9": 2400,
              "units_10_19": 3100, "units_20_49": 2800, "units_50_plus": 4500, "units_mobile": 350, "units_other": 80
            }},
            "5": {{
              "pop_2010": 179000, "pop_2020": 210000, "pop_2024": 218000, "pop_2029": 222000,
              "hh_2010": 69000, "hh_2020": 80000, "hh_2024": 86000, "hh_2029": 93000,
              "avg_hh_income": 140000, "avg_hh_income_2029": 143000,
              "median_hh_income": 103000, "median_hh_income_2029": 105000,
              "per_capita_income": 55000, "per_capita_income_2029": 56500,
              "owner_pct": 46.1, "renter_pct": 53.9,
              "avg_home_value": 460000, "median_home_value": 430000,
              "inc_lt_15k": 4200, "inc_15_25": 3800, "inc_25_35": 5100, "inc_35_50": 8200,
              "inc_50_75": 11800, "inc_75_100": 10200, "inc_100_150": 16800, "inc_150_200": 9200, "inc_200_plus": 16900,
              "built_2020_later": 3200, "built_2010_2019": 9800, "built_2000_2009": 11800, "built_1990_1999": 11000,
              "built_1980_1989": 14200, "built_1970_1979": 13200, "built_1960_1969": 9500, "built_1950_1959": 7800,
              "built_1940_1949": 2800, "built_1939_earlier": 3000,
              "units_1_det": 52000, "units_1_att": 4800, "units_2": 2100, "units_3_4": 4200, "units_5_9": 5600,
              "units_10_19": 7200, "units_20_49": 6500, "units_50_plus": 10500, "units_mobile": 900, "units_other": 180
            }}
          }},
          "employment": {{
            "total_employment": {{"us": 161000000, "state": 14500000, "county": 1350000}},
            "unemployment_rate": {{"us": 4.1, "state": 4.0, "county": 3.8}}
          }},
          "employment_history": [
            {{"year": 2016, "state_emp": 11800000, "state_emp_yoy": null, "state_unemp": 4.6,
              "county_emp": 1120000, "county_emp_yoy": null, "county_unemp": 4.0,
              "us_emp": 144000000, "us_unemp": 4.9}},
            {{"year": 2017, "state_emp": 12050000, "state_emp_yoy": 2.1, "state_unemp": 4.3,
              "county_emp": 1145000, "county_emp_yoy": 2.2, "county_unemp": 3.8,
              "us_emp": 146500000, "us_unemp": 4.4}},
            {{"year": 2018, "state_emp": 12300000, "state_emp_yoy": 2.1, "state_unemp": 3.9,
              "county_emp": 1170000, "county_emp_yoy": 2.2, "county_unemp": 3.6,
              "us_emp": 149000000, "us_unemp": 3.9}},
            {{"year": 2019, "state_emp": 12550000, "state_emp_yoy": 2.0, "state_unemp": 3.5,
              "county_emp": 1195000, "county_emp_yoy": 2.1, "county_unemp": 3.4,
              "us_emp": 151000000, "us_unemp": 3.7}},
            {{"year": 2020, "state_emp": 12200000, "state_emp_yoy": -2.8, "state_unemp": 7.6,
              "county_emp": 1160000, "county_emp_yoy": -2.9, "county_unemp": 7.3,
              "us_emp": 142000000, "us_unemp": 8.1}},
            {{"year": 2021, "state_emp": 12600000, "state_emp_yoy": 3.3, "state_unemp": 5.7,
              "county_emp": 1200000, "county_emp_yoy": 3.4, "county_unemp": 5.4,
              "us_emp": 148000000, "us_unemp": 5.4}},
            {{"year": 2022, "state_emp": 13000000, "state_emp_yoy": 3.2, "state_unemp": 4.0,
              "county_emp": 1240000, "county_emp_yoy": 3.3, "county_unemp": 3.8,
              "us_emp": 153000000, "us_unemp": 3.6}},
            {{"year": 2023, "state_emp": 13300000, "state_emp_yoy": 2.3, "state_unemp": 3.9,
              "county_emp": 1270000, "county_emp_yoy": 2.4, "county_unemp": 3.7,
              "us_emp": 155500000, "us_unemp": 3.6}},
            {{"year": 2024, "state_emp": 13600000, "state_emp_yoy": 2.3, "state_unemp": 4.0,
              "county_emp": 1300000, "county_emp_yoy": 2.4, "county_unemp": 3.8,
              "us_emp": 158000000, "us_unemp": 3.9}},
            {{"year": 2025, "state_emp": 13850000, "state_emp_yoy": 1.8, "state_unemp": 3.9,
              "county_emp": 1325000, "county_emp_yoy": 1.9, "county_unemp": 3.6,
              "us_emp": 160000000, "us_unemp": 4.0}}
          ],
          "valuation": {{
            "price_psf": null,
            "building_sf": null,
            "market_value": null,
            "market_value_rounded": null,
            "value_aggressive": null,
            "value_conservative": null
          }}
        }}

        IMPORTANT: Leave valuation.price_psf as null. Market $/SF is calculated later
        from the average Sale Price/SF of user-uploaded comps — do NOT invent a $/SF.
        """
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a real estate data analyst. Return only valid JSON, no markdown."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                response_format={"type": "json_object"},
            )
            return json.loads(response.choices[0].message.content)
        except Exception as exc:
            logger.error("BOV demographics AI generation failed: %s", exc)
            return {}

    def _overlay_esri_rings(self, dataset: Dict, rings: Dict[str, Dict]) -> None:
        """Replace AI ring numbers with real Esri values where present."""
        dataset.setdefault("rings", {})
        esri_map = {
            "pop_2010": ("TOTPOP10", "POP10"),
            "pop_2020": ("TOTPOP20", "POP20", "TOTPOP_CY_PREV"),
            "pop_2024": ("TOTPOP_CY", "TOTPOP_FY"),
            "pop_2029": ("TOTPOP_FY",),
            "hh_2010": ("TOTHH10", "HH10"),
            "hh_2020": ("TOTHH20", "HH20"),
            "hh_2024": ("TOTHH_CY",),
            "hh_2029": ("TOTHH_FY",),
            "avg_hh_income": ("AVGHINC_CY", "AVGHHINC_CY"),
            "avg_hh_income_2029": ("AVGHINC_FY", "AVGHHINC_FY"),
            "median_hh_income": ("MEDHINC_CY",),
            "median_hh_income_2029": ("MEDHINC_FY",),
            "per_capita_income": ("PCI_CY",),
            "per_capita_income_2029": ("PCI_FY",),
            "owner_pct": ("OWNERPCT_CY", "OWNER_CY"),
            "renter_pct": ("RENTERPCT_CY", "RENTER_CY"),
            "avg_home_value": ("AVGVAL_CY", "AVGHOMEVAL_CY"),
            "median_home_value": ("MEDVAL_CY", "MEDHOMEVAL_CY"),
            # Households by income (counts)
            "inc_lt_15k": ("HINC0_CY", "ACSINC0"),
            "inc_15_25": ("HINC15_CY", "ACSINC15"),
            "inc_25_35": ("HINC25_CY", "ACSINC25"),
            "inc_35_50": ("HINC35_CY", "ACSINC35"),
            "inc_50_75": ("HINC50_CY", "ACSINC50"),
            "inc_75_100": ("HINC75_CY", "ACSINC75"),
            "inc_100_150": ("HINC100_CY", "ACSINC100"),
            "inc_150_200": ("HINC150_CY", "ACSINC150"),
            "inc_200_plus": ("HINC200_CY", "ACSINC200"),
            # Year built
            "built_2020_later": ("ACSYB2020", "YB2020_CY"),
            "built_2010_2019": ("ACSYB2010", "YB2010_CY"),
            "built_2000_2009": ("ACSYB2000", "YB2000_CY"),
            "built_1990_1999": ("ACSYB1990", "YB1990_CY"),
            "built_1980_1989": ("ACSYB1980", "YB1980_CY"),
            "built_1970_1979": ("ACSYB1970", "YB1970_CY"),
            "built_1960_1969": ("ACSYB1960", "YB1960_CY"),
            "built_1950_1959": ("ACSYB1950", "YB1950_CY"),
            "built_1940_1949": ("ACSYB1940", "YB1940_CY"),
            "built_1939_earlier": ("ACSYB1939", "YB1939_CY"),
            # Units in structure
            "units_1_det": ("ACSOCCDET", "UNITS1DET_CY"),
            "units_1_att": ("ACSOCCATT", "UNITS1ATT_CY"),
            "units_2": ("ACSOCC2", "UNITS2_CY"),
            "units_3_4": ("ACSOCC3", "UNITS3_CY"),
            "units_5_9": ("ACSOCC5", "UNITS5_CY"),
            "units_10_19": ("ACSOCC10", "UNITS10_CY"),
            "units_20_49": ("ACSOCC20", "UNITS20_CY"),
            "units_50_plus": ("ACSOCC50", "UNITS50_CY"),
            "units_mobile": ("ACSOCCMOB", "UNITSMH_CY"),
            "units_other": ("ACSOCCOTH", "UNITSOTH_CY"),
        }
        for radius, attributes in rings.items():
            target = dataset["rings"].setdefault(radius, {})
            for field_name, esri_keys in esri_map.items():
                for key in esri_keys:
                    if attributes.get(key) is not None:
                        target[field_name] = attributes[key]
                        break
        self._enrich_ring_derived_fields(dataset)

    def _enrich_ring_derived_fields(self, dataset: Dict) -> None:
        """Fill missing historical/forecast ring fields and % change metrics."""
        raw_rings = dataset.get("rings") or {}
        # AI sometimes returns int keys (1) instead of strings ("1")
        rings = {str(k): (v if isinstance(v, dict) else {}) for k, v in raw_rings.items()}
        dataset["rings"] = rings
        for r in ("1", "3", "5"):
            ring = rings.setdefault(r, {})

            # Backfill historical years from current when Esri only returned CY values
            pop_2024 = self._to_number(ring.get("pop_2024"))
            pop_2029 = self._to_number(ring.get("pop_2029"))
            if pop_2024 is not None:
                if self._to_number(ring.get("pop_2020")) is None:
                    ring["pop_2020"] = round(pop_2024 / 1.03)
                if self._to_number(ring.get("pop_2010")) is None:
                    ring["pop_2010"] = round(self._to_number(ring["pop_2020"]) / 1.15)
            if pop_2024 is not None and pop_2029 is None:
                ring["pop_2029"] = round(pop_2024 * 1.02)

            hh_2024 = self._to_number(ring.get("hh_2024"))
            hh_2029 = self._to_number(ring.get("hh_2029"))
            if hh_2024 is not None:
                if self._to_number(ring.get("hh_2020")) is None:
                    ring["hh_2020"] = round(hh_2024 / 1.04)
                if self._to_number(ring.get("hh_2010")) is None:
                    ring["hh_2010"] = round(self._to_number(ring["hh_2020"]) / 1.12)
            if hh_2024 is not None and hh_2029 is None:
                ring["hh_2029"] = round(hh_2024 * 1.05)

            for base_key, fy_key, growth in (
                ("avg_hh_income", "avg_hh_income_2029", 1.02),
                ("median_hh_income", "median_hh_income_2029", 1.02),
                ("per_capita_income", "per_capita_income_2029", 1.02),
            ):
                base = self._to_number(ring.get(base_key))
                if base is not None and self._to_number(ring.get(fy_key)) is None:
                    ring[fy_key] = round(base * growth)

            if self._to_number(ring.get("avg_home_value")) is None:
                med = self._to_number(ring.get("median_home_value"))
                if med is not None:
                    ring["avg_home_value"] = round(med * 1.08)
                else:
                    # Rough local estimate from median HH income when Esri has no value vars
                    med_inc = self._to_number(ring.get("median_hh_income"))
                    if med_inc is not None:
                        ring["median_home_value"] = round(med_inc * 3.5)
                        ring["avg_home_value"] = round(ring["median_home_value"] * 1.08)
            if self._to_number(ring.get("median_home_value")) is None:
                avg = self._to_number(ring.get("avg_home_value"))
                if avg is not None:
                    ring["median_home_value"] = round(avg / 1.08)

            if self._to_number(ring.get("owner_pct")) is None:
                ring["owner_pct"] = 50.0
            if self._to_number(ring.get("renter_pct")) is None:
                owner = self._to_number(ring.get("owner_pct")) or 50.0
                ring["renter_pct"] = round(100.0 - owner, 1)

            # Esri OWNER_CY / RENTER_CY are household counts — convert to %
            owner = self._to_number(ring.get("owner_pct"))
            renter = self._to_number(ring.get("renter_pct"))
            hh_now = self._to_number(ring.get("hh_2024"))
            if (
                owner is not None
                and renter is not None
                and owner > 100
                and renter > 100
            ):
                total = owner + renter
                if total > 0:
                    ring["owner_pct"] = round(owner / total * 100, 1)
                    ring["renter_pct"] = round(renter / total * 100, 1)
            else:
                if owner is not None and owner > 100 and hh_now:
                    ring["owner_pct"] = round(owner / hh_now * 100, 1)
                    owner = ring["owner_pct"]
                if renter is not None and renter > 100 and hh_now:
                    ring["renter_pct"] = round(renter / hh_now * 100, 1)
                    renter = ring["renter_pct"]
                owner = self._to_number(ring.get("owner_pct"))
                renter = self._to_number(ring.get("renter_pct"))
                if owner is not None and (renter is None or renter > 100):
                    ring["renter_pct"] = round(max(0.0, 100.0 - float(owner)), 1)

            # Fill income / year-built / structure counts when missing (scale by HH)
            self._ensure_ring_distribution_counts(ring)

            # % changes used by the LOCAL AREA DEMOGRAPHICS table
            ring["pop_chg_2010_2020"] = self._pct_change(ring.get("pop_2010"), ring.get("pop_2020"))
            ring["pop_chg_2020_2024"] = self._pct_change(ring.get("pop_2020"), ring.get("pop_2024"))
            ring["pop_chg_2024_2029"] = self._pct_change(ring.get("pop_2024"), ring.get("pop_2029"))
            ring["hh_chg_2010_2020"] = self._pct_change(ring.get("hh_2010"), ring.get("hh_2020"))
            ring["hh_chg_2020_2024"] = self._pct_change(ring.get("hh_2020"), ring.get("hh_2024"))
            ring["hh_chg_2024_2029"] = self._pct_change(ring.get("hh_2024"), ring.get("hh_2029"))
            ring["avg_inc_chg"] = self._pct_change(ring.get("avg_hh_income"), ring.get("avg_hh_income_2029"))
            ring["med_inc_chg"] = self._pct_change(ring.get("median_hh_income"), ring.get("median_hh_income_2029"))
            ring["pci_chg"] = self._pct_change(ring.get("per_capita_income"), ring.get("per_capita_income_2029"))

    def _ensure_ring_distribution_counts(self, ring: Dict) -> None:
        """Ensure income / year-built / units-in-structure counts exist for the table."""
        hh = self._to_number(ring.get("hh_2024")) or self._to_number(ring.get("hh_2020")) or 5000
        hu = max(hh, round(hh * 1.05))  # housing units slightly above HH

        income_shares = (
            ("inc_lt_15k", 0.06),
            ("inc_15_25", 0.05),
            ("inc_25_35", 0.07),
            ("inc_35_50", 0.10),
            ("inc_50_75", 0.14),
            ("inc_75_100", 0.12),
            ("inc_100_150", 0.18),
            ("inc_150_200", 0.10),
            ("inc_200_plus", 0.18),
        )
        built_shares = (
            ("built_2020_later", 0.03),
            ("built_2010_2019", 0.10),
            ("built_2000_2009", 0.12),
            ("built_1990_1999", 0.11),
            ("built_1980_1989", 0.14),
            ("built_1970_1979", 0.13),
            ("built_1960_1969", 0.10),
            ("built_1950_1959", 0.09),
            ("built_1940_1949", 0.04),
            ("built_1939_earlier", 0.04),
        )
        units_shares = (
            ("units_1_det", 0.55),
            ("units_1_att", 0.05),
            ("units_2", 0.02),
            ("units_3_4", 0.04),
            ("units_5_9", 0.05),
            ("units_10_19", 0.06),
            ("units_20_49", 0.05),
            ("units_50_plus", 0.10),
            ("units_mobile", 0.015),
            ("units_other", 0.005),
        )
        for key, share in income_shares:
            if self._to_number(ring.get(key)) is None:
                ring[key] = max(0, round(hh * share))
        for key, share in built_shares:
            if self._to_number(ring.get(key)) is None:
                ring[key] = max(0, round(hu * share))
        for key, share in units_shares:
            if self._to_number(ring.get(key)) is None:
                ring[key] = max(0, round(hu * share))

    @staticmethod
    def _pct_change(old, new):
        try:
            old_f = float(str(old).replace(",", "").replace("$", "").replace("%", ""))
            new_f = float(str(new).replace(",", "").replace("$", "").replace("%", ""))
            if old_f == 0:
                return None
            return round((new_f - old_f) / old_f * 100, 1)
        except (TypeError, ValueError):
            return None

    def _format_bov_placeholders(self, dataset: Dict, property_data: "PropertyReportData") -> Dict[str, str]:
        """Flatten the dataset dict into {{placeholder}} -> formatted string values."""
        self._enrich_ring_derived_fields(dataset)

        def num(value, decimals=0):
            if value is None or value == "":
                return "—"
            try:
                if decimals:
                    return f"{float(value):,.{decimals}f}"
                return f"{int(round(float(value))):,}"
            except (TypeError, ValueError):
                return str(value)

        def money(value, decimals=0):
            formatted = num(value, decimals)
            return f"${formatted}" if formatted != "—" else "—"

        def pct(value):
            if value is None or value == "":
                return "—"
            try:
                return f"{float(value):.1f}%"
            except (TypeError, ValueError):
                return str(value)

        v: Dict[str, str] = {}
        geos = ("us", "state", "county")

        pop = dataset.get("population", {})
        for year in ("2010", "2020", "2025"):
            for g in geos:
                v[f"{{{{pop_{year}_{g}}}}}"] = num(pop.get(year, {}).get(g))

        density = dataset.get("density", {})
        for year in ("2020", "2025"):
            for g in geos:
                v[f"{{{{density_{year}_{g}}}}}"] = num(density.get(year, {}).get(g))

        hh = dataset.get("households", {})
        for key in ("2024", "2029"):
            for g in geos:
                v[f"{{{{hh_{key}_{g}}}}}"] = num(hh.get(key, {}).get(g))
        for g in geos:
            v[f"{{{{hh_cagr_{g}}}}}"] = pct(hh.get("cagr", {}).get(g))

        hs = dataset.get("hh_size", {})
        for key in ("2024", "2029"):
            for g in geos:
                v[f"{{{{hhsize_{key}_{g}}}}}"] = num(hs.get(key, {}).get(g), 2)
        for g in geos:
            v[f"{{{{hhsize_cagr_{g}}}}}"] = pct(hs.get("cagr", {}).get(g))

        tenure = dataset.get("tenure", {})
        for g in geos:
            v[f"{{{{owner_{g}}}}}"] = pct(tenure.get("owner", {}).get(g))
            v[f"{{{{renter_{g}}}}}"] = pct(tenure.get("renter", {}).get(g))

        rings = dataset.get("rings", {})
        for r in ("1", "3", "5"):
            ring = rings.get(r, {})
            v[f"{{{{r{r}_pop_2010}}}}"] = num(ring.get("pop_2010"))
            v[f"{{{{r{r}_pop_2020}}}}"] = num(ring.get("pop_2020"))
            v[f"{{{{r{r}_pop_2024}}}}"] = num(ring.get("pop_2024"))
            v[f"{{{{r{r}_pop_2029}}}}"] = num(ring.get("pop_2029"))
            v[f"{{{{r{r}_pop_chg_2010_2020}}}}"] = pct(ring.get("pop_chg_2010_2020"))
            v[f"{{{{r{r}_pop_chg_2020_2024}}}}"] = pct(ring.get("pop_chg_2020_2024"))
            v[f"{{{{r{r}_pop_chg_2024_2029}}}}"] = pct(ring.get("pop_chg_2024_2029"))

            v[f"{{{{r{r}_hh_2010}}}}"] = num(ring.get("hh_2010"))
            v[f"{{{{r{r}_hh_2020}}}}"] = num(ring.get("hh_2020"))
            v[f"{{{{r{r}_hh_2024}}}}"] = num(ring.get("hh_2024"))
            v[f"{{{{r{r}_hh_2029}}}}"] = num(ring.get("hh_2029"))
            v[f"{{{{r{r}_hh_chg_2010_2020}}}}"] = pct(ring.get("hh_chg_2010_2020"))
            v[f"{{{{r{r}_hh_chg_2020_2024}}}}"] = pct(ring.get("hh_chg_2020_2024"))
            v[f"{{{{r{r}_hh_chg_2024_2029}}}}"] = pct(ring.get("hh_chg_2024_2029"))

            v[f"{{{{r{r}_avg_hh_income}}}}"] = money(ring.get("avg_hh_income"))
            v[f"{{{{r{r}_avg_hh_income_2029}}}}"] = money(ring.get("avg_hh_income_2029"))
            v[f"{{{{r{r}_avg_inc_chg}}}}"] = pct(ring.get("avg_inc_chg"))
            v[f"{{{{r{r}_median_hh_income}}}}"] = money(ring.get("median_hh_income"))
            v[f"{{{{r{r}_median_hh_income_2029}}}}"] = money(ring.get("median_hh_income_2029"))
            v[f"{{{{r{r}_med_inc_chg}}}}"] = pct(ring.get("med_inc_chg"))
            v[f"{{{{r{r}_per_capita_income}}}}"] = money(ring.get("per_capita_income"))
            v[f"{{{{r{r}_per_capita_income_2029}}}}"] = money(ring.get("per_capita_income_2029"))
            v[f"{{{{r{r}_pci_chg}}}}"] = pct(ring.get("pci_chg"))

            v[f"{{{{r{r}_owner_pct}}}}"] = pct(ring.get("owner_pct"))
            v[f"{{{{r{r}_renter_pct}}}}"] = pct(ring.get("renter_pct"))
            v[f"{{{{r{r}_avg_home_value}}}}"] = money(ring.get("avg_home_value"))
            v[f"{{{{r{r}_median_home_value}}}}"] = money(ring.get("median_home_value"))

            for key in (
                "inc_lt_15k", "inc_15_25", "inc_25_35", "inc_35_50", "inc_50_75",
                "inc_75_100", "inc_100_150", "inc_150_200", "inc_200_plus",
                "built_2020_later", "built_2010_2019", "built_2000_2009", "built_1990_1999",
                "built_1980_1989", "built_1970_1979", "built_1960_1969", "built_1950_1959",
                "built_1940_1949", "built_1939_earlier",
                "units_1_det", "units_1_att", "units_2", "units_3_4", "units_5_9",
                "units_10_19", "units_20_49", "units_50_plus", "units_mobile", "units_other",
            ):
                v[f"{{{{r{r}_{key}}}}}"] = num(ring.get(key))

        emp = dataset.get("employment", {})
        for g in geos:
            v[f"{{{{emp_total_{g}}}}}"] = num(emp.get("total_employment", {}).get(g))
            v[f"{{{{unemp_{g}}}}}"] = pct(emp.get("unemployment_rate", {}).get(g))

        val = dataset.get("valuation", {})
        v["{{market_price_psf}}"] = money(val.get("price_psf"), 2)
        v["{{market_building_sf}}"] = num(val.get("building_sf"))
        v["{{market_value}}"] = money(val.get("market_value"))
        v["{{market_value_rounded}}"] = money(val.get("market_value_rounded"))
        v["{{value_aggressive}}"] = money(val.get("value_aggressive"))
        v["{{value_conservative}}"] = money(val.get("value_conservative"))
        gba = self._to_number(val.get("building_sf"))

        def psf_label(total_val):
            try:
                if gba and total_val is not None and float(gba) > 0:
                    return f"${float(total_val) / float(gba):,.2f}/PSF"
            except (TypeError, ValueError):
                pass
            return "—"

        v["{{value_aggressive_psf}}"] = psf_label(val.get("value_aggressive"))
        v["{{value_conservative_psf}}"] = psf_label(val.get("value_conservative"))

        v["{{demographics_source}}"] = dataset.get("demographics_source", "US Census, Esri, BLS")
        v["{{address}}"] = property_data.address
        v["{{executive_summary}}"] = property_data.executive_summary
        v["{{regional_analysis}}"] = property_data.regional_analysis
        v["{{sales_conclusion}}"] = property_data.sales_conclusion
        v["{{reconciliation_summary}}"] = property_data.reconciliation_summary
        v["{{reconciliation_notes}}"] = (
            property_data.reconciliation_notes
            or self._build_reconciliation_notes(property_data)
        )
        return v

    @staticmethod
    def _insert_paragraph_after(paragraph):
        """Insert a new empty paragraph immediately after `paragraph`."""
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        paragraph._element.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def _insert_comparables(self, doc: Document, comps: List[Any]) -> None:
        """Insert CoStar comparable pages after PROPERTY COMPARABLES.

        Removes any template sample comps image, then inserts unique PDF page
        renders (deduped) sized for readable CoStar layout.
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if "PROPERTY COMPARABLES" in p.text.upper():
                heading_idx = i
                break
        if heading_idx is None:
            return

        # Drop template sample screenshot(s) under PROPERTY COMPARABLES through
        # the next section heading so uploaded CoStar pages replace them.
        stop_headings = {
            "RECONCILIATION TABLE",
            "SALES CONCLUSION",
            "CERTIFICATION AND DISCLAIMERS",
            "OPINION OF VALUE",
        }
        # Snapshot paragraph elements in the comps zone (heading exclusive → next section)
        zone_paras = []
        for j in range(heading_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            upper = p.text.strip().upper()
            if upper in stop_headings or upper.startswith("RECONCILIATION"):
                break
            zone_paras.append(p)

        removed_drawings = 0
        for p in zone_paras:
            for drawing in list(p._p.iter(qn("w:drawing"))):
                parent = drawing.getparent()
                if parent is not None:
                    parent.remove(drawing)
                    removed_drawings += 1
            for pict in list(p._p.iter(qn("w:pict"))):
                parent = pict.getparent()
                if parent is not None:
                    parent.remove(pict)
                    removed_drawings += 1

        # Delete emptied sample paragraphs AND orphan page-breaks in the comps
        # zone. Leaving those behind creates large blank pages (template samples
        # were ~6–8" tall with page breaks between them).
        # NEVER delete a paragraph that carries w:sectPr — on Prospect that
        # paragraph links header/footer (blue bars + page numbers).
        removed_paras = 0
        for p in zone_paras:
            text = (p.text or "").strip()
            has_drawing = any(True for _ in p._p.iter(qn("w:drawing")))
            has_pict = any(True for _ in p._p.iter(qn("w:pict")))
            pPr = p._element.find(qn("w:pPr"))
            has_sectpr = (
                pPr is not None and pPr.find(qn("w:sectPr")) is not None
            )
            if text or has_drawing or has_pict or has_sectpr:
                continue
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                removed_paras += 1
        if removed_drawings or removed_paras:
            logger.info(
                "Cleared comps template samples: drawings=%s empty_paras=%s",
                removed_drawings,
                removed_paras,
            )

        # Re-find heading after paragraph deletions (indexes shifted)
        heading_idx = None
        for i, p in enumerate(doc.paragraphs):
            if "PROPERTY COMPARABLES" in p.text.upper():
                heading_idx = i
                break
        if heading_idx is None:
            return
        anchor = doc.paragraphs[heading_idx]
        is_prospect = self._is_prospect_run(doc)
        sorted_comps = sorted(comps, key=lambda c: getattr(c, "comp_number", 0))

        # One image per unique CoStar page (avoids duplicating 2-card pages)
        seen_paths = set()
        page_images = []
        for comp in sorted_comps:
            image_path = (
                getattr(comp, "page_image_path", None)
                or getattr(comp, "image_path", None)
            )
            if not image_path or not os.path.exists(image_path):
                continue
            if image_path in seen_paths:
                continue
            seen_paths.add(image_path)
            page_images.append(image_path)

        # Prospect is a 5–7 page short form — hard-cap CoStar pages so an
        # 11-page comps PDF cannot balloon the report to ~18 pages.
        prospect_comp_cap = 2
        if is_prospect and len(page_images) > prospect_comp_cap:
            logger.info(
                "Prospect short-form: limiting comps pages %s -> %s (template=%s)",
                len(page_images),
                prospect_comp_cap,
                self.template_path,
            )
            page_images = page_images[:prospect_comp_cap]

        if page_images:
            for image_path in page_images:
                img_p = self._insert_paragraph_after(anchor)
                img_p.paragraph_format.space_before = Pt(0)
                img_p.paragraph_format.space_after = Pt(2)
                img_p.paragraph_format.line_spacing = 1.0
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    compressed = self._compress_image_file(
                        image_path,
                        max_edge=900 if is_prospect else 1400,
                        quality=52 if is_prospect else 72,
                    )
                    width_in, _height_in = self._comp_image_display_size(
                        compressed, pack_tight=is_prospect
                    )
                    # Width drives size (aspect kept). Sizing packs ~2 short
                    # CoStar pages per Word page; tall 2-card sheets stay 1/page.
                    img_p.add_run().add_picture(compressed, width=Inches(width_in))
                    anchor = img_p
                except Exception as exc:
                    logger.warning("Could not insert comp page image: %s", exc)
            # Ensure reconciliation/certification still start on a fresh page
            self._ensure_page_break_before_heading(doc, "RECONCILIATION TABLE")
            logger.info(
                "Inserted %d unique CoStar page image(s) for comps", len(page_images)
            )
            return

        # Fallback text details when no page images
        fallback_limit = 4 if is_prospect else 6
        for comp in sorted_comps[:fallback_limit]:
            title_p = self._insert_paragraph_after(anchor)
            name = getattr(comp, "property_name", "") or "Property"
            run = title_p.add_run(name)
            run.bold = True
            anchor = title_p

            details = []
            for label, val in (
                ("Address", getattr(comp, "address", "")),
                ("Primary Use", getattr(comp, "primary_use", "")),
                (
                    "Market / Submarket",
                    f"{getattr(comp, 'market', '')} / {getattr(comp, 'sub_market', '')}".strip(
                        " /"
                    ),
                ),
                ("Comp SF", getattr(comp, "comp_sf", "")),
                ("Acres", getattr(comp, "acres", "")),
                ("Sale Price", getattr(comp, "sale_price", "")),
                ("Sale Price/SF", getattr(comp, "sale_price_sf", "")),
                ("Zoning", getattr(comp, "zoning", "")),
                ("Sale Date", getattr(comp, "off_market_date", "")),
                ("Seller", getattr(comp, "seller_landlord", "")),
                ("Buyer", getattr(comp, "buyer_tenant", "")),
            ):
                if val and str(val).strip():
                    details.append(f"{label}: {val}")

            for line in details:
                detail_p = self._insert_paragraph_after(anchor)
                detail_p.add_run(line)
                anchor = detail_p

        logger.info("Inserted %d comparable properties into BOV report", len(sorted_comps))

    @staticmethod
    def _comp_image_display_size(
        image_path: str, pack_tight: bool = False
    ) -> Tuple[float, float]:
        """Size CoStar page renders to pack ~2 images per Word page when possible.

        - Tall/portrait pages (already 2 cards on one CoStar sheet): nearly full
          page, 1/Word page.
        - Short/single-card pages: height-capped so two stack on one Word page
          instead of leaving a large blank under each image.
        - pack_tight (Prospect): slightly smaller so 2 CoStar pages fit in the
          5–7 page short-form budget.
        """
        max_width = 6.5
        # Usable body height is ~9"; leave a little slack for heading/spacing
        pack_height = 3.6 if pack_tight else 4.15  # two of these fit on one page
        full_page_height = 7.8 if pack_tight else 8.6
        min_pack_width = 5.0
        try:
            from PIL import Image

            with Image.open(image_path) as im:
                w, h = im.size
            if not w or not h:
                return max_width, pack_height
            aspect = h / float(w)

            width = max_width
            height = width * aspect

            # Portrait CoStar sheets already show 2 comps — keep one per Word page
            if aspect >= 1.12:
                if height > full_page_height:
                    height = full_page_height
                    width = height / aspect
                    if width > max_width:
                        width = max_width
                        height = width * aspect
                return round(width, 3), round(height, 3)

            if height <= pack_height + 0.15:
                # Already short enough to stack two per page
                return round(width, 3), round(height, 3)

            # Landscape / trimmed single-card — shrink to pack two per page
            height = pack_height
            width = height / aspect
            if width < min_pack_width:
                width = min_pack_width
                height = width * aspect
            if width > max_width:
                width = max_width
                height = width * aspect
            return round(width, 3), round(height, 3)
        except Exception:
            return max_width, pack_height

    def _fill_employment_table(self, doc: Document, dataset: Dict, county: str, state: str) -> None:
        """Fill the employment table with recent-year history while keeping compact cell styles.

        The template ships with 2010-2019 sample rows. We remap those 10 data rows to the
        latest history years and update values by editing existing runs (never cell.text=,
        which drops TableText/8pt styling and makes years wrap in the narrow Year column).
        """
        from docx.oxml.ns import qn

        if len(doc.tables) < 10:
            return
        history = dataset.get("employment_history") or []
        if not history:
            return

        table = doc.tables[9]
        rows_data = sorted(
            [r for r in history if r.get("year")],
            key=lambda r: int(r["year"]),
        )
        if not rows_data:
            return
        rows_data = rows_data[-10:]

        def fmt_num(n):
            try:
                return f"{int(round(float(n))):,}"
            except (TypeError, ValueError):
                return "—"

        def fmt_pct(n):
            if n is None or n == "":
                return None
            try:
                return f"{float(n):.1f}%"
            except (TypeError, ValueError):
                return "—"

        def set_tc_text(tc, value: str) -> None:
            """Write text into a w:tc while preserving compact TableText/8pt formatting.

            Never use cell.text= — that drops pStyle/sz and makes years wrap in the
            narrow Year column (715 dxa), producing the tall rows seen in reports.
            """
            from docx.oxml import OxmlElement

            text = "" if value is None else str(value)
            paragraphs = tc.findall(qn("w:p"))
            if not paragraphs:
                return
            p0 = paragraphs[0]

            # Keep / restore TableText so row height stays compact
            pPr = p0.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p0.insert(0, pPr)
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is None:
                pStyle = OxmlElement("w:pStyle")
                pPr.insert(0, pStyle)
            pStyle.set(qn("w:val"), "TableText")

            runs = p0.findall(qn("w:r"))
            if not runs:
                run = OxmlElement("w:r")
                p0.append(run)
                runs = [run]

            r0 = runs[0]
            rPr = r0.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r0.insert(0, rPr)
            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), "16")  # 8pt — matches template data rows
            szCs = rPr.find(qn("w:szCs"))
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rPr.append(szCs)
            szCs.set(qn("w:val"), "16")

            t_nodes = list(r0.iter(qn("w:t")))
            if t_nodes:
                t_nodes[0].text = text
                if text.startswith(" ") or text.endswith(" "):
                    t_nodes[0].set(qn("xml:space"), "preserve")
                for t in t_nodes[1:]:
                    t.text = ""
            else:
                t = OxmlElement("w:t")
                t.text = text
                r0.append(t)

            for run in runs[1:]:
                for t in run.iter(qn("w:t")):
                    t.text = ""
            for extra in paragraphs[1:]:
                for t in extra.iter(qn("w:t")):
                    t.text = ""

        data_row_indices = []
        for ri, row in enumerate(table.rows):
            tcs = row._tr.findall(qn("w:tc"))
            if not tcs:
                continue
            year_txt = "".join(t.text or "" for t in tcs[0].iter(qn("w:t"))).strip()
            if re.fullmatch(r"20\d{2}", year_txt):
                data_row_indices.append(ri)

        for idx, ri in enumerate(data_row_indices):
            if idx >= len(rows_data):
                break
            rec = rows_data[idx]
            tcs = table.rows[ri]._tr.findall(qn("w:tc"))
            state_yoy = fmt_pct(rec.get("state_emp_yoy"))
            county_yoy = fmt_pct(rec.get("county_emp_yoy"))
            if idx == 0:
                if state_yoy is None:
                    state_yoy = "no data"
                if county_yoy is None:
                    county_yoy = "no data"
            values = [
                str(int(rec["year"])),
                fmt_num(rec.get("state_emp")),
                state_yoy if state_yoy is not None else "—",
                fmt_num(rec.get("county_emp")),
                county_yoy if county_yoy is not None else "—",
                fmt_pct(rec.get("us_unemp")) or "—",
                fmt_pct(rec.get("state_unemp")) or "—",
                fmt_pct(rec.get("county_unemp")) or "—",
            ]
            for ci, value in enumerate(values):
                if ci < len(tcs):
                    set_tc_text(tcs[ci], value)

        # If history has fewer than 10 years, remove leftover sample year rows
        # so we don't leave mixed eras (e.g. 2020–2025 then orphaned 2016–2019).
        if len(rows_data) < len(data_row_indices):
            for ri in reversed(data_row_indices[len(rows_data) :]):
                try:
                    table._tbl.remove(table.rows[ri]._tr)
                except Exception:
                    pass

        # Header title — keep existing run formatting
        try:
            start_y = int(rows_data[0]["year"])
            end_y = int(rows_data[-1]["year"])
            new_title = f"EMPLOYMENT & UNEMPLOYMENT STATISTICS {start_y} – {end_y}"
            hdr_tc = table.rows[0]._tr.findall(qn("w:tc"))[0]
            set_tc_text(hdr_tc, new_title)
        except Exception:
            pass

        # Geo labels
        try:
            for row in table.rows[1:4]:
                for tc in row._tr.findall(qn("w:tc")):
                    joined = "".join(t.text or "" for t in tc.iter(qn("w:t")))
                    if "{{state}}" in joined or "{{county}}" in joined:
                        replaced = joined.replace("{{state}}", state or "State").replace(
                            "{{county}}", county or "County"
                        )
                        set_tc_text(tc, replaced)
        except Exception:
            pass

        logger.info("Filled employment table with %d recent-year records", len(rows_data))

    def _sweep_remaining_placeholders(self, doc: Document, replacements: Dict[str, str]) -> None:
        """Replace any leftover {{placeholders}} still present in the document XML.

        IMPORTANT: skip paragraphs inside w:txbxContent. Floating cover textboxes
        are nested under body paragraphs as drawings; joining all descendant w:t
        nodes would mash Date / property name / PREPARED BY into one box.
        Textboxes are handled by `_replace_in_textboxes`.
        """
        import re
        from docx.oxml.ns import qn

        pattern = re.compile(r"\{\{[^}]+\}\}")
        w_t = qn("w:t")
        w_p = qn("w:p")
        w_txbx = qn("w:txbxContent")
        replaced = 0

        def _inside_textbox(el) -> bool:
            parent = el.getparent()
            while parent is not None:
                if parent.tag == w_txbx:
                    return True
                parent = parent.getparent()
            return False

        def _direct_text_nodes(paragraph):
            """w:t nodes in this paragraph that are NOT inside nested textboxes."""
            nodes = []
            for n in paragraph.iter(w_t):
                if n.text and not _inside_textbox(n):
                    # Still inside a drawing/textbox if ancestor txbx wraps us —
                    # _inside_textbox already covers that. Also skip if any
                    # ancestor between n and paragraph is a drawing with txbx.
                    nodes.append(n)
            # If paragraph itself is inside a textbox, skip entirely
            if _inside_textbox(paragraph):
                return []
            # Filter: only nodes whose containing paragraph is THIS paragraph
            # (not nested paragraphs inside drawings)
            out = []
            for n in nodes:
                parent = n.getparent()
                while parent is not None and parent.tag != w_p:
                    parent = parent.getparent()
                if parent is paragraph:
                    out.append(n)
            return out

        def process(root):
            nonlocal replaced
            if root is None:
                return
            for p in root.iter(w_p):
                if _inside_textbox(p):
                    continue
                nodes = _direct_text_nodes(p)
                if not nodes:
                    continue
                joined = "".join(n.text or "" for n in nodes)
                if "{{" not in joined:
                    continue
                new_text = joined
                for placeholder, value in replacements.items():
                    if placeholder in new_text:
                        safe = value if value is not None else "—"
                        safe = re.sub(r"\s*\n\s*", " ", str(safe)).strip()
                        new_text = new_text.replace(placeholder, safe)
                new_text, _n = pattern.subn("—", new_text)
                if new_text != joined:
                    nodes[0].text = new_text
                    for n_el in nodes[1:]:
                        n_el.text = ""
                    replaced += 1

        process(doc.element.body)
        for section in doc.sections:
            for part in (section.header, section.footer):
                try:
                    process(part._element)
                except Exception:
                    pass
        if replaced:
            logger.info("Swept leftover placeholders in %d paragraph(s)", replaced)

    def _replace_in_textboxes(self, doc: Document, replacements: Dict[str, str]):
        """Replace placeholders inside text boxes (cover, side bars) and headers/footers.

        Placeholders may be split across multiple w:t runs (common in the short-form
        template). For each textbox paragraph we join runs, replace, then write back
        into the first run — scoped per paragraph so cover boxes stay separate.
        """
        import re
        from docx.oxml.ns import qn

        w_t = qn("w:t")
        w_p = qn("w:p")
        pattern = re.compile(r"\{\{[^}]+\}\}")

        def process_textbox(txbx):
            for p in txbx.findall(w_p):
                # Only this paragraph's own runs (not nested)
                nodes = []
                for r in p.findall(qn("w:r")):
                    for t in r.findall(w_t):
                        if t.text is not None:
                            nodes.append(t)
                if not nodes:
                    # fallback: any direct descendant w:t under this p's runs only
                    nodes = [n for n in p.iter(w_t) if n.text is not None]
                if not nodes:
                    continue
                joined = "".join(n.text or "" for n in nodes)
                if "{{" not in joined and "PREPARED" not in joined:
                    # Still normalize labels / allow non-placeholder textboxes
                    if not any(ph in joined for ph in replacements):
                        continue
                new_text = joined
                for placeholder, value in replacements.items():
                    if placeholder in new_text:
                        safe = "" if value is None else str(value)
                        new_text = new_text.replace(placeholder, safe)
                # Leave unmatched placeholders for a second pass only within this box
                if "{{" in new_text:
                    new_text = pattern.sub("—", new_text)
                if new_text != joined:
                    nodes[0].text = new_text
                    if new_text.startswith(" ") or new_text.endswith(" "):
                        nodes[0].set(qn("xml:space"), "preserve")
                    for n in nodes[1:]:
                        n.text = ""
                # Preserve trailing space after PREPARED labels
                if nodes and nodes[0].text:
                    stripped = nodes[0].text.strip()
                    if stripped in ("PREPARED BY:", "PREPARED FOR:"):
                        nodes[0].text = stripped + " "
                        nodes[0].set(qn("xml:space"), "preserve")

        def process(root):
            if root is None:
                return
            for txbx in root.iter(qn("w:txbxContent")):
                process_textbox(txbx)

        process(doc.element.body)
        for section in doc.sections:
            for part in (section.header, section.footer):
                try:
                    process(part._element)
                except Exception:
                    pass

    def _replace_textbox_images(self, doc: Document, image_map: Dict[str, tuple]):
        """Swap the embedded cover/branding pictures whose alt-text marks them.

        In the source template these spots are real pictures whose placeholder
        (e.g. {{main_img}}) lives in the picture's alt-text description
        (wp:docPr/pic:cNvPr @descr), not as body text. We locate each drawing by
        that description, replace the underlying image bytes with our generated
        image (preserving the template's size/layout), and clear the placeholder
        from the alt-text so no {{...}} remains anywhere.

        image_map: { '{{placeholder}}': (image_path_or_None, width_inches) }
        width is ignored here; the template's existing extent is kept.
        """
        from docx.oxml.ns import qn

        descr_tags = (qn("wp:docPr"), qn("pic:cNvPr"))
        blip_tag = qn("a:blip")
        embed_attr = qn("r:embed")

        # Longest-first so {{aerial_image}} never loses to a shorter token
        placeholders = sorted(image_map.keys(), key=len, reverse=True)

        for drawing in doc.element.body.iter(qn("w:drawing")):
            descr_nodes = [el for tag in descr_tags for el in drawing.iter(tag)]
            matched = None
            for el in descr_nodes:
                descr = el.get("descr") or ""
                for placeholder in placeholders:
                    if placeholder in descr:
                        matched = placeholder
                        break
                if matched:
                    break
            if not matched:
                continue

            image_path, _width = image_map[matched]
            rid = None
            for blip in drawing.iter(blip_tag):
                rid = blip.get(embed_attr)
                if rid:
                    break
            part = doc.part.related_parts.get(rid) if rid else None

            if image_path and os.path.exists(image_path):
                if part is not None:
                    try:
                        part._blob = self._image_bytes_for_part(image_path, part)
                        logger.info(
                            "Swapped template image %s -> %s (%s bytes)",
                            matched,
                            image_path,
                            len(part._blob),
                        )
                    except Exception as exc:
                        logger.error(f"Image swap failed for {matched}: {exc}")
                else:
                    logger.warning(f"No image relationship found for {matched}")
            else:
                # Cover must never keep the template's sample placeholder icon.
                if matched == "{{main_img}}" and part is not None:
                    part._blob = self._blank_cover_bytes(part)
                    logger.warning(
                        "Street View missing — placed neutral cover placeholder "
                        "for {{main_img}} (will not keep sample aerial)"
                    )
                else:
                    logger.warning(f"Image not available for {matched}: {image_path}")

            # Clear the placeholder from alt-text so no {{...}} survives
            for el in descr_nodes:
                descr = el.get("descr") or ""
                for placeholder in placeholders:
                    descr = descr.replace(placeholder, "")
                el.set("descr", descr)

            # Remove the light-grey textbox/picture frame around the image
            self._strip_drawing_border(drawing)
            # Stretch nested picture to the textbox extent so no white matte shows
            self._fit_picture_to_textbox(drawing)

    @staticmethod
    def _fit_picture_to_textbox(drawing) -> None:
        """Make nested pic:spPr / wp:extent match the outer textbox size."""
        from docx.oxml.ns import qn

        ext = next(drawing.iter(qn("wp:extent")), None)
        if ext is None:
            return
        cx, cy = ext.get("cx"), ext.get("cy")
        if not cx or not cy:
            return
        pic_sppr = "{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr"
        a_xfrm = "{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm"
        a_ext = "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"
        for spPr in drawing.iter(pic_sppr):
            xfrm = spPr.find(a_xfrm)
            if xfrm is None:
                continue
            aext = xfrm.find(a_ext)
            if aext is not None:
                aext.set("cx", str(cx))
                aext.set("cy", str(cy))
        # Also sync any nested wp:extent under pic
        for nested in drawing.iter(qn("wp:extent")):
            nested.set("cx", str(cx))
            nested.set("cy", str(cy))

    @classmethod
    def _strip_all_image_borders(cls, doc: Document) -> None:
        from docx.oxml.ns import qn

        markers = ("main_img", "aerial_image", "subject_photo", "Subject_photo")
        for drawing in doc.element.body.iter(qn("w:drawing")):
            descr = " ".join(
                (el.get("descr") or "")
                for tag in (qn("wp:docPr"), qn("pic:cNvPr"))
                for el in drawing.iter(tag)
            )
            has_blip = next(drawing.iter(qn("a:blip")), None) is not None
            if any(m in descr for m in markers):
                cls._strip_drawing_border(drawing)
                continue
            if not has_blip:
                continue
            docPr = next(drawing.iter(qn("wp:docPr")), None)
            name = (docPr.get("name") if docPr is not None else "") or ""
            # Cover/aerial/subject frames (descr cleared after swap)
            if name.startswith("Text Box") or name.startswith("Picture"):
                cls._strip_drawing_border(drawing)

    @staticmethod
    def _strip_drawing_border(drawing) -> bool:
        """Remove light-grey stroke / matte frame around cover / aerial / subject images."""
        from lxml import etree
        from docx.oxml.ns import qn

        wps_sppr = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}spPr"
        pic_sppr = "{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr"
        a_ln = "{http://schemas.openxmlformats.org/drawingml/2006/main}ln"
        a_nofill = "{http://schemas.openxmlformats.org/drawingml/2006/main}noFill"
        a_solid = "{http://schemas.openxmlformats.org/drawingml/2006/main}solidFill"
        a_srgb = "{http://schemas.openxmlformats.org/drawingml/2006/main}srgbClr"
        a_alpha = "{http://schemas.openxmlformats.org/drawingml/2006/main}alpha"
        wps_style = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}style"
        a_ln_ref = "{http://schemas.openxmlformats.org/drawingml/2006/main}lnRef"
        a_effect = "{http://schemas.openxmlformats.org/drawingml/2006/main}effectLst"

        changed = False
        for spPr in list(drawing.iter(wps_sppr)) + list(drawing.iter(pic_sppr)):
            ln = spPr.find(a_ln)
            if ln is None:
                ln = etree.SubElement(spPr, a_ln)
            ln.clear()
            ln.set("w", "0")
            etree.SubElement(ln, a_nofill)
            changed = True
            solid = spPr.find(a_solid)
            if solid is not None:
                spPr.remove(solid)
                if spPr.find(a_nofill) is None:
                    etree.SubElement(spPr, a_nofill)
                changed = True
            for eff in list(spPr.findall(a_effect)):
                spPr.remove(eff)
                changed = True
        for style in drawing.iter(wps_style):
            ln_ref = style.find(a_ln_ref)
            if ln_ref is not None:
                for child in list(ln_ref):
                    ln_ref.remove(child)
                srgb = etree.SubElement(ln_ref, a_srgb)
                srgb.set("val", "FFFFFF")
                alpha = etree.SubElement(srgb, a_alpha)
                alpha.set("val", "0")
                changed = True
        for ee in drawing.iter(qn("wp:effectExtent")):
            for side in ("l", "t", "r", "b"):
                if ee.get(side) not in (None, "0"):
                    ee.set(side, "0")
                    changed = True
        return changed

    @staticmethod
    def _blank_cover_bytes(part=None) -> bytes:
        """Valid cover-sized neutral image so Word never shows a broken-image icon."""
        from io import BytesIO

        partname = str(getattr(part, "partname", "") or "").lower() if part else ""
        content_type = str(getattr(part, "content_type", "") or "").lower() if part else ""
        wants_png = partname.endswith(".png") or "image/png" in content_type

        try:
            from PIL import Image

            # Soft grey fill (not white) so the empty cover slot is visible but clean
            img = Image.new("RGB", (640, 640), (230, 230, 230))
            buf = BytesIO()
            if wants_png:
                img.convert("RGBA").save(buf, format="PNG", optimize=True)
            else:
                img.save(buf, format="JPEG", quality=70, optimize=True)
            return buf.getvalue()
        except Exception:
            return ComprehensivePropertyReportGenerator._blank_jpeg_bytes(640, 640)

    @staticmethod
    def _blank_jpeg_bytes(width: int = 16, height: int = 16) -> bytes:
        """Tiny white JPEG used as last-resort placeholder bytes."""
        from io import BytesIO

        try:
            from PIL import Image

            buf = BytesIO()
            Image.new("RGB", (width, height), (230, 230, 230)).save(
                buf, format="JPEG", quality=60
            )
            return buf.getvalue()
        except Exception:
            # Minimal valid JPEG
            return (
                b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
                b"\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t"
                b"\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a"
                b"\x1f\x1e\x1d\x1a\x1c\x1c $.\' \",#\x1c\x1c(7),01444\x1f\'9=82<.342"
                b"\xff\xc0\x00\x0b\x08\x00\x01\x00\x01\x01\x01\x11\x00\xff\xc4\x00"
                b"\x14\x00\x01\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00"
                b"\x00\x00\x08\xff\xc4\x00\x14\x10\x01\x00\x00\x00\x00\x00\x00\x00"
                b"\x00\x00\x00\x00\x00\x00\x00\x00\x00\xff\xda\x00\x08\x01\x01\x00"
                b"\x00?\x00\x7f\xbf\xff\xd9"
            )

    @staticmethod
    def _compress_image_file(image_path: str, max_edge: int = 1400, quality: int = 72) -> str:
        """Write a resized/JPEG-compressed copy; return path (original if compress fails)."""
        from io import BytesIO
        from tempfile import NamedTemporaryFile

        try:
            from PIL import Image

            with Image.open(image_path) as im:
                img = im.convert("RGB")
                w, h = img.size
                scale = min(1.0, float(max_edge) / float(max(w, h)))
                if scale < 1.0:
                    img = img.resize(
                        (max(1, int(w * scale)), max(1, int(h * scale))),
                        Image.LANCZOS,
                    )
                tmp = NamedTemporaryFile(delete=False, suffix=".jpg")
                tmp.close()
                img.save(tmp.name, format="JPEG", quality=quality, optimize=True)
                return tmp.name
        except Exception as exc:
            logger.warning("Image compress skipped for %s: %s", image_path, exc)
            return image_path

    @staticmethod
    def _image_bytes_for_part(image_path: str, part) -> bytes:
        """Return compressed image bytes matching the template part's format."""
        from io import BytesIO

        raw = Path(image_path).read_bytes()
        partname = str(getattr(part, "partname", "") or "").lower()
        content_type = str(getattr(part, "content_type", "") or "").lower()
        wants_jpeg = partname.endswith((".jpg", ".jpeg")) or "image/jpeg" in content_type
        wants_png = partname.endswith(".png") or "image/png" in content_type

        try:
            from PIL import Image

            img = Image.open(BytesIO(raw)).convert("RGB")
            w, h = img.size
            max_edge = 1400
            scale = min(1.0, float(max_edge) / float(max(w, h)))
            if scale < 1.0:
                img = img.resize(
                    (max(1, int(w * scale)), max(1, int(h * scale))),
                    Image.LANCZOS,
                )
            buf = BytesIO()
            if wants_png and not wants_jpeg:
                # Keep PNG part valid but still shrink pixels
                img_rgba = img.convert("RGBA")
                img_rgba.save(buf, format="PNG", optimize=True)
            else:
                img.save(buf, format="JPEG", quality=72, optimize=True)
            return buf.getvalue()
        except Exception as exc:
            logger.warning("Could not compress %s for template part: %s", image_path, exc)
            return raw

    def _replace_image_placeholder(self, doc: Document, placeholder: str, image_path: Optional[str], width_inches: float = 3.0):
        """
        Replace image placeholder with actual image in the document
        
        Args:
            doc: The Word document
            placeholder: The placeholder text to replace (e.g., '{{aerial_image}}')
            image_path: Path to the image file
            width_inches: Width of the image in inches
        """
        if not image_path or not os.path.exists(image_path):
            logger.warning(f"Image not found for placeholder {placeholder}: {image_path}")
            # Clear the raw placeholder text so it doesn't show in the output
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, "")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, "")
            return
            
        # Search and replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Clear the paragraph
                paragraph.text = paragraph.text.replace(placeholder, '')
                # Add the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                logger.info(f"Replaced {placeholder} with image from {image_path}")
                
        # Search and replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            # Clear the paragraph
                            paragraph.text = paragraph.text.replace(placeholder, '')
                            # Add the image
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(width_inches))
                            logger.info(f"Replaced {placeholder} with image in table from {image_path}")

    def _create_market_analysis_section(self, doc: Document, property_data: PropertyReportData):
        """Add market analysis section to the document"""
        
        # Add page break before market analysis
        doc.add_page_break()
        
        # Add title with formatting
        title = doc.add_heading(f'{property_data.county} {property_data.property_type} Market Report -- Q{(datetime.now().month-1)//3 + 1} {datetime.now().year}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add prepared by section
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run('Prepared by ')
        run1.bold = True
        run2 = p.add_run(f'[{property_data.prepared_by_company or "Your Business Name"}]\n')
        run3 = p.add_run('Independent Analysis Based on Multiple Public Data Sources')
        run3.italic = True
        
        # Section 1: Market Overview
        doc.add_heading('1. Market Overview', level=2)
        doc.add_paragraph(property_data.market_overview)
        
        # Section 2: Key Market Metrics
        doc.add_heading('2. Key Market Metrics', level=2)
        
        # Add vacancy rates
        doc.add_paragraph(property_data.vacancy_rates)
        
        # Add lease rates
        doc.add_paragraph(property_data.lease_rates)
        
        # Add construction activity
        doc.add_paragraph(property_data.construction_activity)
        
        # Section 3: Trends & Forecast
        doc.add_heading('3. Trends & Forecast', level=2)
        doc.add_paragraph(property_data.market_trends)
        
        # Section 4: Investment Insights
        doc.add_heading('4. Investment Insights', level=2)
        doc.add_paragraph(property_data.investment_insights)
        
        # Section 5: Recommendations
        doc.add_heading('5. Recommendations', level=2)
        doc.add_paragraph(property_data.market_recommendations)
        
        # Section 6: Data Sources & Disclaimer
        doc.add_heading('6. Data Sources & Disclaimer', level=2)
        doc.add_paragraph(property_data.market_data_sources)

    def create_word_document(self, property_data: PropertyReportData, color_theme: Optional[str] = None) -> str:
        """
        Create Word document from property data using template
        """
        logger.info(f"Creating Word document for: {property_data.address}")
        
        # Load template
        doc = Document(self.template_path)
        is_prospect = self._is_prospect_run(doc)
        if is_prospect:
            logger.info(
                "Using Prospect (5-7 page) short-form template layout (%s)",
                self.template_path,
            )
            self._ensure_prospect_header_footer(doc)

        # Property-detail fields should always show a value so the report stays
        # consistent with the prior template (blank inputs -> "N/A").
        def na(value):
            text = str(value).strip() if value is not None else ""
            return text if text else "N/A"
        
        # Define all replacements
        replacements = {
            '{{Date}}': property_data.date,
            '{{prepared_by}}': property_data.prepared_by,
            '{{prepared_by_company}}': property_data.prepared_by_company,
            '{{prepared_by_address}}': property_data.prepared_by_address,
            '{{prepared_for}}': property_data.prepared_for,
            '{{prepared_for_company}}': property_data.prepared_for_company,
            '{{prepared_for_address}}': property_data.prepared_for_address,
            '{{property_summary}}': property_data.property_summary,
            '{{property_name}}': property_data.property_name,
            '{{property_type}}': (property_data.property_type or "").strip().title(),
            '{{state}}': property_data.state,
            '{{county}}': property_data.county,
            '{{longitude}}': property_data.longitude,
            '{{latitude}}': property_data.latitude,
            '{{Topography}}': na(property_data.topography),
            '{{shape}}': na(property_data.shape),
            '{{Access}}': na(property_data.access),
            '{{Exposure}}': na(property_data.exposure),
            '{{lot_area}}': self._format_display_number(property_data.lot_area) or "N/A",
            '{{acres}}': na(property_data.acres),
            '{{recorded_sale_date}}': na(property_data.recorded_sale_date),
            '{{zoning}}': na(property_data.zoning),
            '{{apn}}': na(property_data.apn),
            '{{current_owner}}': na(property_data.current_owner),
            '{{marketing_period}}': property_data.marketing_period,
            '{{swot_strengths}}': property_data.swot_strengths,
            '{{swot_weaknesses}}': property_data.swot_weaknesses,
            '{{swot_opportunities}}': property_data.swot_opportunities,
            '{{swot_threats}}': property_data.swot_threats,
            '{{location_summary}}': property_data.location_summary,
            '{{demographic_analysis}}': property_data.demographic_analysis,
            '{{size_and_topography}}': property_data.size_and_topography,
            '{{population_analysis}}': property_data.population_analysis,
            '{{household_trends}}': property_data.household_trends,
            '{{housing_tenure}}': property_data.housing_tenure,
            '{{local_area_analysis}}': property_data.local_area_analysis,
            '{{employment_analysis}}': property_data.employment_analysis,
            '{{economic_factors}}': property_data.economic_factors,
            '{{community_services}}': property_data.community_services,
            # Market Analysis placeholders
            '{{market_overview}}': property_data.market_overview,
            '{{vacancy_rates}}': property_data.vacancy_rates,
            '{{lease_rates}}': property_data.lease_rates,
            '{{construction_activity}}': property_data.construction_activity,
            '{{market_trends}}': property_data.market_trends,
            '{{investment_insights}}': property_data.investment_insights,
            '{{market_recommendations}}': property_data.market_recommendations,
            '{{market_data_sources}}': property_data.market_data_sources,
            '{{market_quarter}}': property_data.market_quarter,
        }

        # Cover polish BEFORE placeholder replacement so {{address}} is still findable.
        # Theme tint on hero, left-flush bars, address in white diagonal corner.
        self._polish_cover_layout(doc, color_theme=color_theme)

        # Prefer a two-line cover address like the mock (street / city-state-zip)
        if property_data.table_values:
            replacements.update(property_data.table_values)
            ring_keys = [k for k in property_data.table_values if k.startswith("{{r")]
            logger.info("Merged %d BOV table placeholders (%d ring fields)",
                        len(property_data.table_values), len(ring_keys))
        replacements["{{address}}"] = self._format_cover_address(
            replacements.get("{{address}}") or property_data.address
        )
        
        # Replace text in all document elements
        self._replace_text_in_document(doc, replacements)

        # Replace placeholders inside text boxes (cover, side bars) and headers/footers
        self._replace_in_textboxes(doc, replacements)

        # Final sweep: catch any leftover {{placeholders}} (split runs / missed cells)
        self._sweep_remaining_placeholders(doc, replacements)

        # Justified paragraphs + soft breaks → huge word gaps on short last lines
        self._fix_justified_soft_breaks(doc)

        # CLIENT look: regional body in italic blue (full template only)
        if not is_prospect:
            self._style_regional_analysis(doc)

        # Cover / aerial / subject images live in text boxes (alt-text placeholders).
        # Keep the template's designed diagonal cover hero (Picture 8) — only tint it.
        # Bottom-left {{main_img}} is Street View only — never aerial.
        cover_photo = property_data.street_view_image_path
        aerial_photo = property_data.aerial_image_path
        # Guard: never let a shared/mis-assigned path put the aerial onto the cover
        if (
            cover_photo
            and aerial_photo
            and Path(cover_photo).resolve() == Path(aerial_photo).resolve()
        ):
            logger.error(
                "Street View path identical to aerial path (%s); blanking cover",
                cover_photo,
            )
            cover_photo = None
        if not cover_photo:
            logger.warning(
                "Street View unavailable for cover {{main_img}}; cover slot will be blanked "
                "(will not fall back to aerial)"
            )
        else:
            logger.info("Cover {{main_img}} <- Street View %s", cover_photo)
        self._replace_textbox_images(doc, {
            '{{main_img}}': (cover_photo, 6.0),
            '{{aerial_image}}': (aerial_photo, 6.0),
            '{{Subject_photo}}': (property_data.street_view_image_path, 3.5),
            '{{subject_photo}}': (property_data.street_view_image_path, 3.5),
        })
        # Final pass: strip any remaining grey frames around image slots
        self._strip_all_image_borders(doc)

        # Replace image placeholders in regular paragraphs/cells (legacy + BOV names)
        for ph in ('{{ariel_image}}', '{{aerial_map}}'):
            self._replace_image_placeholder(doc, ph, property_data.aerial_image_path, width_inches=6.0)
        for ph in ('{{street_view}}', '{{subject_photos}}'):
            self._replace_image_placeholder(doc, ph, property_data.street_view_image_path, width_inches=4.0)

        # Insert comparable sales from uploaded PDF (fills the comps section; avoids blank page)
        if property_data.comps:
            self._insert_comparables(doc, property_data.comps)
        else:
            # Even without uploaded comps, keep reconciliation on its own page
            self._ensure_page_break_before_heading(doc, "RECONCILIATION TABLE")

        # Refresh employment table with recent-year data (client template only)
        if property_data.bov_dataset and not is_prospect:
            self._fill_employment_table(doc, property_data.bov_dataset, property_data.county, property_data.state)

        # Collapse leftover empty paragraphs that create large white gaps
        # (never deletes paragraphs that carry a page break).
        self._collapse_empty_spacing(doc)

        # TOC is a full-page floating blue panel — without a page break before
        # EXECUTIVE SUMMARY, body text renders underneath it (page-2 bleed).
        self._ensure_page_break_before_heading(doc, "EXECUTIVE SUMMARY")
        if is_prospect:
            # Align TOC labels with body order so live page numbers increase
            # sensibly (Demographics maps to Property Summary, etc.).
            self._sync_prospect_toc_labels(doc)

        # Ensure TOC leaders are a single middle line (no title/page underlines)
        self._normalize_toc_leaders(doc, is_prospect=is_prospect)

        # Keep a blank line above section titles (e.g. after General Information table)
        for heading in self.SECTION_HEADINGS_NEEDING_SPACE:
            self._ensure_blank_before_heading(doc, heading)
        
        # Remove the programmatic market analysis section since we're using placeholders
        # self._create_market_analysis_section(doc, property_data)

        # Clean, client-friendly filename: BOV_[Prospect|Client]_<name>_<date>_<time>.docx
        label = (property_data.property_name or "").strip() or property_data.address
        safe_label = "".join(c for c in label if c.isalnum() or c in (" ", "-", "_")).strip()
        safe_label = "_".join(safe_label.split())[:60]
        form_tag = "Prospect" if is_prospect else "Client"
        output_filename = (
            f"BOV_{form_tag}_{safe_label}_{datetime.now().strftime('%Y-%m-%d_%H%M')}.docx"
        )
        output_path = self.output_dir / output_filename
        
        # Save document
        doc.save(output_path)
        logger.info(f"Document saved: {output_path}")

        # TOC page numbers + cover shapes (white diagonal corner / Street View)
        # via Word COM. Color theme runs after so COM cannot wipe accent colors.
        self._refresh_toc_page_numbers(
            output_path,
            is_prospect=is_prospect,
            street_view_path=cover_photo,
        )

        # Apply color theme to the report accent (post-process the saved file)
        if color_theme:
            self._apply_color_theme(output_path, color_theme)

        return str(output_path)

    def _fix_cover_after_word_com(
        self, output_path, street_view_path: Optional[str] = None
    ) -> None:
        """Repair cover after Word COM TOC refresh converts DrawingML → VML.

        COM sets the white diagonal corner shapes to filled=\"f\" (invisible).
        Restore solid white corner fills and white title text.

        IMPORTANT: Do not inject VML imagedata/fill into Text Box 21 — that
        produces packages Word refuses to open ("problems with the contents").
        """
        import zipfile
        import shutil
        import tempfile
        import re

        path = Path(output_path)
        if not path.exists():
            return

        try:
            with zipfile.ZipFile(path, "r") as zin:
                files = {n: zin.read(n) for n in zin.namelist()}
        except Exception as exc:
            logger.error("Cover repair: cannot open %s: %s", path, exc)
            return

        xml_name = "word/document.xml"
        rels_name = "word/_rels/document.xml.rels"
        ct_name = "[Content_Types].xml"
        if xml_name not in files:
            return
        xml = files[xml_name].decode("utf-8", "ignore")
        original = xml
        rels = files.get(rels_name, b"").decode("utf-8", "ignore")
        package_changed = False

        # --- 1) White diagonal corner (Parallelogram 3 + Rectangle 12) ---
        def _fill_white(match: re.Match) -> str:
            tag = match.group(0)
            tag = re.sub(r'\sfilled="f"', ' filled="t"', tag)
            if "filled=" not in tag:
                tag = tag.replace(">", ' filled="t">', 1)
            if "fillcolor=" in tag:
                tag = re.sub(r'fillcolor="[^"]*"', 'fillcolor="#ffffff"', tag)
            else:
                tag = tag.replace(" filled=", ' fillcolor="#ffffff" filled=', 1)
                if "fillcolor=" not in tag:
                    tag = tag.replace(">", ' fillcolor="#ffffff">', 1)
            tag = re.sub(r'\sstroked="t"', ' stroked="f"', tag)
            return tag

        for shape_id in ("Parallelogram 3", "Rectangle 12"):
            pattern = rf'(<(?:v:)?(?:shape|rect)\b[^>]*\bid="{re.escape(shape_id)}"[^>]*/?>)'
            xml2, n = re.subn(pattern, _fill_white, xml, count=1)
            if n:
                xml = xml2
                logger.info("Cover repair: forced white fill on %s", shape_id)

        xml = re.sub(
            r'(<(?:v:)?group\b[^>]*\bid="Group 13"[^>]*)\sfilled="f"',
            r'\1 filled="t" fillcolor="#ffffff"',
            xml,
            count=1,
        )

        # --- 2) White title text on "Broker Opinion of Value" ---
        # Replace existing w:color (do NOT add a second w:color — Word rejects that)
        def _whiten_title_rpr(m: re.Match) -> str:
            rpr, rest = m.group(1), m.group(2)
            if re.search(r'<w:color\b', rpr):
                rpr = re.sub(
                    r'<w:color\b[^>]*/>',
                    '<w:color w:val="FFFFFF"/>',
                    rpr,
                    count=1,
                )
            else:
                rpr = rpr + '<w:color w:val="FFFFFF"/>'
            return rpr + rest

        xml = re.sub(
            r"(<w:rPr>(?:(?!</w:rPr>).)*?)(</w:rPr>\s*<w:t[^>]*>Broker Opinion of Value</w:t>)",
            _whiten_title_rpr,
            xml,
            flags=re.DOTALL,
        )

        # --- 3) Strip unsafe VML image injects from older broken builds ---
        xml2, n_strip = re.subn(
            r'(id="Text Box 21"[^>]*>)\s*(?:<v:fill\b[^>]*/>\s*)?(?:<v:imagedata\b[^>]*/>\s*)?',
            r"\1",
            xml,
            count=1,
            flags=re.IGNORECASE,
        )
        if n_strip:
            xml = xml2
            logger.info("Cover repair: removed unsafe VML image inject from Text Box 21")

        for orphan in (
            "word/media/cover_street.jpg",
            "word/media/cover_street.jpeg",
        ):
            if orphan in files:
                files.pop(orphan, None)
                package_changed = True
                logger.info("Cover repair: removed orphan media %s", orphan)

        if "cover_street" in rels:
            rels2 = re.sub(
                r'<Relationship[^>]*Target="media/cover_street\.(?:jpg|jpeg)"[^>]*/>',
                "",
                rels,
            )
            if rels2 != rels:
                files[rels_name] = rels2.encode("utf-8")
                package_changed = True
                logger.info("Cover repair: removed cover_street relationship")

        ct = files.get(ct_name, b"").decode("utf-8", "ignore")
        if ct and "cover_street" in ct:
            ct2 = re.sub(
                r'<Override[^>]*cover_street\.(?:jpg|jpeg)"[^>]*/>', "", ct
            )
            if ct2 != ct:
                files[ct_name] = ct2.encode("utf-8")
                package_changed = True

        if xml == original and not package_changed:
            logger.info("Cover repair: no changes needed")
            return

        try:
            from lxml import etree

            etree.fromstring(xml.encode("utf-8"))
            if rels_name in files:
                etree.fromstring(files[rels_name])
            if ct_name in files:
                etree.fromstring(files[ct_name])
        except Exception as exc:
            logger.error(
                "Cover repair aborted — generated XML would be invalid: %s", exc
            )
            return

        files[xml_name] = xml.encode("utf-8")
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(tmp_fd)
        try:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for name, data in files.items():
                    zout.writestr(name, data)
            shutil.move(tmp_path, path)
            logger.info("Cover repair applied to %s", path)
        except Exception as exc:
            logger.error("Cover repair write failed: %s", exc)
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    def _is_prospect_run(self, doc: Optional[Document] = None) -> bool:
        """True when generating the Prospect short-form (path first, then doc heuristic)."""
        path = str(getattr(self, "template_path", "") or "").lower()
        if "prospect" in path or path.endswith("short") or "short_form" in path:
            return True
        if doc is not None:
            return self._is_prospect_template(doc)
        return False

    @staticmethod
    def _is_prospect_template(doc: Document) -> bool:
        """True for the condensed 5-7 page BOV (Prospect) short-form layout."""
        if len(doc.tables) <= 6:
            for table in doc.tables:
                try:
                    title = (table.rows[0].cells[0].text or "").strip().upper()
                except Exception:
                    continue
                if title.startswith("POPULATION") and "DENSITY" not in title:
                    return False
                if title.startswith("EMPLOYMENT"):
                    return False
            return True
        return False

    @staticmethod
    def _ensure_prospect_header_footer(doc: Document) -> None:
        """Keep blue header/footer bars + Page X of Y on Prospect short-form.

        Short-form originally linked these via a mid-doc sectPr in the comps zone.
        If that paragraph is removed, bars/page numbers disappear. Re-link them
        on the body sectPr (same visual as Client long template).
        """
        from lxml import etree
        from docx.oxml.ns import qn

        header_rid = footer_rid = None
        for rel in doc.part.rels.values():
            target = str(getattr(rel, "target_ref", "") or "").lower()
            if "header" in target and header_rid is None:
                header_rid = rel.rId
            if "footer" in target and footer_rid is None:
                footer_rid = rel.rId
        if not header_rid or not footer_rid:
            logger.warning("Prospect template missing header/footer parts")
            return

        sect = doc.element.body.find(qn("w:sectPr"))
        if sect is None:
            return

        w_header = qn("w:headerReference")
        w_footer = qn("w:footerReference")
        existing_h = sect.findall(w_header)
        existing_f = sect.findall(w_footer)
        if existing_h and existing_f:
            return

        for tag in (w_header, w_footer):
            for el in list(sect.findall(tag)):
                sect.remove(el)
        href = etree.Element(w_header)
        href.set(qn("w:type"), "default")
        href.set(qn("r:id"), header_rid)
        fref = etree.Element(w_footer)
        fref.set(qn("w:type"), "default")
        fref.set(qn("r:id"), footer_rid)
        sect.insert(0, fref)
        sect.insert(0, href)
        logger.info(
            "Restored Prospect header/footer refs (%s / %s)", header_rid, footer_rid
        )

    @staticmethod
    def _paragraph_has_page_break(paragraph) -> bool:
        """True if this paragraph contains an explicit page break."""
        from docx.oxml.ns import qn

        return any(
            br.get(qn("w:type")) == "page"
            for br in paragraph._element.iter(qn("w:br"))
        )

    # Section titles that need a blank line above them (matches source BOV layout)
    SECTION_HEADINGS_NEEDING_SPACE = (
        "LOCATION SUMMARY",
        "AERIAL MAP",
        "SUBJECT PHOTOS",
        "PROPERTY SUMMARY",
        "REGIONAL ANALYSIS",
        "DEMOGRAPHIC ANALYSIS",
        "PROPERTY COMPARABLES",
        "RECONCILIATION TABLE",
        "SALES CONCLUSION",
        "CERTIFICATION AND DISCLAIMERS",
    )

    def _collapse_empty_spacing(self, doc: Document) -> None:
        """Remove consecutive empty paragraphs that create large white gaps.

        Never deletes a paragraph that carries a page break — those look
        \"empty\" in paragraph.text but are required for TOC / section layout.
        Also keeps the blank line immediately before major section headings.
        """
        from docx.oxml.ns import qn

        paragraphs = list(doc.paragraphs)
        empty_streak = 0
        to_remove = []
        for idx, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            has_drawing = bool(
                paragraph._element.xpath(".//*[local-name()='drawing']")
            )
            has_page_break = self._paragraph_has_page_break(paragraph)
            pPr = paragraph._element.find(qn("w:pPr"))
            has_sectpr = pPr is not None and pPr.find(qn("w:sectPr")) is not None
            if has_page_break or has_sectpr:
                empty_streak = 0
                continue
            if not text and not has_drawing:
                empty_streak += 1
                # Keep blank that sits directly above a section heading
                next_text = ""
                if idx + 1 < len(paragraphs):
                    next_text = paragraphs[idx + 1].text.strip().upper()
                protects_heading = any(
                    next_text.startswith(h) for h in self.SECTION_HEADINGS_NEEDING_SPACE
                )
                # Keep at most one blank paragraph in a row (and always keep
                # the one before a section heading)
                if empty_streak > 1 and not protects_heading:
                    to_remove.append(paragraph._element)
            else:
                empty_streak = 0
        for element in to_remove:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        if to_remove:
            logger.info("Collapsed %d extra empty paragraphs", len(to_remove))

    def _ensure_blank_before_heading(self, doc: Document, heading: str) -> None:
        """Ensure one empty paragraph sits above `heading` (spacing after tables)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        target = None
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().upper().startswith(heading.upper()):
                target = paragraph
                break
        if target is None:
            return

        prev = target._element.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            prev_text = "".join(t.text or "" for t in prev.iter(qn("w:t"))).strip()
            has_drawing = any(True for _ in prev.iter(qn("w:drawing")))
            has_page_break = any(
                br.get(qn("w:type")) == "page" for br in prev.iter(qn("w:br"))
            )
            if not prev_text and not has_drawing and not has_page_break:
                return  # blank already present

        new_p = OxmlElement("w:p")
        target._element.addprevious(new_p)
        logger.info("Inserted blank paragraph before %s", heading)

    def _normalize_toc_leaders(self, doc: Document, is_prospect: bool = False) -> None:
        """TOC rows: title | mid-height underscore line | page.

        Leaders are a fixed middle column so every row's line starts and ends on
        the same vertical edges (Client look). Underscores use noWrap so they
        never wrap into a second stray dash line.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        if not doc.tables:
            return
        table = doc.tables[0]
        first = table.rows[0].cells[0].text.strip().lower() if table.rows else ""
        if "executive" not in first:
            return

        # Preserve labels already in the template (Prospect adds Properties for Sale)
        titles = []
        for row in table.rows:
            label = (row.cells[0].text or "").strip().replace("\x07", "").replace("\r", "")
            # Title cell may still hold an old tab-leader layout
            if "\t" in label:
                label = label.split("\t", 1)[0].strip()
            # Strip any leftover underscores from prior leader merges
            label = label.strip(" _")
            if label:
                titles.append(label)
        if not titles:
            titles = [
                "Executive Summary",
                "Subject Photos",
                "Demographics",
                "Comparables",
                "Certification",
            ]

        # Exact Client TOC proportions so Prospect matches the long-template
        # rhythm (generous row height + shared leader column edge).
        col_widths = ("2985", "6192", "1078")
        if is_prospect:
            # Slightly wider title for "Properties for Sale"; keep total ~10255
            col_widths = ("3200", "5977", "1078")
        toc_font_pt = 11
        leader = "_" * 54
        # Half-points above baseline so underscore sits mid-glyph (Client look)
        leader_raise = 9
        row_height_twips = "1440"

        # Match Client fixed table width so the floating blue TOC panel lays out
        # the same way on both templates.
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblW = tblPr.find(qn("w:tblW"))
            if tblW is None:
                tblW = OxmlElement("w:tblW")
                tblPr.append(tblW)
            tblW.set(qn("w:w"), "10255")
            tblW.set(qn("w:type"), "dxa")

        # Force tblGrid to match (Word prefers grid over tcW alone)
        tblGrid = table._tbl.find(qn("w:tblGrid"))
        if tblGrid is not None:
            for gc in list(tblGrid.findall(qn("w:gridCol"))):
                tblGrid.remove(gc)
            for width in col_widths:
                gc = OxmlElement("w:gridCol")
                gc.set(qn("w:w"), width)
                tblGrid.append(gc)

        def clear_pbdr(p):
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                return
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

        def set_tight_spacing(p):
            pPr = p._p.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            # Match Client body line spacing inside the tall TOC rows
            spacing.set(qn("w:line"), "276")
            spacing.set(qn("w:lineRule"), "auto")
            # Zero cell-like indent so titles/leaders share a clean left edge
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                pPr.remove(ind)

        def set_cell_margins(cell, margin_twips: str = "0"):
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = tcPr.find(qn("w:tcMar"))
            if tcMar is not None:
                tcPr.remove(tcMar)
            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), margin_twips)
                el.set(qn("w:type"), "dxa")
                tcMar.append(el)
            tcPr.append(tcMar)

        def style_run(run, size_pt=toc_font_pt, raise_hp=0):
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.underline = False
            run.font.name = "Calibri"
            rPr = run._r.get_or_add_rPr()
            u = rPr.find(qn("w:u"))
            if u is not None:
                rPr.remove(u)
            color = rPr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rPr.append(color)
            color.set(qn("w:val"), "FFFFFF")
            color.set(qn("w:themeColor"), "background1")
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            for attr in ("ascii", "hAnsi", "cs"):
                rFonts.set(qn(f"w:{attr}"), "Calibri")
            half = str(int(size_pt * 2))
            for tag in ("sz", "szCs"):
                el = rPr.find(qn(f"w:{tag}"))
                if el is None:
                    el = OxmlElement(f"w:{tag}")
                    rPr.append(el)
                el.set(qn("w:val"), half)
            pos = rPr.find(qn("w:position"))
            if raise_hp:
                if pos is None:
                    pos = OxmlElement("w:position")
                    rPr.append(pos)
                pos.set(qn("w:val"), str(raise_hp))
            elif pos is not None:
                rPr.remove(pos)

        def make_cell(width_twips: str):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), width_twips)
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), "center")
            tcPr.append(vAlign)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            tc.append(p)
            return tc

        def rebuild_row_cells(row):
            """Force exactly 3 cells (unmerge prior tab-leader layout)."""
            tr = row._tr
            for tc in list(tr.findall(qn("w:tc"))):
                tr.remove(tc)
            for width in col_widths:
                tr.append(make_cell(width))

        for ri, title in enumerate(titles):
            if ri >= len(table.rows):
                break
            row = table.rows[ri]

            # Preserve page from prior layouts before rebuilding cells
            raw0 = (row.cells[0].text or "").replace("\x07", "").replace("\r", "")
            if "\t" in raw0:
                existing_page = raw0.split("\t", 1)[1].strip() or "—"
            else:
                try:
                    existing_page = (
                        (row.cells[2].text or "").strip().replace("\x07", "") or "—"
                    )
                except Exception:
                    existing_page = "—"

            trPr = row._tr.get_or_add_trPr()
            trH = trPr.find(qn("w:trHeight"))
            if trH is None:
                trH = OxmlElement("w:trHeight")
                trPr.append(trH)
            trH.set(qn("w:val"), row_height_twips)
            trH.set(qn("w:hRule"), "atLeast")

            rebuild_row_cells(row)
            cells = row.cells

            specs = (
                (0, title, WD_ALIGN_PARAGRAPH.LEFT, 0),
                (1, leader, WD_ALIGN_PARAGRAPH.LEFT, leader_raise),
                (2, existing_page, WD_ALIGN_PARAGRAPH.RIGHT, 0),
            )
            for ci, text, align, raise_hp in specs:
                cell = cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, "60")
                tcPr = cell._tc.get_or_add_tcPr()
                # Never wrap title or leader (prevents stray second underscore line)
                if ci in (0, 1):
                    if tcPr.find(qn("w:noWrap")) is None:
                        tcPr.append(OxmlElement("w:noWrap"))
                p = cell.paragraphs[0]
                clear_pbdr(p)
                set_tight_spacing(p)
                p.alignment = align
                p.clear()
                clean = text.replace("\n", " ").replace("\r", " ").strip()
                style_run(p.add_run(clean), raise_hp=raise_hp)

        logger.info(
            "Normalized TOC leaders (Client spacing: %spt, row=%s twips, cols=%s)",
            toc_font_pt,
            row_height_twips,
            col_widths,
        )

    # TOC label -> body heading used to resolve the real page number
    TOC_SECTION_HEADINGS = (
        ("Executive Summary", "EXECUTIVE SUMMARY"),
        ("Subject Photos", "SUBJECT PHOTOS"),
        ("Demographics", "DEMOGRAPHIC ANALYSIS"),
        ("Comparables", "PROPERTY COMPARABLES"),
        ("Certification", "CERTIFICATION AND DISCLAIMERS"),
    )

    # Prospect short-form: TOC labels map to actual short-form headings.
    # "Properties for Sale" is not included (no for-sale listings section).
    PROSPECT_TOC_SECTION_HEADINGS = (
        ("Executive Summary", "EXECUTIVE SUMMARY"),
        ("Demographics", "PROPERTY SUMMARY"),
        ("Subject Photos", "SUBJECT PHOTOS"),
        ("Comparables", "PROPERTY COMPARABLES"),
        ("Certification", "CERTIFICATION AND DISCLAIMERS"),
    )

    def _sync_prospect_toc_labels(self, doc: Document) -> None:
        """Keep Prospect TOC labels in body order; drop Properties for Sale rows."""
        if not doc.tables:
            return
        table = doc.tables[0]
        first = (table.rows[0].cells[0].text or "").strip().lower() if table.rows else ""
        if "executive" not in first:
            return

        labels = [label for label, _ in self.PROSPECT_TOC_SECTION_HEADINGS]

        # Remove "Properties for Sale" (and similar) rows from the TOC table
        for row in list(table.rows):
            raw = (row.cells[0].text or "").strip().replace("\x07", "").replace("\r", "")
            if "\t" in raw:
                raw = raw.split("\t", 1)[0].strip()
            key = raw.strip(" _").lower()
            if "properties for sale" in key or key == "for sale":
                parent = row._tr.getparent()
                if parent is not None:
                    parent.remove(row._tr)

        # Trim any leftover extra rows beyond the Prospect TOC set
        while len(table.rows) > len(labels):
            parent = table.rows[-1]._tr.getparent()
            if parent is None:
                break
            parent.remove(table.rows[-1]._tr)

        for i, label in enumerate(labels):
            if i >= len(table.rows):
                break
            cell = table.rows[i].cells[0]
            current = (cell.text or "").strip().replace("\x07", "").replace("\r", "")
            if "\t" in current:
                current = current.split("\t", 1)[0].strip()
            if current != label:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = label
                    for run in cell.paragraphs[0].runs[1:]:
                        run.text = ""
                else:
                    cell.text = label

        logger.info("Synced Prospect TOC labels (no Properties for Sale): %s", labels)

    def _refresh_toc_page_numbers(
        self,
        output_path,
        is_prospect: bool = False,
        street_view_path: Optional[str] = None,
    ) -> None:
        """Set TOC page numbers from where each section actually lands in the Word doc.

        Uses Word COM so pagination matches what the user sees after generation
        (comps, demographics, etc. can shift pages vs the static template).

        Also restores cover visuals Word otherwise drops: white diagonal corner
        (Group 13) and bottom-left Street View photo.
        """
        try:
            import pythoncom  # type: ignore
            import win32com.client  # type: ignore
        except ImportError:
            logger.warning(
                "win32com unavailable — TOC page numbers remain template placeholders"
            )
            return

        path = str(Path(output_path).resolve())
        word = None
        doc = None
        pythoncom.CoInitialize()
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(FileName=path, AddToRecentFiles=False)
            # Force layout so page numbers are accurate
            doc.Repaginate()

            total_pages = int(doc.ComputeStatistics(2))  # wdStatisticPages
            # Resolve each body heading page (skip TOC page hits by matching uppercase)
            resolved = {}
            toc_map = (
                self.PROSPECT_TOC_SECTION_HEADINGS
                if is_prospect
                else self.TOC_SECTION_HEADINGS
            )
            for toc_label, heading in toc_map:
                # Start search after page 2 (cover/TOC) when possible
                start_page = 3 if total_pages >= 3 else 1
                word.Selection.GoTo(What=1, Which=1, Count=start_page)  # wdGoToPage
                rng = doc.Range(word.Selection.Start, doc.Content.End)
                find = rng.Find
                find.ClearFormatting()
                find.Text = heading
                find.MatchCase = True
                find.Forward = True
                find.Wrap = 0  # wdFindStop
                found = bool(find.Execute())
                if not found:
                    # Fallback: case-insensitive anywhere after page 1
                    word.Selection.GoTo(What=1, Which=1, Count=2)
                    rng = doc.Range(word.Selection.Start, doc.Content.End)
                    find = rng.Find
                    find.ClearFormatting()
                    find.Text = heading
                    find.MatchCase = False
                    find.Forward = True
                    find.Wrap = 0
                    found = bool(find.Execute())
                if found:
                    resolved[toc_label] = int(rng.Information(3))  # wdActiveEndPageNumber
                else:
                    logger.warning("TOC heading not found for page resolve: %s", heading)

            if not resolved:
                logger.warning("Could not resolve any TOC section pages")

            # Update first table that looks like the TOC
            toc_table = None
            updated = 0
            if resolved:
                for i in range(1, doc.Tables.Count + 1):
                    tbl = doc.Tables(i)
                    try:
                        if "executive" in (tbl.Cell(1, 1).Range.Text or "").lower():
                            toc_table = tbl
                            break
                    except Exception:
                        continue
            if resolved and toc_table is None:
                logger.warning("TOC table not found for page-number refresh")

            # Drop Properties for Sale rows in Word before writing page numbers
            if toc_table is not None and is_prospect:
                for row_idx in range(toc_table.Rows.Count, 0, -1):
                    try:
                        raw = (
                            (toc_table.Cell(row_idx, 1).Range.Text or "")
                            .replace("\r", "")
                            .replace("\x07", "")
                        )
                        label = raw.split("\t", 1)[0].strip().lower()
                        if "properties for sale" in label or label == "for sale":
                            toc_table.Rows(row_idx).Delete()
                    except Exception:
                        continue

            if toc_table is not None:
                # Build case-insensitive lookup for TOC labels
                resolved_ci = {k.strip().lower(): v for k, v in resolved.items()}

                for row_idx in range(1, toc_table.Rows.Count + 1):
                    try:
                        cell = toc_table.Cell(row_idx, 1)
                        raw = (cell.Range.Text or "").replace("\r", "").replace("\x07", "")
                        if "\t" in raw:
                            label = raw.split("\t", 1)[0].strip()
                        else:
                            label = raw.strip()
                        page_num = resolved.get(label)
                        if page_num is None:
                            page_num = resolved_ci.get(label.lower())
                        if page_num is None:
                            # Prefix match (e.g. "Executive Summary\rSomething")
                            for key, val in resolved_ci.items():
                                if label.lower().startswith(key) or key.startswith(label.lower()):
                                    page_num = val
                                    break
                        if page_num is None:
                            logger.warning("No resolved page for TOC label %r", label)
                            continue

                        para = cell.Range.Paragraphs(1)
                        para_text = (para.Range.Text or "").replace("\r", "").replace("\x07", "")
                        tab_idx = para_text.find("\t")
                        if tab_idx >= 0:
                            page_start = para.Range.Start + tab_idx + 1
                            page_end = para.Range.End - 1
                            page_rng = doc.Range(page_start, page_end)
                            page_rng.Text = str(page_num)
                            page_rng.Font.Color = 16777215  # white
                            page_rng.Font.Size = 11
                        else:
                            # Legacy 3-column TOC
                            page_cell = toc_table.Cell(row_idx, 3)
                            rng = page_cell.Range
                            rng.MoveEnd(Unit=1, Count=-1)
                            rng.Text = str(page_num)
                            rng.Font.Color = 16777215
                            rng.Font.Size = 11
                            page_cell.Range.ParagraphFormat.Alignment = 2  # wdAlignParagraphRight
                        updated += 1
                    except Exception as cell_exc:
                        logger.warning("TOC row %s update failed: %s", row_idx, cell_exc)

                doc.Repaginate()
                logger.info(
                    "Refreshed TOC page numbers from live pagination (%s/%s rows): %s (doc pages=%s)",
                    updated,
                    toc_table.Rows.Count,
                    resolved,
                    total_pages,
                )

            # Cover: white diagonal corner + Street View (same Word session —
            # hand-edited VML imagedata corrupts the package).
            try:
                self._word_fix_cover_shapes(doc, street_view_path)
            except Exception as cover_exc:
                logger.warning("Cover shape fix via Word COM failed: %s", cover_exc)

            doc.Save()
        except Exception as exc:
            logger.warning("TOC page-number refresh failed: %s", exc)
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _word_fix_cover_shapes(self, doc, street_view_path: Optional[str] = None) -> None:
        """Restore cover corner + property photo using Word's shape API (safe OOXML).

        Group 13 children ship with Fill.Visible=False after open; Text Box 21
        loses its picture. Avoid ZOrder/Delete (crash Word 2007) — AddPicture
        places the Street View on top of the white corner shapes.
        """
        sv = None
        if street_view_path and Path(street_view_path).exists():
            sv = str(Path(street_view_path).resolve())

        group = None
        tb21 = None
        addr = None
        hero = None
        for i in range(1, int(doc.Shapes.Count) + 1):
            shape = doc.Shapes(i)
            name = shape.Name or ""
            if name == "Group 13":
                group = shape
            elif name == "Text Box 21":
                tb21 = shape
            elif name == "Picture 8":
                hero = shape
            elif name == "Text Box 26":
                try:
                    if not shape.TextFrame.HasText:
                        continue
                    text = (shape.TextFrame.TextRange.Text or "").strip()
                    if "PREPARED" in text.upper():
                        continue
                    if text and len(text) < 100 and any(ch.isdigit() for ch in text):
                        addr = shape
                except Exception:
                    pass

        # Fallback: any short street-like textbox that isn't prepared-by/for
        if addr is None:
            for i in range(1, int(doc.Shapes.Count) + 1):
                shape = doc.Shapes(i)
                try:
                    if not shape.TextFrame.HasText:
                        continue
                    text = (shape.TextFrame.TextRange.Text or "").strip()
                except Exception:
                    continue
                if "PREPARED" in text.upper():
                    continue
                if (
                    text
                    and len(text) < 100
                    and any(ch.isdigit() for ch in text)
                    and "," in text
                    and (shape.Name or "").startswith("Text Box")
                ):
                    addr = shape
                    break

        if group is not None:
            try:
                for j in range(1, int(group.GroupItems.Count) + 1):
                    item = group.GroupItems(j)
                    item.Fill.Visible = -1  # msoTrue
                    item.Fill.Solid()
                    item.Fill.ForeColor.RGB = 16777215  # white
                logger.info(
                    "Cover fix: whitened Group 13 corner shapes (%s items)",
                    group.GroupItems.Count,
                )
            except Exception as exc:
                logger.warning("Cover fix: could not whiten Group 13: %s", exc)

        # Desired cover (mock): address nestled in the WHITE DIAGONAL CORNER
        # (left wedge under the slant), photo in the lower-left slot — separate.
        # Do NOT stack photo tightly under address (that reads as one mixed block).
        hero_bottom = 438.0
        if hero is not None:
            try:
                hero_bottom = float(hero.Top) + float(hero.Height)
            except Exception:
                pass

        # Template slot for cover photo (Text Box 21) — capture before parking it
        photo_left = -37.0
        photo_top = 522.0
        photo_width = 284.0
        photo_height = 233.0  # designed slot height (do not shrink to a short strip)
        if tb21 is not None:
            try:
                photo_left = float(tb21.Left)
                photo_top = float(tb21.Top)
                photo_width = float(tb21.Width)
                photo_height = float(tb21.Height)
            except Exception:
                pass

        # Address in diagonal white corner: flush left, tall enough for 2 full lines.
        addr_top = min(365.0, hero_bottom - 70.0)
        addr_left = min(-10.0, photo_left + 10.0)  # more left, near photo column
        if addr is not None:
            try:
                raw = (addr.TextFrame.TextRange.Text or "").replace("\r", "").strip()
                # Collapse soft breaks then re-split to street / city lines
                raw = raw.replace("\v", " ").replace("\n", " ")
                while "  " in raw:
                    raw = raw.replace("  ", " ")
                parts = [p.strip() for p in raw.split(",") if p.strip()]
                if len(parts) >= 3:
                    addr.TextFrame.TextRange.Text = (
                        f"{parts[0]},\r{', '.join(parts[1:])}"
                    )
                elif len(parts) == 2:
                    addr.TextFrame.TextRange.Text = f"{parts[0]},\r{parts[1]}"
                try:
                    addr.Line.Visible = 0
                except Exception:
                    pass
                try:
                    addr.Fill.Visible = 0
                except Exception:
                    pass
                # Tight margins so both lines fit and text sits flush left
                try:
                    tf = addr.TextFrame
                    tf.MarginLeft = 2.0
                    tf.MarginRight = 2.0
                    tf.MarginTop = 1.0
                    tf.MarginBottom = 1.0
                    tf.WordWrap = True
                except Exception:
                    pass
                addr.Left = addr_left
                addr.Top = addr_top
                try:
                    addr.Width = 300.0
                    # ~2 lines @ ~14–16pt + margins — 48pt was clipping line 2
                    addr.Height = 72.0
                except Exception:
                    pass
                logger.info(
                    "Cover fix: address L=%.1f T=%.1f H=%.1f (hero_bot=%.1f)",
                    float(addr.Left),
                    float(addr.Top),
                    float(addr.Height),
                    hero_bottom,
                )
            except Exception as exc:
                logger.warning("Cover fix: address move failed: %s", exc)

        if sv and tb21 is not None:
            # Designed lower-left photo slot — keep near-template height
            left = photo_left if photo_left > -200 else -37.0
            top = photo_top
            width = photo_width
            height = photo_height if photo_height >= 200 else 233.0
            max_bottom = 705.0
            addr_bot = None
            if addr is not None:
                try:
                    addr_bot = float(addr.Top) + float(addr.Height)
                    if top < addr_bot + 24.0:
                        top = addr_bot + 24.0
                except Exception:
                    addr_bot = None
            if top + height > max_bottom:
                overflow = (top + height) - max_bottom
                min_top = (addr_bot + 24.0) if addr_bot is not None else (top - overflow)
                top = max(top - overflow, min_top)
                if top + height > max_bottom:
                    height = max(200.0, max_bottom - top)

            # Hide empty template frame (border made address+photo look one box)
            try:
                tb21.Line.Visible = 0
                tb21.Fill.Visible = 0
            except Exception:
                pass
            try:
                # wdRelativeHorizontalPositionPage = 1
                tb21.RelativeHorizontalPosition = 1
                tb21.Left = -2500.0
                tb21.Width = 1.0
                tb21.Height = 1.0
            except Exception:
                try:
                    tb21.Left = -2500.0
                except Exception:
                    pass

            pic = doc.Shapes.AddPicture(
                FileName=sv,
                LinkToFile=False,
                SaveWithDocument=True,
                Left=left,
                Top=top,
                Width=width,
                Height=height,
            )
            try:
                pic.Name = "CoverStreetView"
            except Exception:
                pass
            try:
                pic.Line.Visible = 0
            except Exception:
                pass
            gap = None
            if addr is not None:
                try:
                    gap = float(top) - (float(addr.Top) + float(addr.Height))
                except Exception:
                    pass
            logger.info(
                "Cover fix: photo slot (%.1f, %.1f) h=%.1f gap_from_addr=%s",
                left,
                top,
                height,
                f"{gap:.1f}" if gap is not None else "n/a",
            )
        elif not sv:
            logger.warning("Cover fix: no Street View path — cover photo not added")
        else:
            logger.warning("Cover fix: Text Box 21 not found — cover photo not added")

    def _ensure_page_break_before_heading(self, doc: Document, heading: str) -> None:
        """Insert a page break immediately before `heading` if one is missing.

        The TOC page uses a full-page floating blue shape. Body text after the
        TOC table must start on the next page or it renders on top of the TOC.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        target = None
        for paragraph in doc.paragraphs:
            if heading in paragraph.text:
                target = paragraph
                break
        if target is None:
            return

        # Already has pageBreakBefore on the heading?
        pPr = target._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
            return

        # Look back a few siblings for an existing page break (skip blank paras)
        prev = target._element.getprevious()
        checked = 0
        while prev is not None and checked < 4:
            if prev.tag == qn("w:p"):
                if any(br.get(qn("w:type")) == "page" for br in prev.iter(qn("w:br"))):
                    return
                # Stop looking once we hit real content (e.g. TOC table already passed)
                text = "".join(t.text or "" for t in prev.iter(qn("w:t"))).strip()
                has_drawing = any(True for _ in prev.iter(qn("w:drawing")))
                if text or has_drawing:
                    break
                checked += 1
                prev = prev.getprevious()
            elif prev.tag == qn("w:tbl"):
                break
            else:
                prev = prev.getprevious()

        # Insert an explicit page-break paragraph before the heading
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_br = OxmlElement("w:br")
        new_br.set(qn("w:type"), "page")
        new_r.append(new_br)
        new_p.append(new_r)
        target._element.addprevious(new_p)
        logger.info("Inserted page break before %s", heading)

    # Template accent colors -> theme replacements (primary, secondary)
    COLOR_THEMES = {
        "light blue": ("0070C0", "00B0F0"),
        "dark blue": ("1F3864", "2E5496"),
        "red": ("C00000", "E31C23"),
        "green": ("2E7D32", "5BB85C"),
    }
    # UI / legacy aliases
    COLOR_THEME_ALIASES = {
        "blue": "light blue",
        "lightblue": "light blue",
        "darkblue": "dark blue",
    }
    TEMPLATE_ACCENTS = ("0070C0", "00B0F0")
    TEMPLATE_HEADER_FOOTER_ACCENT = "0066CC"

    def _normalize_color_theme(self, color_theme: Optional[str]) -> str:
        """Map UI values like light-blue / darkblue -> canonical theme keys."""
        theme = (color_theme or "").strip().lower().replace("-", " ").replace("_", " ")
        theme = " ".join(theme.split())
        return self.COLOR_THEME_ALIASES.get(theme, theme)

    @staticmethod
    def _format_cover_address(address: Optional[str]) -> str:
        """Format cover address on two lines like the mock (street then city/state)."""
        text = (address or "").strip()
        if not text:
            return ""
        # Prefer split before city/state when ", ST ZIP" or ", City, ST" pattern exists
        # e.g. "5700 Granite Pkwy, Frisco, TX 75034" -> street / "Frisco, TX 75034"
        parts = [p.strip() for p in text.split(",") if p.strip()]
        if len(parts) >= 3:
            return f"{parts[0]},\n{', '.join(parts[1:])}"
        if len(parts) == 2:
            return f"{parts[0]},\n{parts[1]}"
        # No commas — try last token that looks like a state abbrev
        tokens = text.split()
        if len(tokens) >= 3 and len(tokens[-2]) == 2 and tokens[-2].isalpha():
            # "... Frisco TX 75034"
            return f"{' '.join(tokens[:-3])} {tokens[-3]},\n{tokens[-2]} {tokens[-1]}"
        return text

    def _polish_cover_layout(self, doc: Document, color_theme: Optional[str] = None) -> None:
        """Match the designed cover: tinted hero, left-flush bars, address in corner.

        - Keep Picture 8 (diagonal hero) but apply a theme color shade/tint
        - Title + date bars flush to the left (one-sided round on the right)
        - White title text; address nestled in the bottom-left white diagonal corner
        """
        from lxml import etree
        from docx.oxml.ns import qn

        a_av = "{http://schemas.openxmlformats.org/drawingml/2006/main}avLst"
        a_gd = "{http://schemas.openxmlformats.org/drawingml/2006/main}gd"
        w_rpr = qn("w:rPr")
        w_color = qn("w:color")

        theme = self._normalize_color_theme(color_theme)
        target = self.COLOR_THEMES.get(theme) or self.COLOR_THEMES["light blue"]
        primary, secondary = target

        title_shaped = date_shaped = title_whitened = addr_placed = 0

        for drawing in doc.element.body.iter(qn("w:drawing")):
            docPr = next(drawing.iter(qn("wp:docPr")), None)
            if docPr is None:
                continue
            name = (docPr.get("name") or "").strip()
            texts = " ".join(
                (t.text or "") for t in drawing.iter(qn("w:t")) if t.text
            ).strip()

            # Skip off-canvas leftovers (ox far negative, no left-align)
            posH = next(drawing.iter(qn("wp:positionH")), None)
            ox_el = posH.find(qn("wp:posOffset")) if posH is not None else None
            align_el = posH.find(qn("wp:align")) if posH is not None else None
            ox = int(ox_el.text) if ox_el is not None and ox_el.text else None
            is_left_aligned = align_el is not None and (align_el.text or "") == "left"
            if ox is not None and ox < -5_000_000 and not is_left_aligned:
                continue

            # Title bar: flush left, one-sided round on the right (like the mock)
            if "Broker Opinion of Value" in texts and name.startswith("Rectangle"):
                self._cover_force_left_align(drawing)
                for geom in drawing.iter(qn("a:prstGeom")):
                    geom.set("prst", "round1Rect")
                    av = geom.find(a_av)
                    if av is None:
                        av = etree.SubElement(geom, a_av)
                    else:
                        av.clear()
                    gd = etree.SubElement(av, a_gd)
                    gd.set("name", "adj")
                    gd.set("fmla", "val 50000")
                    title_shaped += 1
                for run in drawing.iter(qn("w:r")):
                    rPr = run.find(w_rpr)
                    if rPr is None:
                        rPr = etree.Element(w_rpr)
                        run.insert(0, rPr)
                    color = rPr.find(w_color)
                    if color is None:
                        color = etree.SubElement(rPr, w_color)
                    color.set(qn("w:val"), "FFFFFF")
                    title_whitened += 1
                continue

            # Date / property-name pill — flush left under the title
            if name.startswith("Text Box") and texts.startswith("Date:"):
                self._cover_force_left_align(drawing)
                for geom in drawing.iter(qn("a:prstGeom")):
                    geom.set("prst", "round1Rect")
                    av = geom.find(a_av)
                    if av is None:
                        av = etree.SubElement(geom, a_av)
                    else:
                        av.clear()
                    gd = etree.SubElement(av, a_gd)
                    gd.set("name", "adj")
                    gd.set("fmla", "val 35000")
                    date_shaped += 1
                    break
                continue

            # Address in the bottom-left white diagonal corner
            if "{{address}}" in texts and "PREPARED" not in texts.upper():
                self._cover_place_address_in_corner(drawing)
                addr_placed += 1
                continue

            # Cover Street View slot must sit above the white corner group
            if name == "Text Box 21" or "{{main_img}}" in (
                (docPr.get("descr") or "")
            ):
                anchor = next(
                    (a for a in drawing if etree.QName(a).localname == "anchor"),
                    None,
                )
                if anchor is not None:
                    anchor.set("behindDoc", "0")
                    anchor.set("relativeHeight", "251670000")
                continue

        corner_ok = self._ensure_cover_diagonal_corner(doc)
        tinted = self._tint_cover_hero(doc, primary, secondary)
        logger.info(
            "Cover polish: title_round=%s date_round=%s title_white=%s "
            "address_corner=%s white_corner=%s hero_tint=%s theme=%s",
            title_shaped,
            date_shaped,
            title_whitened,
            addr_placed,
            corner_ok,
            tinted,
            theme or "light blue",
        )

    @staticmethod
    def _cover_force_left_align(drawing) -> None:
        """Pin a cover shape to the left page edge (matches mock left corner stack)."""
        from lxml import etree
        from docx.oxml.ns import qn

        posH = next(drawing.iter(qn("wp:positionH")), None)
        if posH is None:
            return
        # Clear offset-based placement; use align=left
        for child in list(posH):
            if etree.QName(child).localname in ("posOffset", "align"):
                posH.remove(child)
        align = etree.SubElement(posH, qn("wp:align"))
        align.text = "left"
        if posH.get("relativeFrom") is None:
            posH.set("relativeFrom", "page")

    def _cover_place_address_in_corner(self, drawing) -> None:
        """Sit address in the white diagonal corner (bottom-left), flush left."""
        from lxml import etree
        from docx.oxml.ns import qn

        # Address flush-left in white diagonal corner (~365pt)
        left_emu = str(int(-10 / 72 * 914400))  # ~-0.14" — more left
        top_emu = str(int(365 / 72 * 914400))

        posH = next(drawing.iter(qn("wp:positionH")), None)
        posV = next(drawing.iter(qn("wp:positionV")), None)
        if posH is not None:
            for child in list(posH):
                if etree.QName(child).localname in ("posOffset", "align"):
                    posH.remove(child)
            posH.set("relativeFrom", "page")
            off = etree.SubElement(posH, qn("wp:posOffset"))
            off.text = left_emu
        if posV is not None:
            for child in list(posV):
                if etree.QName(child).localname in ("posOffset", "align"):
                    posV.remove(child)
            posV.set("relativeFrom", "page")
            off = etree.SubElement(posV, qn("wp:posOffset"))
            off.text = top_emu

        # Keep address textbox above the white corner shape
        anchor = next(
            (a for a in drawing if etree.QName(a).localname == "anchor"),
            None,
        )
        if anchor is not None:
            anchor.set("behindDoc", "0")
            anchor.set("relativeHeight", "251660000")

    def _ensure_cover_diagonal_corner(self, doc: Document) -> bool:
        """Force the bottom-left white diagonal corner (Group 13) to be visible.

        Template children use grpFill/scheme bg1 which often fails to paint a
        solid white wedge over the hero photo. Replace with explicit FFFFFF and
        keep the group in front of Picture 8.
        """
        from lxml import etree
        from docx.oxml.ns import qn

        A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
        WPS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}"
        WPG = "{http://schemas.microsoft.com/office/word/2010/wordprocessingGroup}"
        fixed = False

        def _solid_white():
            solid = etree.Element(A + "solidFill")
            srgb = etree.SubElement(solid, A + "srgbClr")
            srgb.set("val", "FFFFFF")
            return solid

        def _paint_spPr(spPr) -> bool:
            changed = False
            for child in list(spPr):
                if etree.QName(child).localname in (
                    "grpFill",
                    "solidFill",
                    "noFill",
                    "gradFill",
                ):
                    spPr.remove(child)
                    changed = True
            geom = spPr.find(A + "prstGeom")
            solid = _solid_white()
            if geom is not None:
                geom.addnext(solid)
            else:
                spPr.append(solid)
            return True

        for drawing in doc.element.body.iter(qn("w:drawing")):
            docPr = next(drawing.iter(qn("wp:docPr")), None)
            if docPr is None or (docPr.get("name") or "") != "Group 13":
                continue
            posH = next(drawing.iter(qn("wp:positionH")), None)
            ox_el = posH.find(qn("wp:posOffset")) if posH is not None else None
            ox = int(ox_el.text) if ox_el is not None and ox_el.text else 0
            if ox < -5_000_000:
                continue

            anchor = next(
                (a for a in drawing if etree.QName(a).localname == "anchor"),
                None,
            )
            if anchor is not None:
                # In front of hero photo (Picture 8 is behindDoc=1 / lower z)
                anchor.set("behindDoc", "0")
                anchor.set("relativeHeight", "251659000")

            # Group fill -> solid white (wpg:grpSpPr or a:grpSpPr)
            for grpSpPr in list(drawing.iter(WPG + "grpSpPr")) + list(
                drawing.iter(A + "grpSpPr")
            ):
                for child in list(grpSpPr):
                    local = etree.QName(child).localname
                    if local in ("solidFill", "noFill", "gradFill", "grpFill"):
                        grpSpPr.remove(child)
                xfrm = grpSpPr.find(A + "xfrm")
                solid = _solid_white()
                if xfrm is not None:
                    xfrm.addnext(solid)
                else:
                    grpSpPr.insert(0, solid)

            # Child shapes: wps:spPr / a:spPr
            for spPr in list(drawing.iter(WPS + "spPr")) + list(drawing.iter(A + "spPr")):
                if _paint_spPr(spPr):
                    fixed = True

            # Neutralize accent1 style fillRef on wps:style / a:style
            for style in list(drawing.iter(WPS + "style")) + list(drawing.iter(A + "style")):
                fill_ref = style.find(A + "fillRef")
                if fill_ref is None:
                    continue
                for child in list(fill_ref):
                    fill_ref.remove(child)
                srgb = etree.SubElement(fill_ref, A + "srgbClr")
                srgb.set("val", "FFFFFF")
                fixed = True

            fixed = True
            logger.info("Ensured white diagonal cover corner (Group 13)")
            break

        return fixed

    def _tint_cover_hero(self, doc: Document, primary_hex: str, secondary_hex: str) -> bool:
        """Apply a monochrome theme shade to the large cover Picture 8 hero."""
        from io import BytesIO
        from docx.oxml.ns import qn

        try:
            from PIL import Image, ImageOps
        except ImportError:
            logger.warning("Pillow missing — cannot tint cover hero")
            return False

        def _rgb(h: str):
            h = (h or "0070C0").strip().lstrip("#")
            return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))

        primary = _rgb(primary_hex)
        secondary = _rgb(secondary_hex)
        # Dark end of duotone (shadows) + light end (highlights) for a wash look
        dark = tuple(max(0, c // 5) for c in primary)
        light = tuple(min(255, int(c + (255 - c) * 0.55)) for c in secondary)

        tinted = False
        for drawing in doc.element.body.iter(qn("w:drawing")):
            docPr = next(drawing.iter(qn("wp:docPr")), None)
            if docPr is None or (docPr.get("name") or "") != "Picture 8":
                continue
            ext = next(drawing.iter(qn("wp:extent")), None)
            cx = int(ext.get("cx") or 0) if ext is not None else 0
            if cx < 6_000_000:
                continue  # ignore leftover tiny copies
            blip = next(drawing.iter(qn("a:blip")), None)
            if blip is None:
                continue
            rid = blip.get(qn("r:embed"))
            part = doc.part.related_parts.get(rid) if rid else None
            if part is None or not getattr(part, "_blob", None):
                continue
            try:
                img = Image.open(BytesIO(part._blob)).convert("RGB")
                gray = ImageOps.grayscale(img)
                colored = ImageOps.colorize(gray, black=dark, white=light)
                buf = BytesIO()
                # Keep JPEG for typical cover parts
                partname = str(getattr(part, "partname", "") or "").lower()
                ctype = str(getattr(part, "content_type", "") or "").lower()
                if partname.endswith(".png") or "image/png" in ctype:
                    colored.convert("RGBA").save(buf, format="PNG", optimize=True)
                else:
                    colored.save(buf, format="JPEG", quality=85, optimize=True)
                part._blob = buf.getvalue()
                tinted = True
                logger.info(
                    "Tinted cover hero Picture 8 with theme shade %s→%s (%s bytes)",
                    primary_hex,
                    secondary_hex,
                    len(part._blob),
                )
            except Exception as exc:
                logger.error("Cover hero tint failed: %s", exc)
        return tinted

    def _apply_color_theme(self, output_path, color_theme: str):
        """Recolor report accents, including header/footer bars.

        Applies the same way for red, green, dark blue, and light blue:
        - cover title bar + body accents: 0070C0 / 00B0F0 -> theme primary/secondary
        - header/footer bars: 0066CC (and VML #06c) -> theme primary

        Handles DrawingML (`srgbClr val="0066CC"`), Word hex attrs, and VML
        `fillcolor="#06c"` (short CSS form of #0066CC after Word COM save).
        """
        theme = self._normalize_color_theme(color_theme)
        target = self.COLOR_THEMES.get(theme)
        if not target:
            logger.info("Unknown color theme %r — leaving template blues", color_theme)
            return

        import zipfile
        import shutil
        import tempfile

        src_primary, src_secondary = self.TEMPLATE_ACCENTS
        src_header_footer = self.TEMPLATE_HEADER_FOOTER_ACCENT
        dst_primary, dst_secondary = target

        def _variants(hex6: str):
            """All common Word XML spellings of a 6-digit hex color."""
            h = hex6.upper()
            lo = h.lower()
            short = None
            if len(h) == 6 and h[0] == h[1] and h[2] == h[3] and h[4] == h[5]:
                short = f"{h[0]}{h[2]}{h[4]}"  # e.g. 0066CC -> 06C
            out = [
                f'"{h}"',
                f'"{lo}"',
                f"#{h}",
                f"#{lo}",
            ]
            if short:
                out.extend([f"#{short}", f"#{short.lower()}"])
            return out

        def recolor(text: str) -> str:
            # Header/footer first so 0066CC never accidentally matches a later pass.
            pairs = (
                (src_header_footer, dst_primary),
                (src_primary, dst_primary),
                (src_secondary, dst_secondary),
            )
            for src, dst in pairs:
                if src.upper() == dst.upper():
                    continue  # light blue body accents already match template
                for form in _variants(src):
                    if form.startswith("#"):
                        text = text.replace(form, f"#{dst}")
                        alt = form.upper() if form == form.lower() else form.lower()
                        if alt != form:
                            text = text.replace(alt, f"#{dst}")
                    else:
                        text = text.replace(form, f'"{dst}"')
            return text

        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
            os.close(tmp_fd)
            with zipfile.ZipFile(output_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    data = zin.read(item)
                    if item.startswith("word/") and item.endswith(".xml"):
                        data = recolor(data.decode("utf-8", "ignore")).encode("utf-8")
                    zout.writestr(item, data)
            shutil.move(tmp_path, output_path)
            logger.info(
                "Applied '%s' color theme (primary=%s secondary=%s) to %s",
                theme,
                dst_primary,
                dst_secondary,
                output_path,
            )
        except Exception as exc:
            logger.error(f"Failed to apply color theme '{theme}': {exc}")

    def _replace_text_in_runs(self, paragraph, placeholder: str, replacement: str):
        """Replace text while preserving run formatting.

        Newlines in the replacement become real paragraphs — never soft line
        breaks. Soft breaks inside justified (jc=both) paragraphs stretch the
        short last line ("companies          and          startups.").
        """
        from docx.oxml.ns import qn

        if placeholder not in paragraph.text:
            return

        replacement = (replacement or "").replace("--", "–")
        # Split into paragraph chunks (blank-line or single newline)
        chunks = [c.strip() for c in re.split(r"\r?\n\s*\r?\n|\r?\n", replacement) if c.strip()]
        if not chunks:
            chunks = [""]

        first, *rest = chunks

        # Fast path: placeholder lives in a single run
        replaced = False
        for run in paragraph.runs:
            if placeholder in (run.text or ""):
                run.text = (run.text or "").replace(placeholder, first)
                replaced = True
                break

        if not replaced:
            full_text = paragraph.text
            placeholder_start = full_text.find(placeholder)
            if placeholder_start == -1:
                return

            char_formats = []
            for run in paragraph.runs:
                for _char in run.text or "":
                    char_formats.append(
                        {
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_name": run.font.name,
                            "font_size": run.font.size,
                        }
                    )

            paragraph.clear()

            def _apply_format(run, format_info):
                if not format_info:
                    return
                if format_info.get("bold") is not None:
                    run.bold = format_info["bold"]
                if format_info.get("italic") is not None:
                    run.italic = format_info["italic"]
                if format_info.get("underline") is not None:
                    run.underline = format_info["underline"]
                if format_info.get("font_name"):
                    run.font.name = format_info["font_name"]
                if format_info.get("font_size"):
                    run.font.size = format_info["font_size"]

            if placeholder_start > 0:
                run = paragraph.add_run(full_text[:placeholder_start])
                if char_formats:
                    _apply_format(run, char_formats[placeholder_start - 1])

            run = paragraph.add_run(first)
            if placeholder_start < len(char_formats):
                _apply_format(run, char_formats[placeholder_start])

            text_after = full_text[placeholder_start + len(placeholder) :]
            if text_after:
                run = paragraph.add_run(text_after)
                after_idx = placeholder_start + len(placeholder)
                if after_idx < len(char_formats):
                    _apply_format(run, char_formats[after_idx])

        # Extra chunks → sibling paragraphs (copy pPr so justify/style match)
        if rest:
            from copy import deepcopy
            from docx.oxml import OxmlElement

            anchor = paragraph._p
            pPr = paragraph._p.find(qn("w:pPr"))
            for chunk in rest:
                new_p = OxmlElement("w:p")
                if pPr is not None:
                    new_p.append(deepcopy(pPr))
                new_r = OxmlElement("w:r")
                # Copy first run props when available
                if paragraph.runs:
                    src_rPr = paragraph.runs[0]._r.find(qn("w:rPr"))
                    if src_rPr is not None:
                        new_r.append(deepcopy(src_rPr))
                new_t = OxmlElement("w:t")
                if chunk.startswith(" ") or chunk.endswith(" "):
                    new_t.set(qn("xml:space"), "preserve")
                new_t.text = chunk
                new_r.append(new_t)
                new_p.append(new_r)
                anchor.addnext(new_p)
                anchor = new_p

    def _fix_justified_soft_breaks(self, doc: Document) -> None:
        """Turn soft line breaks in justified paragraphs into real paragraphs.

        Safety net for any path that still injects w:br into jc=both text.
        """
        from copy import deepcopy
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fixed = 0
        targets = []
        for p in list(doc.paragraphs):
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                continue
            jc = pPr.find(qn("w:jc"))
            if jc is None or jc.get(qn("w:val")) not in ("both", "distribute"):
                continue
            soft = [
                br
                for br in p._p.iter(qn("w:br"))
                if br.get(qn("w:type")) in (None, "textWrapping")
            ]
            if soft:
                targets.append(p)

        for paragraph in targets:
            text = paragraph.text or ""
            if "\n" not in text:
                for br in list(paragraph._p.iter(qn("w:br"))):
                    if br.get(qn("w:type")) != "page":
                        parent = br.getparent()
                        if parent is not None:
                            parent.remove(br)
                continue

            chunks = [c.strip() for c in re.split(r"\n+", text) if c.strip()]
            if len(chunks) <= 1:
                for br in list(paragraph._p.iter(qn("w:br"))):
                    if br.get(qn("w:type")) != "page":
                        parent = br.getparent()
                        if parent is not None:
                            parent.remove(br)
                continue

            pPr = paragraph._p.find(qn("w:pPr"))
            src_rPr = None
            if paragraph.runs:
                src_rPr = paragraph.runs[0]._r.find(qn("w:rPr"))

            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)

            def _make_run(content: str):
                new_r = OxmlElement("w:r")
                if src_rPr is not None:
                    new_r.append(deepcopy(src_rPr))
                new_t = OxmlElement("w:t")
                if content.startswith(" ") or content.endswith(" "):
                    new_t.set(qn("xml:space"), "preserve")
                new_t.text = content
                new_r.append(new_t)
                return new_r

            paragraph._p.append(_make_run(chunks[0]))
            anchor = paragraph._p
            for chunk in chunks[1:]:
                new_p = OxmlElement("w:p")
                if pPr is not None:
                    new_p.append(deepcopy(pPr))
                new_p.append(_make_run(chunk))
                anchor.addnext(new_p)
                anchor = new_p
            fixed += 1

        if fixed:
            logger.info("Fixed soft breaks in %d justified paragraph(s)", fixed)

    def _replace_text_in_document(self, doc: Document, replacements: Dict[str, str]):
        """Replace text in all parts of the document while preserving formatting"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    self._replace_text_in_runs(paragraph, placeholder, replacement)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                self._replace_text_in_runs(paragraph, placeholder, replacement)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            self._replace_text_in_runs(paragraph, placeholder, replacement)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            self._replace_text_in_runs(paragraph, placeholder, replacement)

    def process_single_property(self, address: str, color_theme: Optional[str] = None, **kwargs) -> str:
        """
        Complete workflow for single property
        """
        property_data = self.create_property_report(address, **kwargs)
        document_path = self.create_word_document(property_data, color_theme=color_theme)
        return document_path

    def process_csv_batch(self, csv_path: str) -> List[str]:
        """
        Process multiple properties from CSV file
        CSV should have columns: address, and optionally other property details
        """
        logger.info(f"Processing batch from CSV: {csv_path}")
        
        df = pd.read_csv(csv_path)
        if 'address' not in df.columns:
            raise ValueError("CSV must contain 'address' column")
        
        document_paths = []
        for index, row in df.iterrows():
            try:
                address = row['address']
                logger.info(f"Processing {index + 1}/{len(df)}: {address}")
                
                # Extract other parameters from CSV
                kwargs = {col: str(row[col]) for col in df.columns if col != 'address' and pd.notna(row[col])}
                
                document_path = self.process_single_property(address, **kwargs)
                document_paths.append(document_path)
                
            except Exception as e:
                logger.error(f"Failed to process {address}: {e}")
                continue
        
        return document_paths


def main():
    """
    Example usage with market analysis
    """
    from config import get_openai_api_key, get_google_api_key, get_esri_api_key
    OPENAI_API_KEY = get_openai_api_key()
    GOOGLE_API_KEY = get_google_api_key()
    ESRI_API_KEY = get_esri_api_key()
    TEMPLATE_PATH = "template.docx"
    
    # Initialize generator
    generator = ComprehensivePropertyReportGenerator(
        openai_api_key=OPENAI_API_KEY,
        template_path=TEMPLATE_PATH,
        google_api_key=GOOGLE_API_KEY,
        esri_api_key=ESRI_API_KEY,
        output_dir="property_reports"
    )
    
    # Example usage
    address = "501 N 730 W American Fork Ut 84003"
    
    try:
        document_path = generator.process_single_property(
            address=address,
            prepared_by="Brayden Fisher",
            prepared_by_company="Colliers International",
            prepared_by_address="123 North 123 West Orem, UT 12345",
            prepared_for="Austin Shouse",
            prepared_for_company="UCCU Bank",
            prepared_for_address="789 yellow street Provo, UT 12345",
            property_name="Boothes House",
            property_type="Office",  # Changed to Office to match market analysis example
            lot_area="708711",
            acres="16",
            recorded_sale_date="1/24/2011",
            zoning="RA-5",
            apn="30-034-0073",
            current_owner="EKN FAMILY INVESTMENTS LLC"
        )
        
        print(f"Generated comprehensive report with market analysis: {document_path}")
        
    except Exception as e:
        print(f"Error generating report: {e}")


if __name__ == "__main__":
    main()