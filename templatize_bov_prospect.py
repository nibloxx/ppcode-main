"""Convert 'New Template (BOV) - Short Form.docx' into bov_prospect_template.docx.

Preserves the 5-7 page Prospect layout and heals split {{placeholders}} so
both4.py can fill cover textboxes correctly.

Run:  python templatize_bov_prospect.py
"""
from docx import Document
from docx.oxml.ns import qn

SOURCE = "New Template (BOV) - Short Form.docx"
OUTPUT = "bov_prospect_template.docx"

W_T = qn("w:t")
W_P = qn("w:p")
W_R = qn("w:r")


def set_runs_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def set_cell(cell, text):
    paragraph = cell.paragraphs[0]
    set_runs_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def set_paragraph_nodes_text(p, text: str) -> None:
    """Write `text` into the first w:t of paragraph XML; clear the rest."""
    nodes = []
    for r in p.findall(W_R):
        nodes.extend(r.findall(W_T))
    if not nodes:
        from docx.oxml import OxmlElement

        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        return
    nodes[0].text = text
    if text.startswith(" ") or text.endswith(" "):
        nodes[0].set(qn("xml:space"), "preserve")
    for n in nodes[1:]:
        n.text = ""


def heal_textbox_placeholders(doc) -> int:
    """Collapse split placeholders like {{ / p / roperty_name / }} into one run."""
    fixed = 0
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p in txbx.findall(W_P):
            nodes = []
            for r in p.findall(W_R):
                nodes.extend(list(r.findall(W_T)))
            if not nodes:
                continue
            joined = "".join(n.text or "" for n in nodes)
            # Normalize known broken / sample forms
            replacements = [
                ("{{Insert Date}}", "{{Date}}"),
                ("{{Insert Property Name}}", "{{property_name}}"),
                ("{{InsertDate}}", "{{Date}}"),
                ("{{InsertPropertyName}}", "{{property_name}}"),
            ]
            new_text = joined
            for old, new in replacements:
                if old in new_text:
                    new_text = new_text.replace(old, new)

            # If this paragraph only has Date label leftovers after a prior wipe,
            # restore Date + property name layout for header bars.
            stripped = new_text.strip()
            if stripped in ("{{Date}}", "Date: {{Date}}"):
                # Keep Date on this paragraph; property name is usually next para
                new_text = "Date:  {{Date}}" if stripped == "{{Date}}" else stripped

            # Heal any joinable {{...}} that was split across runs
            if "{{" in joined and new_text != joined:
                set_paragraph_nodes_text(p, new_text)
                fixed += 1
            elif "{{" in joined and any(
                (n.text or "") in ("{{", "}}", "{", "}") or (n.text or "").startswith("prepared_")
                or (n.text or "").startswith("roperty")
                or (n.text or "") == "p"
                for n in nodes
            ):
                # Force collapse split placeholder into first run
                set_paragraph_nodes_text(p, joined)
                fixed += 1
            elif new_text != joined:
                set_paragraph_nodes_text(p, new_text)
                fixed += 1

            # Ensure PREPARED labels keep a trailing space
            texts = list(p.iter(W_T))
            if texts and texts[0].text:
                t0 = texts[0].text
                if t0.strip() in ("PREPARED BY:", "PREPARED FOR:") and not t0.endswith(" "):
                    texts[0].text = t0.strip() + " "
                    texts[0].set(qn("xml:space"), "preserve")
                    fixed += 1
    return fixed


def heal_header_date_property_pairs(doc) -> int:
    """Header bars that lost property_name after templatizing Insert* tokens."""
    fixed = 0
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        paras = list(txbx.findall(W_P))
        if len(paras) < 2:
            continue
        p0 = "".join(t.text or "" for t in paras[0].iter(W_T)).strip()
        p1 = "".join(t.text or "" for t in paras[1].iter(W_T)).strip()
        if p0.startswith("Date:") and (not p1 or p1 in ("—", "{{Date}}")):
            set_paragraph_nodes_text(paras[0], "Date:  {{Date}}")
            set_paragraph_nodes_text(paras[1], "{{property_name}}")
            fixed += 1
        elif p0 == "{{Date}}" and not p1:
            set_paragraph_nodes_text(paras[0], "Date:  {{Date}}")
            set_paragraph_nodes_text(paras[1], "{{property_name}}")
            fixed += 1
    return fixed


TABLE_MAP = {
    (1, 1, 1): "{{market_price_psf}}/PSF X {{market_building_sf}}",
    (1, 1, 2): "{{market_value}}",
    (1, 2, 2): "{{market_value_rounded}}",
    (1, 5, 0): "{{value_aggressive}}",
    (1, 5, 1): "{{market_value_rounded}}",
    (1, 5, 2): "{{value_conservative}}",
    (2, 1, 1): "{{property_name}}",
    (2, 2, 1): "{{property_type}}",
    (2, 3, 1): "{{state}}",
    (2, 4, 1): "{{county}}",
    (2, 5, 1): "{{longitude}}",
    (2, 6, 1): "{{latitude}}",
    (2, 8, 1): "{{Topography}}",
    (2, 9, 1): "{{shape}}",
    (2, 10, 1): "{{Access}}",
    (2, 11, 1): "{{Exposure}}",
    (2, 13, 1): "{{lot_area}}",
    (2, 14, 1): "{{acres}}",
    (2, 15, 1): "{{recorded_sale_date}}",
    (2, 16, 1): "{{zoning}}",
    (2, 17, 1): "{{apn}}",
    (2, 18, 1): "{{current_owner}}",
    (3, 1, 1): "{{market_value_rounded}}",
    (3, 1, 2): "{{reconciliation_notes}}",
    (3, 2, 1): "{{market_value_rounded}}",
    (3, 2, 2): "{{reconciliation_notes}}",
    (4, 1, 1): "{{market_price_psf}}/SF X {{market_building_sf}}",
    (4, 1, 2): "{{market_value}}",
    (4, 2, 2): "{{market_value_rounded}}",
    (4, 5, 0): "{{value_aggressive}}",
    (4, 5, 1): "{{market_value_rounded}}",
    (4, 5, 2): "{{value_conservative}}",
}

# Labels kept from short-form design; order matches body sections so live
# page-number refresh produces sensible values down the TOC.
# Static pages are fallbacks only — both4 refreshes them via Word COM.
# "Properties for Sale" intentionally omitted (not a Prospect short-form section).
TOC_ROWS = [
    ("Executive Summary", "3"),
    ("Demographics", "3"),
    ("Subject Photos", "4"),
    ("Comparables", "5"),
    ("Certification", "7"),
]

NARRATIVE_PREFIX = [
    ("An executive summary provides a concise overview", "{{executive_summary}}"),
    ("{{The property located at", "{{property_summary}}"),
    ("{{225 State Hwy 121 sits", "{{location_summary}}"),
    ("{{The sales comparables analyzed", "{{reconciliation_summary}}"),
    ("The property located at", "{{property_summary}}"),
    ("225 State Hwy 121 sits", "{{location_summary}}"),
    ("The sales comparables analyzed", "{{reconciliation_summary}}"),
]


def replace_narratives(doc):
    seen = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for prefix, placeholder in NARRATIVE_PREFIX:
            if text.startswith(prefix):
                if placeholder in seen:
                    set_runs_text(paragraph, "")
                else:
                    set_runs_text(paragraph, placeholder)
                    seen.add(placeholder)
                break


def fix_general_info_braces(doc):
    brace_map = {
        "{Steele Oral Surgery Center}": "{{property_name}}",
        "{Retail Property}": "{{property_type}}",
        "{Texas}": "{{state}}",
        "{Dallas County}": "{{county}}",
        "{-96.9915956}": "{{longitude}}",
        "{32.9849477}": "{{latitude}}",
        "{Level at Street Grade}": "{{Topography}}",
        "{Irregular}": "{{shape}}",
        "{Good}": "{{Access}}",
        "{Average/Good}": "{{Exposure}}",
        "{87991}": "{{lot_area}}",
        "{2.02}": "{{acres}}",
        "{1/24/2010}": "{{recorded_sale_date}}",
        "{Z104}": "{{zoning}}",
        "{18-90123-00A-001}": "{{apn}}",
        "{SWSH LLC}": "{{current_owner}}",
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw = cell.text.strip()
                if raw in brace_map:
                    set_cell(cell, brace_map[raw])


def fix_toc_table(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.tables[0]
    # Match Client TOC proportions + generous row height
    col_widths = ("3200", "5977", "1078")
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "10255")
        tblW.set(qn("w:type"), "dxa")
    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gc in list(tblGrid.findall(qn("w:gridCol"))):
            tblGrid.remove(gc)
        for width in col_widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), width)
            tblGrid.append(gc)
    for ri, (label, page) in enumerate(TOC_ROWS):
        try:
            row = table.rows[ri]
            trPr = row._tr.get_or_add_trPr()
            trH = trPr.find(qn("w:trHeight"))
            if trH is None:
                trH = OxmlElement("w:trHeight")
                trPr.append(trH)
            trH.set(qn("w:val"), "1440")
            trH.set(qn("w:hRule"), "atLeast")
            for ci, width in enumerate(col_widths):
                if ci >= len(row.cells):
                    break
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), width)
                tcW.set(qn("w:type"), "dxa")
                if ci in (0, 1) and tcPr.find(qn("w:noWrap")) is None:
                    tcPr.append(OxmlElement("w:noWrap"))
            set_cell(row.cells[0], label)
            set_cell(row.cells[1], "_" * 54)
            set_cell(row.cells[2], page)
        except IndexError:
            pass

    # Drop leftover rows (e.g. old "Properties for Sale") beyond Prospect TOC
    while len(table.rows) > len(TOC_ROWS):
        parent = table.rows[-1]._tr.getparent()
        if parent is None:
            break
        parent.remove(table.rows[-1]._tr)


def tag_image_placeholders(doc) -> int:
    """Mark short-form pictures so both4 can swap aerial / street-view bytes.

    Client template already uses {{main_img}} / {{aerial_image}} / {{Subject_photo}}
    in picture alt-text. The short-form source uses sample images without those
    tags — including tiny broken placeholder PNGs that should be removed.
    """
    from docx.oxml.ns import qn

    emu_per_inch = 914400.0
    tagged = 0
    seen_rids = {}
    removed_tiny = 0
    for drawing in doc.element.body.iter(qn("w:drawing")):
        docPr = next(drawing.iter(qn("wp:docPr")), None)
        name = (docPr.get("name") if docPr is not None else "") or ""
        blip = next(drawing.iter(qn("a:blip")), None)
        if blip is None:
            continue
        rid = blip.get(qn("r:embed"))
        if not rid or rid not in doc.part.related_parts:
            continue
        part = doc.part.related_parts[rid]
        size = len(part.blob)
        ext = next(drawing.iter(qn("wp:extent")), None)
        width_in = height_in = None
        if ext is not None and ext.get("cx") and ext.get("cy"):
            width_in = int(ext.get("cx")) / emu_per_inch
            height_in = int(ext.get("cy")) / emu_per_inch
        descr_nodes = [
            el
            for tag in (qn("wp:docPr"), qn("pic:cNvPr"))
            for el in drawing.iter(tag)
        ]
        current = " ".join((el.get("descr") or "") for el in descr_nodes).lower()

        placeholder = None
        if "{{main_img}}" in current or "{{aerial_image}}" in current or "{{subject_photo}}" in current.lower():
            continue
        if "screenshot of" in current:
            # Source doc includes decorative sample screenshots that are not part
            # of the final BOV layout; remove them so they never surface as icons.
            anchor = next(drawing.iter(qn("wp:anchor")), None)
            inline = next(drawing.iter(qn("wp:inline")), None)
            holder = anchor if anchor is not None else inline
            parent = holder.getparent() if holder is not None else drawing.getparent()
            target = holder if holder is not None else drawing
            if parent is not None and target is not None:
                parent.remove(target)
                removed_tiny += 1
            continue
        if (
            "Picture" in name
            and width_in is not None
            and height_in is not None
            and width_in <= 1.6
            and height_in <= 1.6
        ):
            # Tiny duplicate picture layers render as broken image icons in Word.
            # Keep the large textbox slot, remove the tiny overlay picture.
            anchor = next(drawing.iter(qn("wp:anchor")), None)
            inline = next(drawing.iter(qn("wp:inline")), None)
            holder = anchor if anchor is not None else inline
            parent = holder.getparent() if holder is not None else drawing.getparent()
            target = holder if holder is not None else drawing
            if parent is not None and target is not None:
                parent.remove(target)
                removed_tiny += 1
            continue
        if size < 5000 and "Text Box" in name:
            # Large cover image slot in the source short-form uses a tiny broken
            # sample blob but the correct textbox geometry.
            placeholder = "{{main_img}}"
        elif "aerial" in current:
            placeholder = "{{aerial_image}}"
        elif (
            "building" in current
            or "cars parked" in current
            or "home address" in current
            or "street" in current
        ):
            placeholder = "{{Subject_photo}}"
        elif rid in seen_rids:
            placeholder = seen_rids[rid]

        if not placeholder:
            continue
        seen_rids[rid] = placeholder
        for el in descr_nodes:
            el.set("descr", placeholder)
        tagged += 1
    if removed_tiny:
        print(f"  removed_tiny_placeholders={removed_tiny}")
    return tagged


def _cover_classify(drawing):
    texts = "".join(t.text or "" for t in drawing.iter(W_T))
    descr_parts = []
    for tag in (qn("wp:docPr"), qn("pic:cNvPr")):
        for el in drawing.iter(tag):
            if el.get("descr"):
                descr_parts.append(el.get("descr"))
    descr = " ".join(descr_parts)
    docPr = next(drawing.iter(qn("wp:docPr")), None)
    name = (docPr.get("name") if docPr is not None else "") or ""

    if "{{address}}" in texts:
        return "address"
    if "PREPARED BY" in texts.upper():
        return "prep_by"
    if "PREPARED FOR" in texts.upper():
        return "prep_for"
    if "{{main_img}}" in descr:
        if "Picture" in name or next(drawing.iter(qn("a:blip")), None) is not None:
            if "Text Box" not in name:
                return "main_img_pic"
        return "main_img_box"
    return None


def _cover_get_offsets(drawing):
    ox = oy = cx = cy = None
    posH = next(drawing.iter(qn("wp:positionH")), None)
    posV = next(drawing.iter(qn("wp:positionV")), None)
    if posH is not None:
        off = posH.find(qn("wp:posOffset"))
        if off is not None and off.text:
            ox = off.text
    if posV is not None:
        off = posV.find(qn("wp:posOffset"))
        if off is not None and off.text:
            oy = off.text
    ext = next(drawing.iter(qn("wp:extent")), None)
    if ext is not None:
        cx, cy = ext.get("cx"), ext.get("cy")
    return ox, oy, cx, cy


def _cover_set_offsets(drawing, ox=None, oy=None, cx=None, cy=None):
    wps_sppr = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}spPr"
    a_xfrm = "{http://schemas.openxmlformats.org/drawingml/2006/main}xfrm"
    a_ext = "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"
    pic_sppr = "{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr"

    if ox is not None:
        posH = next(drawing.iter(qn("wp:positionH")), None)
        if posH is not None:
            off = posH.find(qn("wp:posOffset"))
            if off is not None:
                off.text = str(ox)
    if oy is not None:
        posV = next(drawing.iter(qn("wp:positionV")), None)
        if posV is not None:
            off = posV.find(qn("wp:posOffset"))
            if off is not None:
                off.text = str(oy)
    if cx is None and cy is None:
        return
    for ext in drawing.iter(qn("wp:extent")):
        if cx is not None:
            ext.set("cx", str(cx))
        if cy is not None:
            ext.set("cy", str(cy))
    for spPr in list(drawing.iter(wps_sppr)) + list(drawing.iter(pic_sppr)):
        xfrm = spPr.find(a_xfrm)
        if xfrm is None:
            continue
        aext = xfrm.find(a_ext)
        if aext is None:
            continue
        if cx is not None:
            aext.set("cx", str(cx))
        if cy is not None:
            aext.set("cy", str(cy))


# Extra cover polish after aligning to Client (review comments on Prospect cover).
COVER_NUDGE_DOWN_EMU = int(0.28 * 914400)  # ~0.28" — address + prep boxes


def _cover_nudge_down(drawing, delta_emu: int = COVER_NUDGE_DOWN_EMU) -> bool:
    ox, oy, cx, cy = _cover_get_offsets(drawing)
    if oy is None:
        return False
    _cover_set_offsets(drawing, oy=str(int(oy) + delta_emu))
    return True


def _cover_remove_shape_border(drawing) -> bool:
    """Remove light-grey line around the cover aerial textbox/image."""
    from lxml import etree

    wps_sppr = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}spPr"
    pic_sppr = "{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr"
    a_ln = "{http://schemas.openxmlformats.org/drawingml/2006/main}ln"
    a_nofill = "{http://schemas.openxmlformats.org/drawingml/2006/main}noFill"
    a_srgb = "{http://schemas.openxmlformats.org/drawingml/2006/main}srgbClr"
    a_alpha = "{http://schemas.openxmlformats.org/drawingml/2006/main}alpha"
    wps_style = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}style"
    a_ln_ref = "{http://schemas.openxmlformats.org/drawingml/2006/main}lnRef"

    changed = False
    for spPr in list(drawing.iter(wps_sppr)) + list(drawing.iter(pic_sppr)):
        ln = spPr.find(a_ln)
        if ln is not None:
            ln.clear()
            ln.set("w", "0")
            etree.SubElement(ln, a_nofill)
            changed = True
    for style in drawing.iter(wps_style):
        ln_ref = style.find(a_ln_ref)
        if ln_ref is not None:
            for child in list(ln_ref):
                ln_ref.remove(child)
            srgb = etree.SubElement(ln_ref, a_srgb)
            srgb.set("val", "FFFFFF")
            alpha = etree.SubElement(srgb, a_alpha)
            alpha.set("val", "0")
            changed = True
    for ee in drawing.iter(qn("wp:effectExtent")):
        for side in ("l", "t", "r", "b"):
            if ee.get(side) not in (None, "0"):
                ee.set(side, "0")
                changed = True
    return changed


def align_cover_to_client(doc, client_path: str = "bov_template.docx") -> int:
    """Match Client cover: address under slant, prep blocks, full-size cover image.

    Short-form source keeps address lower (on the hero) with a tiny {{main_img}}
    square. Align positions/sizes to bov_template.docx and strip image borders.
    """
    client = Document(client_path)
    buckets_c, buckets_p = {}, {}
    for d in client.element.body.iter(qn("w:drawing")):
        key = _cover_classify(d)
        if key:
            buckets_c.setdefault(key, []).append(d)
    for d in doc.element.body.iter(qn("w:drawing")):
        key = _cover_classify(d)
        if key:
            buckets_p.setdefault(key, []).append(d)

    changed = 0
    for key in ("address", "prep_by", "prep_for", "main_img_box"):
        if key not in buckets_c or key not in buckets_p:
            continue
        ox, oy, cx, cy = _cover_get_offsets(buckets_c[key][0])
        for drawing in buckets_p[key]:
            _cover_set_offsets(drawing, ox=ox, oy=oy, cx=cx, cy=cy)
            changed += 1

    if "main_img_pic" in buckets_c and "main_img_pic" in buckets_p:
        _, _, cx, cy = _cover_get_offsets(buckets_c["main_img_pic"][0])
        for drawing in buckets_p["main_img_pic"]:
            _cover_set_offsets(drawing, cx=cx, cy=cy)
            changed += 1

    # Strip light-grey borders around cover / aerial / subject images
    for d in doc.element.body.iter(qn("w:drawing")):
        descr = " ".join(
            (el.get("descr") or "")
            for tag in (qn("wp:docPr"), qn("pic:cNvPr"))
            for el in d.iter(tag)
        ).lower()
        if any(m in descr for m in ("main_img", "aerial_image", "subject_photo")):
            if _cover_remove_shape_border(d):
                changed += 1

    # Match Client cover positions exactly (no extra nudge)
    return changed


def ensure_header_footer_on_body(doc) -> bool:
    """Ensure body sectPr links header/footer (blue bars + Page X of Y).

    Short-form stores those refs on a mid-doc sectPr in the comps zone; if that
    paragraph is cleared, bars/page numbers vanish. Mirror Client by also
    linking them on the body sectPr.
    """
    from lxml import etree

    # Find header/footer rIds from document relationships
    header_rid = footer_rid = None
    for rel in doc.part.rels.values():
        target = str(getattr(rel, "target_ref", "") or "").lower()
        if "header" in target and header_rid is None:
            header_rid = rel.rId
        if "footer" in target and footer_rid is None:
            footer_rid = rel.rId
    if not header_rid or not footer_rid:
        return False

    sect = doc.element.body.find(qn("w:sectPr"))
    if sect is None:
        return False

    w_header = qn("w:headerReference")
    w_footer = qn("w:footerReference")
    for tag in (w_header, w_footer):
        for el in list(sect.findall(tag)):
            sect.remove(el)

    href = etree.Element(w_header)
    href.set(qn("w:type"), "default")
    href.set(qn("r:id"), header_rid)
    fref = etree.Element(w_footer)
    fref.set(qn("w:type"), "default")
    fref.set(qn("r:id"), footer_rid)
    sect.insert(0, fref)
    sect.insert(0, href)
    return True


def compress_embedded_images(doc, max_edge: int = 1200, quality: int = 70) -> int:
    """Shrink large template sample images so outputs are not multi‑MB by default."""
    from io import BytesIO

    try:
        from PIL import Image
    except ImportError:
        return 0

    compressed = 0
    seen = set()
    for rid, part in list(doc.part.related_parts.items()):
        if rid in seen:
            continue
        name = str(getattr(part, "partname", "") or "").lower()
        ctype = str(getattr(part, "content_type", "") or "").lower()
        if "image" not in ctype and not name.endswith((".png", ".jpg", ".jpeg")):
            continue
        blob = part.blob
        if len(blob) < 80_000:
            continue
        try:
            img = Image.open(BytesIO(blob)).convert("RGB")
            w, h = img.size
            scale = min(1.0, float(max_edge) / float(max(w, h)))
            if scale < 1.0:
                img = img.resize(
                    (max(1, int(w * scale)), max(1, int(h * scale))),
                    Image.LANCZOS,
                )
            buf = BytesIO()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            part._blob = buf.getvalue()
            seen.add(rid)
            compressed += 1
        except Exception:
            continue
    return compressed


def rebuild_sale_opinion_block(doc: Document) -> int:
    """Normalize OPINIONS / SALE OPINION blocks.

    - Merge & center Market Sale Price formula ($/SF X SF) into one cell
    - Keep Aggressive | Market Value | Conservative as a clean 3-col grid
    """
    from lxml import etree

    def make_tc(text, span, fill=None, bold=False, color=None, sz="20"):
        tc = etree.Element(qn("w:tc"))
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        etree.SubElement(tcPr, qn("w:tcW"), {
            qn("w:w"): str(1200 * span),
            qn("w:type"): "dxa",
        })
        if span > 1:
            etree.SubElement(tcPr, qn("w:gridSpan"), {qn("w:val"): str(span)})
        borders = etree.SubElement(tcPr, qn("w:tcBorders"))
        for edge in ("top", "left", "bottom", "right"):
            etree.SubElement(borders, qn(f"w:{edge}"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): "000000",
            })
        if fill:
            etree.SubElement(tcPr, qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): fill,
            })
        etree.SubElement(tcPr, qn("w:vAlign"), {qn("w:val"): "center"})
        p = etree.SubElement(tc, qn("w:p"))
        pPr = etree.SubElement(p, qn("w:pPr"))
        etree.SubElement(pPr, qn("w:jc"), {qn("w:val"): "center"})
        # Match VALUATION METHOD / SUMMARY tables (10pt)
        rPr_p = etree.SubElement(pPr, qn("w:rPr"))
        etree.SubElement(rPr_p, qn("w:sz"), {qn("w:val"): sz})
        etree.SubElement(rPr_p, qn("w:szCs"), {qn("w:val"): sz})
        r = etree.SubElement(p, qn("w:r"))
        rPr = etree.SubElement(r, qn("w:rPr"))
        if bold:
            etree.SubElement(rPr, qn("w:b"))
        if color:
            etree.SubElement(rPr, qn("w:color"), {qn("w:val"): color})
        etree.SubElement(rPr, qn("w:sz"), {qn("w:val"): sz})
        etree.SubElement(rPr, qn("w:szCs"), {qn("w:val"): sz})
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        return tc

    def make_tr(cells, height="288"):
        tr = etree.Element(qn("w:tr"))
        trPr = etree.SubElement(tr, qn("w:trPr"))
        etree.SubElement(trPr, qn("w:trHeight"), {
            qn("w:val"): height,
            qn("w:hRule"): "atLeast",
        })
        for c in cells:
            tr.append(c)
        return tr

    fixed = 0
    for table in doc.tables:
        flat = " ".join(c.text for r in table.rows for c in r.cells).upper()
        if "OPINIONS OF VALUE" not in flat or "MARKET SALE PRICE" not in flat:
            continue

        fill = "0070C0"
        for c in table.rows[0].cells:
            shd = c._tc.find(qn("w:tcPr"))
            if shd is not None:
                s = shd.find(qn("w:shd"))
                if s is not None and s.get(qn("w:fill")):
                    fill = s.get(qn("w:fill"))
                    break

        # Prospect short-form uses an 8-col grid; Client opinions is 5-col.
        grid = table._tbl.find(qn("w:tblGrid"))
        n_cols = len(grid.findall(qn("w:gridCol"))) if grid is not None else 8
        if n_cols >= 8:
            label_span, mid_span, val_span = 2, 4, 2
            agg_span, mkt_span, con_span = 3, 3, 2
            total_span = 8
        else:
            label_span, mid_span, val_span = 1, 3, 1
            agg_span, mkt_span, con_span = 1, 1, 1
            total_span = n_cols

        price_row = make_tr([
            make_tc("Market Sale Price", label_span, bold=True),
            make_tc("{{market_price_psf}}/SF X {{market_building_sf}}", mid_span),
            make_tc("{{market_value}}", val_span),
        ])
        rounded_row = make_tr([
            make_tc("Market Sales Price (Rounded)", label_span, bold=True),
            make_tc("", mid_span),
            make_tc("{{market_value_rounded}}", val_span),
        ])

        tbl = table._tbl
        trs = list(tbl.findall(qn("w:tr")))
        for idx, new_tr in ((1, price_row), (2, rounded_row)):
            if idx < len(trs):
                trs[idx].getparent().replace(trs[idx], new_tr)

        # SALE OPINION block only on Prospect (6+ rows)
        if len(table.rows) >= 6 and "SALE OPINION OF VALUE" in flat:
            trs = list(tbl.findall(qn("w:tr")))
            header = make_tr([
                make_tc(
                    "SALE OPINION OF VALUE",
                    total_span,
                    fill=fill,
                    bold=True,
                    color="FFFFFF",
                )
            ])
            labels = make_tr([
                make_tc("Aggressive", agg_span, bold=True),
                make_tc("Market Value", mkt_span, bold=True),
                make_tc("Conservative", con_span, bold=True),
            ])
            values = make_tr([
                make_tc("{{value_aggressive}}", agg_span),
                make_tc("{{market_value_rounded}}", mkt_span),
                make_tc("{{value_conservative}}", con_span),
            ])
            for idx, new_tr in ((3, header), (4, labels), (5, values)):
                if idx < len(trs):
                    trs[idx].getparent().replace(trs[idx], new_tr)

        fixed += 1
    return fixed


def main():
    doc = Document(SOURCE)
    healed = heal_textbox_placeholders(doc)
    headers = heal_header_date_property_pairs(doc)
    fix_general_info_braces(doc)
    img_tags = tag_image_placeholders(doc)
    img_compressed = compress_embedded_images(doc)

    cells = 0
    for (ti, ri, ci), placeholder in TABLE_MAP.items():
        try:
            set_cell(doc.tables[ti].rows[ri].cells[ci], placeholder)
            cells += 1
        except IndexError:
            print(f"  WARN: cell {(ti, ri, ci)} out of range")

    opinion_fixed = rebuild_sale_opinion_block(doc)
    fix_toc_table(doc)
    replace_narratives(doc)

    # Second heal pass after other edits
    healed += heal_textbox_placeholders(doc)
    headers += heal_header_date_property_pairs(doc)

    cover_aligned = align_cover_to_client(doc)
    hf_ok = ensure_header_footer_on_body(doc)

    doc.save(OUTPUT)
    print(
        f"Wrote {OUTPUT}: healed={healed}, header_pairs={headers}, "
        f"image_tags={img_tags}, images_compressed={img_compressed}, "
        f"cover_aligned={cover_aligned}, header_footer={hf_ok}, "
        f"cells={cells}, opinion_fixed={opinion_fixed}, tables={len(doc.tables)}, "
        f"paragraphs={len(doc.paragraphs)}"
    )


if __name__ == "__main__":
    main()
