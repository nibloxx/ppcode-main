"""Patch bov_template.docx: replace remaining static sample text with placeholders
and normalize body placeholder styling so filled AI text doesn't inherit blue/italic.
"""
from docx import Document
from docx.shared import RGBColor, Pt
from templatize_bov import set_runs_text, OUTPUT

# (paragraph_index, new_text)
PARA_MAP = {
    9: "{{executive_summary}}",
    24: "{{regional_analysis}}",
    53: "{{economic_factors}}",
    55: "{{community_services}}",
    63: "{{vacancy_rates}}",
    69: "{{lease_rates}}",
    71: "{{construction_activity}}",
    73: "{{market_trends}}",
    77: "{{investment_insights}}",
    80: "{{market_recommendations}}",
    85: "{{market_data_sources}}",
    100: "{{reconciliation_summary}}",
    104: "{{sales_conclusion}}",
}

# Duplicate bullets / old static blocks / leftover empty image paras in comps
CLEAR_PARAS = {64, 65, 66, 67, 74, 75, 78, 81, 82, 84, 86, 87, 88, 89, 90, 93, 94, 95, 96, 97, 98}

BODY_PLACEHOLDERS = {
    "{{executive_summary}}",
    "{{regional_analysis}}",
    "{{economic_factors}}",
    "{{community_services}}",
    "{{vacancy_rates}}",
    "{{lease_rates}}",
    "{{construction_activity}}",
    "{{market_trends}}",
    "{{investment_insights}}",
    "{{market_recommendations}}",
    "{{market_data_sources}}",
    "{{reconciliation_summary}}",
    "{{sales_conclusion}}",
    "{{property_summary}}",
    "{{location_summary}}",
    "{{demographic_analysis}}",
    "{{size_and_topography}}",
    "{{population_analysis}}",
    "{{household_trends}}",
    "{{employment_analysis}}",
}


def _style_body_placeholder(paragraph):
    """Body placeholders should be black, non-italic so fills match the original body look."""
    for run in paragraph.runs:
        run.italic = False
        try:
            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        if run.font.size is None:
            run.font.size = Pt(11)


def fix_template(path: str = OUTPUT):
    doc = Document(path)

    # Discover economic / community indices dynamically if fixed indices drifted
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "Economic Factors" and i + 1 < len(doc.paragraphs):
            PARA_MAP[i + 1] = "{{economic_factors}}"
        if text == "Community Services" and i + 1 < len(doc.paragraphs):
            PARA_MAP[i + 1] = "{{community_services}}"

    for idx, text in PARA_MAP.items():
        if idx < len(doc.paragraphs):
            set_runs_text(doc.paragraphs[idx], text)
            if text in BODY_PLACEHOLDERS:
                _style_body_placeholder(doc.paragraphs[idx])

    for idx in CLEAR_PARAS:
        if idx < len(doc.paragraphs):
            set_runs_text(doc.paragraphs[idx], "")
            p = doc.paragraphs[idx]._element
            for drawing in list(p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")):
                drawing.getparent().remove(drawing)

    # Also normalize any remaining body placeholders found by text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in BODY_PLACEHOLDERS:
            _style_body_placeholder(paragraph)

    if len(doc.tables) > 9:
        t = doc.tables[9]
        if t.rows:
            set_runs_text(
                t.rows[0].cells[0].paragraphs[0],
                "EMPLOYMENT & UNEMPLOYMENT STATISTICS (Recent)",
            )

    doc.save(path)
    print(f"Patched {path}: {len(PARA_MAP)} placeholders, cleared {len(CLEAR_PARAS)} static/blank blocks")


if __name__ == "__main__":
    fix_template(OUTPUT)
    try:
        fix_template("bov_prospect_template.docx")
    except Exception as exc:
        print(f"Prospect template skip: {exc}")
